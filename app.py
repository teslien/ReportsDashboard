from flask import Flask, render_template, request, jsonify, make_response, has_request_context, redirect, url_for, send_file, session, abort
from werkzeug.local import LocalProxy
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from authlib.integrations.flask_client import OAuth
from functools import wraps
from collections import defaultdict
from dotenv import load_dotenv
import io
import re
import json
import hashlib
import os
import requests
import base64
import secrets
import mysql.connector
from mysql.connector import Error
from datetime import datetime, timedelta, timezone
from urllib.parse import unquote
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from xhtml2pdf import pisa
from pypdf import PdfWriter

load_dotenv()

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev_secret_key_change_in_production")

import time

# =========================
# SIMPLE IN-MEMORY CACHE
# =========================
JIRA_CACHE = {}
CACHE_TTL = 300 # 5 minutes default
CUSTOMER_CACHE = {"expires_at": 0, "values": []}
CUSTOMER_CACHE_TTL = 1800  # 30 minutes

def get_jira_cached(url, headers, params, ttl=CACHE_TTL, force_refresh=False):
    """
    Simple caching wrapper for Jira GET requests.
    Keys are based on URL + sorted params.
    """
    # Create a unique key for this request
    # Convert params to a sorted tuple of items to be hashable
    param_key = tuple(sorted(params.items())) if params else ()
    cache_key = (url, param_key)
    
    now = time.time()
    
    if not force_refresh and cache_key in JIRA_CACHE:
        timestamp, data = JIRA_CACHE[cache_key]
        if now - timestamp < ttl:
            print(f"DEBUG: Serving from cache: {url}")
            return data
            
    # Fetch fresh data
    try:
        resp = requests.get(url, headers=headers, params=params)
        if resp.status_code == 200:
            data = resp.json()
            JIRA_CACHE[cache_key] = (now, data)
            return data
        else:
            # If error, try to return stale cache if available, otherwise return error
            if cache_key in JIRA_CACHE:
                print(f"DEBUG: Fetch failed ({resp.status_code}), serving stale cache.")
                return JIRA_CACHE[cache_key][1]
            return resp.json() # Return the error response
    except Exception as e:
        if cache_key in JIRA_CACHE:
            print(f"DEBUG: Exception {e}, serving stale cache.")
            return JIRA_CACHE[cache_key][1]
        raise e

# =========================
# 🔐 AUTH & SECURITY SETUP
# =========================
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login_page"

oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=os.environ.get('GOOGLE_CLIENT_ID'),
    client_secret=os.environ.get('GOOGLE_CLIENT_SECRET'),
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={'scope': 'openid email profile'}
)

# User Class
class User(UserMixin):
    def __init__(self, id, email, name, role_id, role_name, permissions, api_token):
        self.id = id
        self.email = email
        self.name = name
        self.role_id = role_id
        self.role_name = role_name
        self.permissions = permissions
        self.api_token = api_token
        
    def has_permission(self, perm):
        if self.role_name == 'Admin': return True
        return self.permissions.get(perm, False)

    def can_view_page(self, page_key):
        if self.role_name == 'Admin': return True
        allowed = self.permissions.get('allowed_pages')
        # Payroll dashboard is opt-in per role (Admin always). Legacy roles with no
        # allowed_pages list otherwise see all pages; exclude this page until granted.
        if page_key == "payroll_dashboard":
            if not isinstance(allowed, list):
                return False
            return page_key in allowed
        # If 'allowed_pages' is not present, default to True (backward compatibility)
        if allowed is None: return True
        return page_key in allowed

    def to_dict(self):
        return {
            "id": self.id,
            "email": self.email,
            "name": self.name,
            "role_id": self.role_id,
            "role_name": self.role_name,
            "permissions": self.permissions
        }

@login_manager.user_loader
def load_user(user_id):
    conn, cursor = get_db_connection()
    if not conn: return None
    cursor.execute("""
        SELECT u.id, u.email, u.name, u.role_id, r.name, r.permissions, u.api_token
        FROM users u
        LEFT JOIN roles r ON u.role_id = r.id
        WHERE u.id = %s
    """, (user_id,))
    row = cursor.fetchone()
    conn.close()
    
    if row:
        perms = json.loads(row[5]) if row[5] else {}
        return User(id=row[0], email=row[1], name=row[2], role_id=row[3], role_name=row[4], permissions=perms, api_token=row[6])
    return None

def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or current_user.role_name != 'Admin':
            return abort(403)
        return f(*args, **kwargs)
    return decorated_function

def permission_required(permission):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not current_user.is_authenticated:
                return abort(403)
            # Admin has all permissions
            if current_user.role_name == 'Admin':
                return f(*args, **kwargs)
            # Check specific permission
            if current_user.has_permission(permission):
                return f(*args, **kwargs)
            return abort(403)
        return decorated_function
    return decorator

def page_permission_required(page_key):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not current_user.is_authenticated:
                return redirect(url_for('login_page'))
            if not current_user.can_view_page(page_key):
                return render_template("403.html"), 403
            return f(*args, **kwargs)
        return decorated_function
    return decorator


def _sprint_tracker_action_allowed(action="edit"):
    if not current_user.is_authenticated:
        return False
    if current_user.role_name == 'Admin':
        return True
    perms = current_user.permissions or {}
    legacy_edit = bool(perms.get("edit"))
    if action == "delete":
        return legacy_edit or bool(perms.get("delete_tickets"))
    if action == "create":
        return legacy_edit or bool(perms.get("create_tickets")) or bool(perms.get("edit_tickets"))
    return legacy_edit or bool(perms.get("edit_tickets"))


def _can_manage_sprint_tracker():
    return any(_sprint_tracker_action_allowed(action) for action in ("create", "edit", "delete"))


def sprint_tracker_write_required(action="edit"):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if not current_user.is_authenticated:
                return jsonify({"error": "Authentication required"}), 401
            if not current_user.can_view_page("sprint_tracker"):
                if str(request.path or "").startswith("/api/"):
                    return jsonify({"error": "You do not have access to sprint tracker."}), 403
                return render_template("403.html"), 403
            if _sprint_tracker_action_allowed(action):
                return f(*args, **kwargs)
            return jsonify({"error": "Read-only access. You can view the sprint tracker but cannot modify it."}), 403
        return decorated_function
    return decorator

# =========================
# 🔴 CONFIG (DYNAMIC)
# =========================

# Defaults
DEFAULT_JIRA_EMAIL = ""
DEFAULT_JIRA_API_TOKEN = ""
DEFAULT_PROJECT_KEY = ""
DEFAULT_JIRA_DOMAIN = "https://lumberfi.atlassian.net"

def _decode_value(value):
    if not value or value.lower() in ("null", "undefined"):
        return ""
    try:
        # Some values might be double quoted or URL encoded
        value = unquote(value).strip()
        # Remove surrounding quotes if present (some browsers/servers add them)
        if value.startswith('"') and value.endswith('"'):
            value = value[1:-1]
        return value.strip()
    except Exception:
        return value.strip()

def add_audit_log(page, item_key, field_name, new_value, old_value=None):
    """Adds a record to the audit_logs table."""
    conn, cursor = get_db_connection()
    if not conn:
        return False
    try:
        cursor.execute("""
            INSERT INTO audit_logs (user_id, user_name, page, item_key, field_name, old_value, new_value)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (current_user.id, current_user.name, page, item_key, field_name, str(old_value) if old_value is not None else None, str(new_value)))
        conn.commit()
        return True
    except Exception as e:
        print(f"Audit Log Error: {e}")
        return False
    finally:
        conn.close()

def _get_jira_config():
    """Fetch Jira config from DB. Returns (email, token, project_key, domain)."""
    try:
        conn, cursor = get_db_connection()
        if not conn: return "", "", "", DEFAULT_JIRA_DOMAIN
        cursor.execute("SELECT email, api_token, project_key, jira_domain FROM jira_config WHERE id = 1")
        row = cursor.fetchone()
        conn.close()
        if row:
            return row[0] or "", row[1] or "", row[2] or "", row[3] or DEFAULT_JIRA_DOMAIN
    except Exception:
        pass
    return "", "", "", DEFAULT_JIRA_DOMAIN

def _get_project_key():
    _, _, db_project, _ = _get_jira_config()
    project = db_project or DEFAULT_PROJECT_KEY
    if has_request_context():
        req_project = request.headers.get("X-Project-Key") or request.args.get("project_key") or request.cookies.get("project_key")
        if req_project:
            project = _decode_value(req_project)
    print(f"DEBUG: _get_project_key decoded: {repr(project)}")
    return project.upper() if project else ""

def _get_jira_domain():
    _, _, _, db_domain = _get_jira_config()
    domain = db_domain or DEFAULT_JIRA_DOMAIN
    if has_request_context():
        req_domain = request.headers.get("X-Jira-Domain") or request.args.get("jira_domain") or request.cookies.get("jira_domain")
        if req_domain:
            domain = _decode_value(req_domain)
    
    if domain and not domain.startswith("http"):
        domain = "https://" + domain
    return domain.rstrip("/")

def _get_jira_headers():
    db_email, db_token, _, _ = _get_jira_config()
    email = db_email or DEFAULT_JIRA_EMAIL
    token = db_token or DEFAULT_JIRA_API_TOKEN
    
    if has_request_context():
        req_email = request.headers.get("X-Jira-Email") or request.args.get("jira_email") or request.cookies.get("jira_email")
        req_token = request.headers.get("X-Jira-Token") or request.args.get("jira_token") or request.cookies.get("jira_token")
        if req_email: email = _decode_value(req_email)
        if req_token: token = _decode_value(req_token)
        
    if not email or not token:
        # Return empty auth if credentials are missing
        return {"Content-Type": "application/json"}
        
    auth_str = f"{email}:{token}"
    auth_b64 = base64.b64encode(auth_str.encode()).decode()
    
    return {
        "Authorization": f"Basic {auth_b64}",
        "Content-Type": "application/json"
    }


def _get_explicit_request_jira_headers():
    """Return Jira auth headers only when the request explicitly supplies user credentials."""
    if not has_request_context():
        return {"Content-Type": "application/json"}

    req_email = request.headers.get("X-Jira-Email") or request.args.get("jira_email") or request.cookies.get("jira_email")
    req_token = request.headers.get("X-Jira-Token") or request.args.get("jira_token") or request.cookies.get("jira_token")
    email = _decode_value(req_email).strip() if req_email else ""
    token = _decode_value(req_token).strip() if req_token else ""

    if not email or not token:
        return {"Content-Type": "application/json"}

    auth_str = f"{email}:{token}"
    auth_b64 = base64.b64encode(auth_str.encode()).decode()
    return {
        "Authorization": f"Basic {auth_b64}",
        "Content-Type": "application/json",
    }

# Proxies to allow dynamic access per request while keeping existing code working
PROJECT_KEY = LocalProxy(_get_project_key)
HEADERS = LocalProxy(_get_jira_headers)
JIRA_DOMAIN = LocalProxy(_get_jira_domain)

@app.context_processor
def inject_config():
    conn, cursor = get_db_connection(dictionary=True)
    public_trackers = []
    header_title = "Jira Analytics"
    if conn:
        try:
            cursor.execute("SELECT id, name FROM trackers_v2 WHERE is_public = TRUE")
            public_trackers = cursor.fetchall()
            
            cursor.execute("SELECT config_value FROM app_config WHERE config_key = 'header_title'")
            row = cursor.fetchone()
            if row:
                header_title = row['config_value']
        except:
            pass
        finally:
            conn.close()
            
    return dict(
        JIRA_DOMAIN=str(JIRA_DOMAIN),
        project=str(PROJECT_KEY),
        public_trackers=public_trackers,
        header_title=header_title
    )

# Legacy auth (deprecated but kept for compatibility if used directly)
AUTH = base64.b64encode(f"{DEFAULT_JIRA_EMAIL}:{DEFAULT_JIRA_API_TOKEN}".encode()).decode()





# =========================
# DATABASE SETUP
# =========================
MYSQL_CONFIG = {
    'host': 'localhost',
    'user': 'rohit',
    'password': 'Rohit',
    'database': 'rb_win'
}

def get_db_connection(dictionary=False):
    """Establish a connection to the MySQL database."""
    try:
        conn = mysql.connector.connect(**MYSQL_CONFIG)
        if dictionary:
            # Equivalent to sqlite3.Row
            cursor = conn.cursor(dictionary=True)
        else:
            cursor = conn.cursor()
        return conn, cursor
    except Error as e:
        print(f"Error connecting to MySQL: {e}")
        return None, None

def init_db():
    # First, try to connect without a database to create it if it doesn't exist
    try:
        conn = mysql.connector.connect(
            host=MYSQL_CONFIG['host'],
            user=MYSQL_CONFIG['user'],
            password=MYSQL_CONFIG['password']
        )
        cursor = conn.cursor()
        cursor.execute(f"CREATE DATABASE IF NOT EXISTS {MYSQL_CONFIG['database']}")
        conn.close()
    except Error as e:
        print(f"Error creating database: {e}")

    conn, cursor = get_db_connection()
    if not conn:
        return

    # Roles Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS roles (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(255) UNIQUE NOT NULL,
            permissions TEXT
        )
    ''')

    # Users Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INT AUTO_INCREMENT PRIMARY KEY,
            email VARCHAR(255) UNIQUE NOT NULL,
            name VARCHAR(255),
            google_id VARCHAR(255) UNIQUE,
            password_hash VARCHAR(255),
            role_id INT,
            api_token VARCHAR(255) UNIQUE,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (role_id) REFERENCES roles(id)
        )
    ''')

    # Seed Roles if empty
    cursor.execute("SELECT count(*) FROM roles")
    if cursor.fetchone()[0] == 0:
        cursor.execute("INSERT INTO roles (name, permissions) VALUES ('Admin', '{\"all\": true}')")
        cursor.execute("INSERT INTO roles (name, permissions) VALUES ('Editor', '{\"view\": true, \"edit\": true}')")
        cursor.execute("INSERT INTO roles (name, permissions) VALUES ('Viewer', '{\"view\": true, \"allowed_pages\": [\"dashboard\", \"sprint_tracker\"]}')")
    else:
        # Ensure Viewer keeps dashboard access and can open the sprint tracker in read-only mode.
        cursor.execute("SELECT permissions FROM roles WHERE name = 'Viewer'")
        row = cursor.fetchone()
        if row:
            perms = json.loads(row[0]) if row[0] else {}
            allowed = perms.get('allowed_pages')
            default_viewer_pages = ['dashboard', 'sprint_tracker']
            if isinstance(allowed, list):
                changed = False
                for page in default_viewer_pages:
                    if page not in allowed:
                        allowed.append(page)
                        changed = True
                if changed:
                    perms['allowed_pages'] = allowed
                    cursor.execute("UPDATE roles SET permissions = %s WHERE name = 'Viewer'", (json.dumps(perms),))
            else:
                perms['allowed_pages'] = default_viewer_pages
                cursor.execute("UPDATE roles SET permissions = %s WHERE name = 'Viewer'", (json.dumps(perms),))

        # Ensure Editor role has access to key pages by default. If the role has
        # an explicit allowed_pages list (added later via the admin UI), make sure
        # these pages are included so existing Editors aren't locked out.
        cursor.execute("SELECT permissions FROM roles WHERE name = 'Editor'")
        row = cursor.fetchone()
        if row:
            perms = json.loads(row[0]) if row[0] else {}
            allowed = perms.get('allowed_pages')
            default_editor_pages = ['sprint_tracker', 'customer_dashboard', 'sprint_dashboard', 'team_productivity', 'customer_closure']
            if isinstance(allowed, list):
                for page in default_editor_pages:
                    if page not in allowed:
                        allowed.append(page)
                # Payroll: Admin-only by default; remove from Editor if it was auto-added earlier.
                if "payroll_dashboard" in allowed:
                    allowed = [p for p in allowed if p != "payroll_dashboard"]
                perms["allowed_pages"] = allowed
                cursor.execute("UPDATE roles SET permissions = %s WHERE name = 'Editor'", (json.dumps(perms),))

    # Trackers Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS trackers (
            id INT AUTO_INCREMENT PRIMARY KEY,
            title TEXT NOT NULL,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    # Tickets Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS tracker_tickets (
            id INT AUTO_INCREMENT PRIMARY KEY,
            tracker_id INT,
            issue_key VARCHAR(255) NOT NULL,
            comment TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (tracker_id) REFERENCES trackers(id) ON DELETE CASCADE
        )
    ''')
    # Todos Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS todos (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT,
            title TEXT NOT NULL,
            description TEXT,
            priority VARCHAR(50) DEFAULT 'Low',
            due_date DATE NOT NULL,
            status VARCHAR(50) DEFAULT 'Pending',
            tags TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS todo_tags (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(255) NOT NULL,
            color VARCHAR(50) DEFAULT 'blue'
        )
    ''')
    # Teams Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS teams (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(255) NOT NULL,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    # Team Members Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS team_members (
            id INT AUTO_INCREMENT PRIMARY KEY,
            team_id INT,
            account_id VARCHAR(255) NOT NULL,
            display_name VARCHAR(255) NOT NULL,
            avatar_url TEXT,
            FOREIGN KEY (team_id) REFERENCES teams(id) ON DELETE CASCADE
        )
    ''')
    # Sprints Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sprints (
            id INT AUTO_INCREMENT PRIMARY KEY,
            team_id INT,
            name VARCHAR(255) NOT NULL,
            state VARCHAR(50) DEFAULT 'active',
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (team_id) REFERENCES teams(id) ON DELETE CASCADE
        )
    ''')
    # Sprint Weeks Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sprint_weeks (
            id INT AUTO_INCREMENT PRIMARY KEY,
            sprint_id INT,
            week_number INT NOT NULL,
            goal TEXT,
            FOREIGN KEY (sprint_id) REFERENCES sprints(id) ON DELETE CASCADE
        )
    ''')
    # Sprint Tickets Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sprint_tickets (
            id INT AUTO_INCREMENT PRIMARY KEY,
            sprint_id INT,
            week_id INT,
            issue_key VARCHAR(255) NOT NULL,
            comment TEXT,
            pr_raised INT DEFAULT 0,
            demo_done INT DEFAULT 0,
            pr_merged INT DEFAULT 0,
            deploy_status VARCHAR(50) DEFAULT 'N/A',
            qa_assignee VARCHAR(255) DEFAULT '',
            qa_status VARCHAR(50) DEFAULT 'Pending',
            bugs_found TEXT,
            requirements_clear VARCHAR(50) DEFAULT 'No',
            completed INT DEFAULT 0,
            is_flagged INT DEFAULT 0
        )
    ''')

    # Ensure is_flagged exists if table already exists
    try:
        cursor.execute("ALTER TABLE sprint_tickets ADD COLUMN is_flagged INT DEFAULT 0")
    except:
        pass

    # Drop legacy foreign keys that cause issues with Jira Sprint IDs
    try:
        cursor.execute("ALTER TABLE sprint_tickets DROP FOREIGN KEY sprint_tickets_ibfk_1")
    except:
        pass
    try:
        cursor.execute("ALTER TABLE sprint_tickets DROP FOREIGN KEY sprint_tickets_ibfk_2")
    except:
        pass

    # Payroll Dashboard - local ticket metadata (does not update Jira)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS payroll_ticket_meta (
            id INT AUTO_INCREMENT PRIMARY KEY,
            issue_key VARCHAR(50) NOT NULL UNIQUE,
            manual_bug_count INT DEFAULT 0,
            manual_bug_links TEXT,
            test_cases_count INT DEFAULT 0,
            notes TEXT,
            updated_by VARCHAR(255),
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
        )
    ''')

    # Scrum Notes Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS scrum_notes (
            id          INT AUTO_INCREMENT PRIMARY KEY,
            `date`      DATE    NOT NULL,
            team_id     INT NOT NULL,
            member_id   VARCHAR(255)    NOT NULL,
            member_name VARCHAR(255)    NOT NULL,
            ticket_key  VARCHAR(255)    NOT NULL,
            comment     TEXT,
            deadline    DATE,
            status      VARCHAR(50)    DEFAULT 'Pending',
            tags        TEXT,
            created_at  DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (team_id) REFERENCES teams(id) ON DELETE CASCADE
        )
    ''')

    # Jira Config Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS jira_config (
            id INT PRIMARY KEY,
            email VARCHAR(255),
            api_token TEXT,
            project_key VARCHAR(255),
            jira_domain VARCHAR(255)
        )
    ''')
    
    # Custom Reports Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS custom_reports (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(255) NOT NULL,
            jql TEXT NOT NULL,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')

    # --- DYNAMIC TRACKERS (V2) ---
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS trackers_v2 (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(255) NOT NULL,
            jql TEXT NOT NULL,
            created_by INT,
            is_public BOOLEAN DEFAULT FALSE,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (created_by) REFERENCES users(id) ON DELETE SET NULL
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS tracker_columns (
            id INT AUTO_INCREMENT PRIMARY KEY,
            tracker_id INT NOT NULL,
            name VARCHAR(255) NOT NULL,
            column_type ENUM('text', 'checkbox', 'select', 'user', 'rca') NOT NULL,
            options TEXT,
            order_index INT DEFAULT 0,
            FOREIGN KEY (tracker_id) REFERENCES trackers_v2(id) ON DELETE CASCADE
        )
    ''')

    # Ensure rca is in the enum if table already exists
    try:
        cursor.execute("ALTER TABLE tracker_columns MODIFY COLUMN column_type ENUM('text', 'checkbox', 'select', 'user', 'rca') NOT NULL")
    except:
        pass

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS tracker_data_v2 (
            id INT AUTO_INCREMENT PRIMARY KEY,
            tracker_id INT NOT NULL,
            issue_key VARCHAR(50) NOT NULL,
            column_id INT NOT NULL,
            value TEXT,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
            UNIQUE KEY unique_cell (tracker_id, issue_key, column_id),
            FOREIGN KEY (tracker_id) REFERENCES trackers_v2(id) ON DELETE CASCADE,
            FOREIGN KEY (column_id) REFERENCES tracker_columns(id) ON DELETE CASCADE
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS tracker_rca (
            id INT AUTO_INCREMENT PRIMARY KEY,
            tracker_id INT NOT NULL,
            issue_key VARCHAR(50) NOT NULL,
            issue_details TEXT,
            rca_text TEXT,
            fix_text TEXT,
            prevention_text TEXT,
            token VARCHAR(100) UNIQUE,
            submitted_at DATETIME,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            UNIQUE KEY unique_rca (tracker_id, issue_key),
            FOREIGN KEY (tracker_id) REFERENCES trackers_v2(id) ON DELETE CASCADE
        )
    ''')

    
    # App Config Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS app_config (
            config_key VARCHAR(255) PRIMARY KEY,
            config_value TEXT,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
        )
    ''')

    # Audit Logs Table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS audit_logs (
            id INT AUTO_INCREMENT PRIMARY KEY,
            user_id INT,
            user_name VARCHAR(255),
            page VARCHAR(100),
            item_key VARCHAR(255),
            field_name VARCHAR(255),
            old_value TEXT,
            new_value TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE SET NULL
        )
    ''')

    cursor.execute('''
        CREATE TABLE IF NOT EXISTS team_diagram_ai_cache (
            cache_key CHAR(64) NOT NULL PRIMARY KEY,
            sub_team_name VARCHAR(512) NOT NULL,
            response_json LONGTEXT NOT NULL,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
        ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
    ''')

    # Seed Default Header Title if empty
    cursor.execute("SELECT count(*) FROM app_config WHERE config_key = 'header_title'")
    if cursor.fetchone()[0] == 0:
        cursor.execute("INSERT INTO app_config (config_key, config_value) VALUES ('header_title', 'Jira Analytics')")

    # Sprint Tracker - Sprints
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sprint_tracker_sprints (
            id INT AUTO_INCREMENT PRIMARY KEY,
            name VARCHAR(100) NOT NULL UNIQUE,
            sprint_goal TEXT,
            goal_edited TINYINT(1) DEFAULT 0,
            divider_index INT DEFAULT 0,
            sort_order INT DEFAULT 0,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP
        )
    ''')

    # Sprint Tracker - Themes
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sprint_tracker_themes (
            id INT AUTO_INCREMENT PRIMARY KEY,
            sprint_id INT NOT NULL,
            theme_key VARCHAR(150) NOT NULL,
            epic_name VARCHAR(500),
            sentence TEXT,
            bullets JSON,
            lb_override TINYINT(1) DEFAULT NULL,
            notes TEXT,
            notes_updated_by VARCHAR(255),
            notes_updated_at DATETIME,
            sort_order INT DEFAULT 0,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
            UNIQUE KEY unique_theme_in_sprint (sprint_id, theme_key),
            FOREIGN KEY (sprint_id) REFERENCES sprint_tracker_sprints(id) ON DELETE CASCADE
        )
    ''')

    # Backward-compatible migration for existing DBs
    try:
        cursor.execute("ALTER TABLE sprint_tracker_themes ADD COLUMN notes_updated_by VARCHAR(255)")
    except Exception:
        pass
    try:
        cursor.execute("ALTER TABLE sprint_tracker_themes ADD COLUMN notes_updated_at DATETIME")
    except Exception:
        pass

    # Sprint Tracker - Tickets
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS sprint_tracker_tickets (
            id INT AUTO_INCREMENT PRIMARY KEY,
            theme_id INT NOT NULL,
            ticket_key VARCHAR(100) NOT NULL,
            summary TEXT,
            status VARCHAR(100),
            customers JSON,
            lb TINYINT(1) DEFAULT 0,
            description_bullets JSON,
            last_synced_at DATETIME,
            sort_order INT DEFAULT 0,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
            updated_at DATETIME DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,
            UNIQUE KEY unique_ticket_in_theme (theme_id, ticket_key),
            FOREIGN KEY (theme_id) REFERENCES sprint_tracker_themes(id) ON DELETE CASCADE
        )
    ''')

    # Seed Sprint Tracker from JSON if empty
    cursor.execute("SELECT COUNT(*) FROM sprint_tracker_sprints")
    if cursor.fetchone()[0] == 0:
        try:
            seed_path = os.path.join(os.path.dirname(__file__), "sprint_tracker_seed.json")
            if os.path.exists(seed_path):
                with open(seed_path, "r", encoding="utf-8") as f:
                    seed = json.load(f)
                for s_idx, sprint in enumerate(seed):
                    themes = sprint.get("themes", [])
                    cursor.execute(
                        """INSERT INTO sprint_tracker_sprints (name, sprint_goal, goal_edited, divider_index, sort_order)
                           VALUES (%s, %s, 0, %s, %s)""",
                        (sprint["name"], sprint.get("sprint_goal", ""), len(themes), s_idx)
                    )
                    sprint_id = cursor.lastrowid
                    for t_idx, theme in enumerate(themes):
                        cursor.execute(
                            """INSERT INTO sprint_tracker_themes
                               (sprint_id, theme_key, epic_name, sentence, bullets, notes, sort_order)
                               VALUES (%s, %s, %s, %s, %s, '', %s)""",
                            (
                                sprint_id,
                                theme["theme_key"],
                                theme.get("epic_name", ""),
                                theme.get("sentence", ""),
                                json.dumps(theme.get("bullets", [])),
                                t_idx,
                            ),
                        )
                        theme_id = cursor.lastrowid
                        for k_idx, ticket in enumerate(theme.get("tickets", [])):
                            cursor.execute(
                                """INSERT INTO sprint_tracker_tickets
                                   (theme_id, ticket_key, summary, status, customers, lb, sort_order)
                                   VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                                (
                                    theme_id,
                                    ticket["ticket_key"],
                                    ticket.get("summary", ""),
                                    ticket.get("status", "Open"),
                                    json.dumps(ticket.get("customers", [])),
                                    1 if ticket.get("lb") else 0,
                                    k_idx,
                                ),
                            )
        except Exception as _seed_err:
            print(f"Sprint Tracker seed error: {_seed_err}")

    conn.commit()
    conn.close()

init_db()

# =========================
# PAGE ROUTE
# =========================
@app.route("/settings")
@login_required
@page_permission_required("settings")
def settings():
    return render_template("settings.html", project=PROJECT_KEY)

@app.route("/api/settings/jira", methods=["GET", "POST"])
@login_required
def jira_settings_api():
    if request.method == "POST" and not current_user.has_permission("manage_settings") and current_user.role_name != 'Admin':
        return jsonify({"error": "Permission denied"}), 403
        
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "POST":
        data = request.json
        email = data.get("email", "").strip()
        token = data.get("token", "").strip()
        project = data.get("project_key", "").strip().upper()
        domain = data.get("jira_domain", "").strip()
        
        if not email or not token or not project:
            return jsonify({"error": "All fields are required"}), 400
            
        cursor.execute("""
            INSERT INTO jira_config (id, email, api_token, project_key, jira_domain)
            VALUES (1, %s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE
                email=VALUES(email),
                api_token=VALUES(api_token),
                project_key=VALUES(project_key),
                jira_domain=VALUES(jira_domain)
        """, (email, token, project, domain))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
        
    else:  # GET
        cursor.execute("SELECT email, api_token, project_key, jira_domain FROM jira_config WHERE id = 1")
        row = cursor.fetchone()
        conn.close()
        if row:
            return jsonify({
                "email": row[0],
                "token": row[1],
                "project_key": row[2],
                "jira_domain": row[3] or DEFAULT_JIRA_DOMAIN
            })
        return jsonify({"email": "", "token": "", "project_key": "", "jira_domain": DEFAULT_JIRA_DOMAIN})

@app.route("/api/settings/app_config", methods=["GET", "POST"])
@login_required
def app_config_settings_api():
    if request.method == "POST" and current_user.role_name != 'Admin':
        return jsonify({"error": "Permission denied"}), 403
        
    conn, cursor = get_db_connection(dictionary=True)
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "POST":
        data = request.json or {}

        if not data:
            conn.close()
            return jsonify({"error": "No config values provided"}), 400

        try:
            for key, value in data.items():
                config_key = str(key).strip()
                if not config_key:
                    continue
                config_value = "" if value is None else str(value).strip()
                cursor.execute("""
                    INSERT INTO app_config (config_key, config_value)
                    VALUES (%s, %s)
                    ON DUPLICATE KEY UPDATE config_value = VALUES(config_value)
                """, (config_key, config_value))
            conn.commit()
            return jsonify({"success": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500
        finally:
            conn.close()
    else: # GET
        try:
            cursor.execute("SELECT config_key, config_value FROM app_config")
            rows = cursor.fetchall()
            config = {row['config_key']: row['config_value'] for row in rows}
            if current_user.role_name != 'Admin' and not current_user.has_permission("manage_settings"):
                masked = dict(config)
                anthropic_value = str(masked.pop("anthropic_api_key", "") or "").strip()
                masked["anthropic_api_key_configured"] = bool(anthropic_value)
                return jsonify(masked)
            return jsonify(config)
        except Exception as e:
            return jsonify({"error": str(e)}), 500
        finally:
            conn.close()

@app.route("/api/settings/todo_tags", methods=["GET", "POST"])
def todo_tags_settings_api():
    conn, cursor = get_db_connection(dictionary=True)
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "POST":
        data = request.json or {}
        name = (data.get("name") or "").strip()
        color = (data.get("color") or "blue").strip()
        
        if not name:
            conn.close()
            return jsonify({"error": "Tag name is required"}), 400
        
        cursor.execute("INSERT INTO todo_tags (name, color) VALUES (%s, %s)", (name, color))
        conn.commit()
        tag_id = cursor.lastrowid
        conn.close()
        return jsonify({"id": tag_id, "name": name, "color": color})
        
    else:
        cursor.execute("SELECT id, name, color FROM todo_tags ORDER BY id ASC")
        # Use column names since we use dictionary=True
        tags = [{"id": r['id'], "name": r['name'], "color": r['color']} for r in cursor.fetchall()]
        conn.close()
        return jsonify(tags)

@app.route("/api/settings/todo_tags/<int:tag_id>", methods=["DELETE"])
@login_required
def delete_todo_tag(tag_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    cursor.execute("DELETE FROM todo_tags WHERE id = %s", (tag_id,))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

# =========================
# 🔑 AUTH ROUTES
# =========================
@app.route("/login")
def login_page():
    conn, cursor = get_db_connection(dictionary=True)
    if not conn: return render_template("login.html", header_title="Reports Dashboard")
    
    cursor.execute("SELECT config_value FROM app_config WHERE config_key = 'header_title'")
    row = cursor.fetchone()
    conn.close()
    
    header_title = row['config_value'] if row else "Reports Dashboard"
    return render_template("login.html", header_title=header_title)

@app.route("/login/google")
def login_google():
    redirect_uri = url_for('auth_callback', _external=True)
    return google.authorize_redirect(redirect_uri)

@app.route("/auth/callback")
def auth_callback():
    try:
        token = google.authorize_access_token()
        user_info = google.parse_id_token(token, nonce=None)
        
        email = user_info.get('email')
        name = user_info.get('name')
        google_id = user_info.get('sub')
        
        conn, cursor = get_db_connection()
        if not conn: return "Database error", 500
        
        # Check if user exists
        cursor.execute("SELECT id FROM users WHERE email = %s", (email,))
        row = cursor.fetchone()
        
        if row:
            user_id = row[0]
            # Update google_id if missing
            cursor.execute("UPDATE users SET google_id = %s, name = %s WHERE id = %s", (google_id, name, user_id))
        else:
            # Create new user
            # Assign 'Viewer' role by default
            cursor.execute("SELECT id FROM roles WHERE name = 'Viewer'")
            role_row = cursor.fetchone()
            role_id = role_row[0] if role_row else 3
            
            api_token = secrets.token_hex(32)
            cursor.execute("INSERT INTO users (email, name, google_id, role_id, api_token) VALUES (%s, %s, %s, %s, %s)",
                           (email, name, google_id, role_id, api_token))
            user_id = cursor.lastrowid
            
        conn.commit()
        conn.close()
        
        # Login the user
        user_obj = load_user(user_id)
        login_user(user_obj)
        
        return redirect(url_for('index'))
        
    except Exception as e:
        print(f"Auth Error: {e}")
        return f"Authentication failed: {e}", 400

@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for('login_page'))

# =========================
# 👑 ADMIN ROUTES
# =========================
@app.route("/admin/users")
@admin_required
def admin_users():
    conn, cursor = get_db_connection(dictionary=True)
    if not conn: return jsonify({"error": "Database error"}), 500
    cursor.execute("""
        SELECT u.*, r.name as role_name 
        FROM users u 
        LEFT JOIN roles r ON u.role_id = r.id
        ORDER BY u.created_at DESC
    """)
    users = cursor.fetchall()
    
    cursor.execute("SELECT * FROM roles")
    roles = cursor.fetchall()
    conn.close()
    return render_template("admin_users.html", users=users, roles=roles)

@app.route("/api/admin/users/<int:user_id>/role", methods=["POST"])
@admin_required
def update_user_role(user_id):
    role_id = request.json.get("role_id")
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    cursor.execute("UPDATE users SET role_id = %s WHERE id = %s", (role_id, user_id))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

PAGE_PERMISSIONS = [
    {"key": "dashboard", "label": "Dashboard"},
    {"key": "scoreboard", "label": "Scoreboard"},
    {"key": "customer_dashboard", "label": "Customer Dashboard"},
    {"key": "team_productivity", "label": "Team Productivity"},
    {"key": "customer_closure", "label": "Customer Closure"},
    {"key": "sprint_dashboard", "label": "Sprint Dashboard"},
    {"key": "explorer", "label": "Explorer"},
    {"key": "custom_reports", "label": "Custom Reports"},
    {"key": "scrum_notes", "label": "Scrum Notes"},
    {"key": "work_report", "label": "Work Report"},
    {"key": "assignee_work", "label": "Assignee Work"},
    {"key": "payroll_dashboard", "label": "Payroll Dashboard"},
    {"key": "planning", "label": "Sprint Planning"},
    {"key": "teams", "label": "Teams"},
    {"key": "settings", "label": "Settings"},
    {"key": "todo", "label": "Todo List"},
    {"key": "tracker", "label": "Tracker"},
    {"key": "status_tracker", "label": "Status Tracker"},
    {"key": "trackers_v2", "label": "Custom Trackers"},
    {"key": "query_builder", "label": "Query Builder"},
    {"key": "bulk_update", "label": "Bulk Update"},
    {"key": "merge_pdf", "label": "Merge PDF"},
    {"key": "sprint_tracker", "label": "Sprint Tracker"},
    {"key": "team_diagram", "label": "Team Diagram Report"}
]

@app.route("/admin/roles")
@admin_required
def admin_roles():
    conn, cursor = get_db_connection(dictionary=True)
    if not conn: return jsonify({"error": "Database error"}), 500
    cursor.execute("SELECT * FROM roles ORDER BY id ASC")
    roles = []
    for row in cursor.fetchall():
        roles.append({
            "id": row["id"],
            "name": row["name"],
            "permissions": json.loads(row["permissions"]) if row["permissions"] else {}
        })
    conn.close()
    
    # Define available permissions for the UI
    available_permissions = [
        {"key": "view_reports", "label": "View Reports & Dashboards"},
        {"key": "create_tickets", "label": "Create Tickets/Todos"},
        {"key": "edit_tickets", "label": "Edit Tickets/Todos"},
        {"key": "delete_tickets", "label": "Delete Tickets/Todos"},
        {"key": "manage_settings", "label": "Manage Settings"},
        {"key": "manage_teams", "label": "Manage Teams"},
        {"key": "manage_trackers", "label": "Manage Custom Trackers (Create/Edit Structure)"},
        {"key": "edit_tracker_data", "label": "Edit Tracker Data (Update Status/Fields)"}
    ]
    
    return render_template("admin_roles.html", roles=roles, permissions=available_permissions, page_permissions=PAGE_PERMISSIONS)

@app.route("/api/admin/roles", methods=["POST"])
@admin_required
def create_role():
    name = request.json.get("name")
    if not name:
        return jsonify({"error": "Role name is required"}), 400
    
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    try:
        cursor.execute("INSERT INTO roles (name, permissions) VALUES (%s, '{}')", (name,))
        conn.commit()
        role_id = cursor.lastrowid
        conn.close()
        return jsonify({"success": True, "id": role_id, "name": name})
    except Error:
        if conn: conn.close()
        return jsonify({"error": "Role name already exists or database error"}), 400

@app.route("/api/admin/roles/<int:role_id>", methods=["PUT", "DELETE"])
@admin_required
def manage_role(role_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "DELETE":
        # Prevent deleting the last Admin role or system roles if necessary
        # For now, just checking ID 1 (Admin)
        if role_id == 1:
            conn.close()
            return jsonify({"error": "Cannot delete the default Admin role"}), 400
            
        cursor.execute("DELETE FROM roles WHERE id = %s", (role_id,))
        # Optional: Reset users with this role to Viewer (ID 3) or NULL
        cursor.execute("UPDATE users SET role_id = 3 WHERE role_id = %s", (role_id,))
        
        conn.commit()
        conn.close()
        return jsonify({"success": True})
        
    elif request.method == "PUT":
        data = request.json
        permissions = data.get("permissions") # Expecting a dict/object
        
        if permissions is not None:
            perm_json = json.dumps(permissions)
            cursor.execute("UPDATE roles SET permissions = %s WHERE id = %s", (perm_json, role_id))
            conn.commit()
            conn.close()
            return jsonify({"success": True})
            
        conn.close()
        return jsonify({"error": "No data provided"}), 400

@app.route("/")
@login_required
@page_permission_required("dashboard")
def index():
    email, token, _, _ = _get_jira_config()
    if not email or not token:
        return redirect(url_for('settings'))
    return render_template("index.html", project=PROJECT_KEY)

# =========================
# FETCH ASSIGNEES
# =========================
@app.route("/api/assignees", methods=["GET", "POST"])
def assignees():
    # If Select2 sends a search query
    data = request.get_json(silent=True) or {}
    query = request.args.get('q') or data.get('q')

    project_key_str = str(PROJECT_KEY)
    if not project_key_str:
        return jsonify({"error": "Missing project key. Save it in Settings first."}), 400
    if "Authorization" not in HEADERS:
        return jsonify({"error": "Missing Jira credentials. Save them in Settings first."}), 401
    
    # Use assignable search to get users who can actually be assigned to the project
    params = {
        "project": PROJECT_KEY,
        "maxResults": 100
    }
    if query:
        params["query"] = query

    try:
        # Explicitly cast LocalProxy objects to ensure they are passed correctly
        headers_dict = dict(HEADERS)
        project_key_str = str(PROJECT_KEY)
        
        print(f"DEBUG: calling Jira assignable/search with project={project_key_str}")
        
        jira_res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/user/assignable/search",
            headers=headers_dict,
            params={
                "project": project_key_str,
                "maxResults": 100,
                "query": query
            } if query else {
                "project": project_key_str,
                "maxResults": 100
            }
        )
        
        if jira_res.status_code != 200:
            print(f"DEBUG: Jira Error {jira_res.status_code}: {jira_res.text}")
            return jsonify({"error": f"Jira API error: {jira_res.text}"}), jira_res.status_code
            
        users = jira_res.json()
        if not isinstance(users, list):
            # Sometimes Jira returns an error object instead of a list if something is wrong
            return jsonify(users), jira_res.status_code

        # Format for Select2
        results = [{"id": u.get("accountId"), "name": u.get("displayName")} for u in users if u.get("accountType") == "atlassian"]
        return jsonify(results)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/jira/search", methods=["GET"])
def jira_search():
    query = request.args.get('q', '')
    if len(query) < 2:
        return jsonify([])
        
    project_key_str = str(PROJECT_KEY)
    headers_dict = dict(HEADERS)
    
    # Search for issues by key or summary in the current project
    jql = f'project = "{project_key_str}" AND (key ~ "{query}*" OR summary ~ "{query}*")'
    
    try:
        res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params={
                "jql": jql,
                "maxResults": 10,
                "fields": "summary,issuetype"
            }
        )
        data = res.json()
        results = []
        for issue in data.get('issues', []):
            results.append({
                "id": issue['key'],
                "text": f"{issue['key']}: {issue['fields']['summary']}",
                "key": issue['key']
            })
        return jsonify(results)
    except Exception as e:
        print(f"DEBUG: Search Error: {str(e)}")
        return jsonify({"error": str(e)}), 500

# =========================
# FETCH CUSTOMERS
# =========================
@app.route("/api/customers", methods=["GET"])
def customers():
    query = request.args.get('q') or request.args.get('term')
    now_ts = time.time()
    q = (query or "").strip().lower()

    def build_results(values):
        out = []
        for c in sorted(values):
            if not q or q in c.lower():
                out.append({"id": c, "name": c})
        return out

    # Serve cached "all customers" for fast dropdown load.
    if CUSTOMER_CACHE["values"] and CUSTOMER_CACHE["expires_at"] > now_ts:
        return jsonify(build_results(CUSTOMER_CACHE["values"]))

    unique_customers = set()
    headers_dict = dict(HEADERS)

    # For free-text search without cache, do a focused Jira query.
    if q:
        jql = f'"Customer" ~ "{q}*"'
        scan_range = [0, 100]
    else:
        jql = 'Customer is not EMPTY'
        # First-time warmup: scan enough pages once, then cache.
        scan_range = range(0, 1500, 100)

    for start_at in scan_range:
        jira_res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params={
                "jql": jql,
                "maxResults": 100,
                "startAt": start_at,
                "fields": "customfield_10077"
            }
        )

        if jira_res.status_code != 200:
            break

        res = jira_res.json()
        issues = res.get("issues", [])
        if not issues:
            break

        for i in issues:
            cust_val = i["fields"].get("customfield_10077")
            if isinstance(cust_val, list):
                for c in cust_val:
                    if isinstance(c, dict):
                        val = c.get("value")
                        if val:
                            unique_customers.add(str(val).strip())
                    elif c:
                        unique_customers.add(str(c).strip())
            elif isinstance(cust_val, dict):
                val = cust_val.get("value")
                if val:
                    unique_customers.add(str(val).strip())
            elif cust_val:
                unique_customers.add(str(cust_val).strip())

    # Cache only full-list refreshes (no query).
    if not q and unique_customers:
        CUSTOMER_CACHE["values"] = sorted(unique_customers)
        CUSTOMER_CACHE["expires_at"] = now_ts + CUSTOMER_CACHE_TTL

    return jsonify(build_results(unique_customers))

# =========================
# JQL SEARCH PROXY
# =========================
@app.route("/api/search", methods=["POST"])
def search():
    # Attempt to handle both JSON and Form data to avoid 415
    data = request.get_json(silent=True) or request.form.to_dict()
    
    # Cast headers
    headers_dict = dict(HEADERS)
    
    jira_res = requests.post(
        f"{JIRA_DOMAIN}/rest/api/3/search/jql",
        headers=headers_dict,
        json=data
    )
    return jsonify(jira_res.json())

# =========================
# SCOREBOARD ROUTE
# =========================
@app.route("/scoreboard")
@page_permission_required("scoreboard")
def scoreboard():
    return render_template("scoreboard.html", project=PROJECT_KEY)

@app.route("/customer_dashboard")
def customer_dashboard_page():
    return render_template("customer_dashboard.html", project=PROJECT_KEY, show_navbar=False)

@app.route("/sprint_dashboard")
@page_permission_required("sprint_dashboard")
def sprint_dashboard_page():
    return render_template("sprint_dashboard.html", project=PROJECT_KEY, show_navbar=False)


@app.route("/team_productivity")
def team_productivity_page():
    return render_template("team_productivity.html", project=PROJECT_KEY, show_navbar=False)


@app.route("/customer_closure")
def customer_closure_page():
    return render_template("customer_closure.html", project=PROJECT_KEY, show_navbar=False)


def _extract_customer_values(raw_value):
    customers = []
    if isinstance(raw_value, list):
        for item in raw_value:
            if isinstance(item, dict):
                val = (item.get("value") or "").strip()
                if val:
                    customers.append(val)
            elif item:
                customers.append(str(item).strip())
    elif isinstance(raw_value, dict):
        val = (raw_value.get("value") or "").strip()
        if val:
            customers.append(val)
    elif raw_value:
        customers.append(str(raw_value).strip())
    # Keep stable order but remove duplicates
    deduped = []
    seen = set()
    for c in customers:
        k = c.lower()
        if k in seen:
            continue
        seen.add(k)
        deduped.append(c)
    return deduped

def _is_launch_blocker(labels):
    normalized = [str(l).strip().lower().replace("-", "_").replace(" ", "_") for l in (labels or [])]
    strong_matches = {"launch_blocker", "launchblocker", "lb", "launch_blocked"}
    if any(v in strong_matches for v in normalized):
        return True
    return any(("launch" in v and "block" in v) for v in normalized)

def _is_high_priority(priority_name):
    p = (priority_name or "").strip().lower()
    if not p:
        return False
    high_tokens = ("highest", "high", "critical", "blocker", "p0", "p1", "sev1", "sev-1")
    return any(tok in p for tok in high_tokens)

def _escape_jql_value(value):
    return str(value or "").replace("\\", "\\\\").replace('"', '\\"')

def _extract_sprint_values(raw_value):
    """Return stable sprint list from Jira sprint custom field."""
    out = []
    seen = set()
    if not isinstance(raw_value, list):
        return out
    for item in raw_value:
        sprint_id = None
        sprint_name = "Unknown Sprint"
        if isinstance(item, dict):
            try:
                sprint_id = int(item.get("id"))
            except Exception:
                sprint_id = None
            sprint_name = (item.get("name") or "").strip() or "Unknown Sprint"
        elif isinstance(item, str):
            id_match = re.search(r"id=(\d+)", item)
            name_match = re.search(r"name=([^,\]]+)", item)
            if id_match:
                try:
                    sprint_id = int(id_match.group(1))
                except Exception:
                    sprint_id = None
            if name_match:
                sprint_name = name_match.group(1).strip() or "Unknown Sprint"
        key = (sprint_id, sprint_name.lower())
        if key in seen:
            continue
        seen.add(key)
        out.append({"id": sprint_id, "name": sprint_name})
    return out

# Lowercase Jira status *names* treated as completed (plus statusCategory "done" below).
_JIRA_COMPLETED_STATUS_NAMES_LOWER = frozenset(
    {
        "done",
        "resolved",
        "closed",
        "ready for qa",
        "deployed",
        "ready for staging",
        "ready for internal demo",
        "qa in progress",
        "not a bug",
        "unable to reproduce",
        # Common variants still seen in workflows
        "staged",
        "fixed",
        "complete",
        "completed",
        "released",
        "won't fix",
        "wont fix",
    }
)


def _is_done_like_status(status_name, status_category=None):
    """
    True if the issue counts as completed for throughput metrics.
    Uses Jira statusCategory key when provided; otherwise status name only.
    """
    if status_category is not None:
        cat = str(status_category).strip().lower()
        if cat == "done":
            return True
    key = (status_name or "").strip().lower()
    return key in _JIRA_COMPLETED_STATUS_NAMES_LOWER

def _is_bug_type(issue_type_name):
    return (issue_type_name or "").strip().lower() == "bug"

def _issue_type_bucket(issue_type_name):
    k = (issue_type_name or "").strip().lower()
    if k == "bug":
        return "Bug"
    if k in ("story", "user story"):
        return "Story"
    if k == "task":
        return "Task"
    return "Other"


def _is_test_case_issue_type(issue_type_name):
    """Treat Zephyr/Xray-style test issues as test cases for team diagram metrics."""
    k = (issue_type_name or "").strip().lower()
    if not k:
        return False
    if k in ("test", "test case", "zephyr test", "test execution"):
        return True
    if "test case" in k:
        return True
    return False


def _jira_issue_is_automated_test(issue_type_name, fields_obj):
    """Automation: issue type name mentions automated, or any label contains 'automated'."""
    k = (issue_type_name or "").strip().lower()
    if "automated" in k:
        return True
    for lab in (fields_obj or {}).get("labels") or []:
        if "automated" in str(lab).lower():
            return True
    return False


_CUSTOMER_RESOLUTION_MILESTONE_STATUSES = _JIRA_COMPLETED_STATUS_NAMES_LOWER


def _reported_by_customer_jql_clause():
    """Extra JQL fragment (without leading AND). Override with JIRA_REPORTED_BY_CUSTOMER_JQL if the field id differs."""
    return (os.environ.get("JIRA_REPORTED_BY_CUSTOMER_JQL") or '"Reported By Customer" = Yes').strip()


_DEFAULT_CUSTOMER_DONE_STATUSES = (
    "Done",
    "Resolved",
    "Closed",
    "Ready for QA",
    "Deployed",
    "Ready for Staging",
    "Ready for Internal Demo",
    "QA In Progress",
    "Not a Bug",
    "Unable to Reproduce",
    "Staged",
    "Fixed",
    "Complete",
    "Completed",
    "Released",
    "Won't Fix",
)


def _normalize_status_list(values):
    if isinstance(values, str):
        raw_items = re.split(r"[\n,]+", values)
    elif isinstance(values, list):
        raw_items = values
    else:
        raw_items = []

    out = []
    seen = set()
    for item in raw_items:
        val = str(item or "").strip()
        if not val:
            continue
        key = val.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(val)
    return out


def _customer_created_range_jql(date_start, date_end_exclusive):
    project_key_str = str(PROJECT_KEY).strip()
    reported_clause = _reported_by_customer_jql_clause()
    return (
        f'project = "{_escape_jql_value(project_key_str)}" AND ({reported_clause}) '
        f'AND created >= "{date_start}" AND created < "{date_end_exclusive}" '
        f"ORDER BY created DESC"
    )


def _parse_jira_datetime(ts_str):
    if not ts_str:
        return None
    s = str(ts_str).strip().replace("Z", "+00:00")
    m = re.search(r"([+-])(\d{2})(\d{2})$", s)
    if m and s[m.start() - 1 : m.start()] != ":":
        s = s[: m.start()] + f"{m.group(1)}{m.group(2)}:{m.group(3)}"
    try:
        return datetime.fromisoformat(s)
    except ValueError:
        return None


def _team_name_from_jira_fields(fields_obj):
    team_raw = fields_obj.get("customfield_10001")
    if isinstance(team_raw, dict):
        team_name = (team_raw.get("name") or team_raw.get("value") or "").strip()
    elif isinstance(team_raw, list) and team_raw:
        first = team_raw[0]
        if isinstance(first, dict):
            team_name = (first.get("name") or first.get("value") or "").strip()
        else:
            team_name = str(first).strip()
    elif team_raw:
        team_name = str(team_raw).strip()
    else:
        team_name = "Unspecified"
    return team_name or "Unspecified"


def _first_customer_resolution_milestone_at(issue, fields_obj):
    """
    First time the issue entered a completed status (from changelog); see _JIRA_COMPLETED_STATUS_NAMES_LOWER.
    Falls back to resolutiondate if no matching status transition is present.
    """
    histories = (issue.get("changelog") or {}).get("histories") or []
    hits = []
    for history in histories:
        ts = _parse_jira_datetime(history.get("created"))
        if not ts:
            continue
        for item in history.get("items") or []:
            if item.get("field") != "status":
                continue
            to_s = item.get("toString") or ""
            if (to_s or "").strip().lower() in _CUSTOMER_RESOLUTION_MILESTONE_STATUSES:
                hits.append(ts)
    if hits:
        return min(hits)
    res_raw = fields_obj.get("resolutiondate")
    if res_raw:
        return _parse_jira_datetime(res_raw)
    return None


@app.route("/api/customer_closure/teams", methods=["POST"])
def customer_closure_teams():
    payload = request.get_json(silent=True) or {}
    date_start = (payload.get("date_start") or "").strip()
    date_end = (payload.get("date_end") or "").strip()
    if not re.match(r"^\d{4}-\d{2}-\d{2}$", date_start) or not re.match(r"^\d{4}-\d{2}-\d{2}$", date_end):
        return jsonify({"error": "date_start and date_end are required (YYYY-MM-DD)"}), 400
    try:
        ds = datetime.strptime(date_start, "%Y-%m-%d").date()
        de = datetime.strptime(date_end, "%Y-%m-%d").date()
    except ValueError:
        return jsonify({"error": "Invalid date format"}), 400
    if de < ds:
        return jsonify({"error": "date_end must be on or after date_start"}), 400

    project_key_str = str(PROJECT_KEY).strip()
    if not project_key_str:
        return jsonify({"error": "Missing project key. Save it in Settings first."}), 400
    headers_dict = _get_explicit_request_jira_headers()
    if "Authorization" not in headers_dict:
        return jsonify({"error": "Add Jira email and API token in local settings first."}), 401

    jql = _customer_created_range_jql(date_start, (de + timedelta(days=1)).strftime("%Y-%m-%d"))
    params = {
        "jql": jql,
        "maxResults": 300,
        "fields": "customfield_10001",
    }
    jira_res = requests.get(
        f"{JIRA_DOMAIN}/rest/api/3/search/jql",
        headers=headers_dict,
        params=params,
    )
    if jira_res.status_code != 200:
        return jsonify({"error": f"Jira API error: {jira_res.text}"}), jira_res.status_code

    teams = set()
    for issue in (jira_res.json().get("issues") or []):
        team_name = _team_name_from_jira_fields(issue.get("fields") or {})
        if team_name:
            teams.add(team_name)
    return jsonify({"teams": sorted(teams, key=lambda x: x.lower())})


@app.route("/api/customer_closure/data", methods=["POST"])
def customer_closure_data():
    payload = request.get_json(silent=True) or {}
    date_start = (payload.get("date_start") or "").strip()
    date_end = (payload.get("date_end") or "").strip()
    team_name_filter = (payload.get("team_name") or "").strip()
    done_statuses = _normalize_status_list(payload.get("done_statuses")) or list(_DEFAULT_CUSTOMER_DONE_STATUSES)

    if not re.match(r"^\d{4}-\d{2}-\d{2}$", date_start) or not re.match(r"^\d{4}-\d{2}-\d{2}$", date_end):
        return jsonify({"error": "date_start and date_end are required (YYYY-MM-DD)"}), 400
    try:
        ds = datetime.strptime(date_start, "%Y-%m-%d").date()
        de = datetime.strptime(date_end, "%Y-%m-%d").date()
    except ValueError:
        return jsonify({"error": "Invalid date format"}), 400
    if de < ds:
        return jsonify({"error": "date_end must be on or after date_start"}), 400

    project_key_str = str(PROJECT_KEY).strip()
    if not project_key_str:
        return jsonify({"error": "Missing project key. Save it in Settings first."}), 400
    headers_dict = _get_explicit_request_jira_headers()
    if "Authorization" not in headers_dict:
        return jsonify({"error": "Add Jira email and API token in local settings first."}), 401

    end_exclusive = (de + timedelta(days=1)).strftime("%Y-%m-%d")
    jql = _customer_created_range_jql(date_start, end_exclusive)
    done_status_lookup = {s.strip().lower() for s in done_statuses}

    fields = "summary,status,created,resolutiondate,customfield_10001,issuetype"
    all_issues = []
    seen_issue_keys = set()
    start_at = 0
    next_page_token = None
    page_safety = 0
    while page_safety < 80:
        page_safety += 1
        params = {"jql": jql, "maxResults": 100, "fields": fields}
        if next_page_token:
            params["nextPageToken"] = next_page_token
        else:
            params["startAt"] = start_at
        jira_res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params=params,
        )
        if jira_res.status_code != 200:
            return jsonify({"error": f"Jira API error: {jira_res.text}"}), jira_res.status_code
        data = jira_res.json()
        issues = data.get("issues", [])
        if not issues:
            break
        new_count = 0
        for issue in issues:
            key = issue.get("key")
            if not key or key in seen_issue_keys:
                continue
            seen_issue_keys.add(key)
            all_issues.append(issue)
            new_count += 1
        if new_count == 0:
            break
        next_page_token = data.get("nextPageToken")
        if next_page_token:
            continue
        if len(issues) < 100:
            break
        start_at += 100

    teams = set()
    rows = []
    closed_count = 0
    for issue in all_issues:
        fields_obj = issue.get("fields") or {}
        team_name = _team_name_from_jira_fields(fields_obj)
        teams.add(team_name)
        if team_name_filter and team_name_filter.lower() != "all teams" and team_name.lower() != team_name_filter.lower():
            continue

        status_obj = fields_obj.get("status") or {}
        status_name = (status_obj.get("name") or "").strip()
        is_closed = status_name.lower() in done_status_lookup
        if is_closed:
            closed_count += 1

        issue_type_obj = fields_obj.get("issuetype") or {}
        rows.append({
            "issue_key": issue.get("key") or "",
            "summary": (fields_obj.get("summary") or "").strip(),
            "status": status_name,
            "team_name": team_name,
            "issue_type": (issue_type_obj.get("name") or "").strip(),
            "issue_type_icon": (issue_type_obj.get("iconUrl") or "").strip(),
            "created": fields_obj.get("created"),
            "resolution_date": fields_obj.get("resolutiondate"),
            "is_closed": is_closed,
        })

    total_count = len(rows)
    open_count = max(0, total_count - closed_count)
    closure_rate = round((closed_count / total_count) * 100, 2) if total_count else 0.0
    rows.sort(key=lambda r: ((not r["is_closed"]), r["team_name"].lower(), r["issue_key"]))

    return jsonify({
        "jql": jql,
        "date_start": date_start,
        "date_end": date_end,
        "team_name": team_name_filter or "",
        "done_statuses": done_statuses,
        "teams": sorted(teams, key=lambda x: x.lower()),
        "total_count": total_count,
        "closed_count": closed_count,
        "open_count": open_count,
        "closure_rate": closure_rate,
        "rows": rows,
    })


@app.route("/api/team_productivity/data", methods=["POST"])
def team_productivity_data():
    payload = request.get_json(silent=True) or {}
    date_start = (payload.get("date_start") or "").strip()
    date_end = (payload.get("date_end") or "").strip()
    if not re.match(r"^\d{4}-\d{2}-\d{2}$", date_start) or not re.match(r"^\d{4}-\d{2}-\d{2}$", date_end):
        return jsonify({"error": "date_start and date_end are required (YYYY-MM-DD)"}), 400

    try:
        ds = datetime.strptime(date_start, "%Y-%m-%d").date()
        de = datetime.strptime(date_end, "%Y-%m-%d").date()
    except ValueError:
        return jsonify({"error": "Invalid date format"}), 400
    if de < ds:
        return jsonify({"error": "date_end must be on or after date_start"}), 400

    end_exclusive = (de + timedelta(days=1)).strftime("%Y-%m-%d")
    project_key_str = str(PROJECT_KEY).strip()
    if not project_key_str:
        return jsonify({"error": "Missing project key. Save it in Settings first."}), 400
    headers_dict = _get_explicit_request_jira_headers()
    if "Authorization" not in headers_dict:
        return jsonify({"error": "Add Jira email and API token in local settings first."}), 401

    reported_clause = _reported_by_customer_jql_clause()
    jql = (
        f'project = "{_escape_jql_value(project_key_str)}" AND ({reported_clause}) '
        f'AND created >= "{date_start}" AND created < "{end_exclusive}" '
        f"ORDER BY created DESC"
    )

    fields = "summary,status,created,resolutiondate,customfield_10001,issuetype"
    all_issues = []
    seen_issue_keys = set()
    start_at = 0
    next_page_token = None
    page_safety = 0
    while page_safety < 80:
        page_safety += 1
        params = {
            "jql": jql,
            "maxResults": 100,
            "fields": fields,
            "expand": "changelog",
        }
        if next_page_token:
            params["nextPageToken"] = next_page_token
        else:
            params["startAt"] = start_at
        jira_res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params=params,
        )
        if jira_res.status_code != 200:
            return jsonify({"error": f"Jira API error: {jira_res.text}"}), jira_res.status_code
        data = jira_res.json()
        issues = data.get("issues", [])
        if not issues:
            break
        new_count = 0
        for issue in issues:
            key = issue.get("key")
            if not key or key in seen_issue_keys:
                continue
            seen_issue_keys.add(key)
            all_issues.append(issue)
            new_count += 1
        if new_count == 0:
            break
        next_page_token = data.get("nextPageToken")
        if next_page_token:
            continue
        if len(issues) < 100:
            break
        start_at += 100

    team_hours = defaultdict(list)
    team_unresolved = defaultdict(int)
    team_rows = defaultdict(list)
    all_hours = []

    for issue in all_issues:
        fields_obj = issue.get("fields") or {}
        team = _team_name_from_jira_fields(fields_obj)
        created_dt = _parse_jira_datetime(fields_obj.get("created"))
        if not created_dt:
            continue
        milestone_dt = _first_customer_resolution_milestone_at(issue, fields_obj)
        issue_key = issue.get("key") or ""
        summary = (fields_obj.get("summary") or "").strip()
        status_name = ((fields_obj.get("status") or {}).get("name") or "").strip()
        issue_type_obj = fields_obj.get("issuetype") or {}
        issue_type_name = (issue_type_obj.get("name") or "").strip()
        issue_type_icon = (issue_type_obj.get("iconUrl") or "").strip()

        row = {
            "issue_key": issue_key,
            "summary": summary,
            "status": status_name,
            "issue_type": issue_type_name,
            "issue_type_icon": issue_type_icon,
            "created": fields_obj.get("created"),
            "resolved_at": milestone_dt.isoformat() if milestone_dt else None,
            "resolution_hours": None,
        }
        if milestone_dt:
            delta_h = (milestone_dt - created_dt).total_seconds() / 3600.0
            if delta_h < 0:
                team_unresolved[team] += 1
                row["resolution_note"] = "negative_elapsed_skipped"
            else:
                row["resolution_hours"] = round(delta_h, 2)
                team_hours[team].append(delta_h)
                all_hours.append(delta_h)
        else:
            team_unresolved[team] += 1
        team_rows[team].append(row)

    def _fmt_hours(h):
        if h is None:
            return None
        if h < 48:
            return f"{h:.1f} h"
        return f"{h / 24:.1f} d"

    team_summaries = []
    for team in sorted(team_rows.keys(), key=lambda x: x.lower()):
        hrs = team_hours[team]
        n_res = len(hrs)
        n_un = team_unresolved.get(team, 0)
        avg_h = sum(hrs) / n_res if n_res else None
        team_summaries.append(
            {
                "team_name": team,
                "resolved_count": n_res,
                "unresolved_count": n_un,
                "total_in_range": n_res + n_un,
                "avg_resolution_hours": round(avg_h, 2) if avg_h is not None else None,
                "avg_resolution_display": _fmt_hours(avg_h) if avg_h is not None else "—",
                "tickets": sorted(team_rows[team], key=lambda r: r["issue_key"] or ""),
            }
        )

    overall_avg = sum(all_hours) / len(all_hours) if all_hours else None

    return jsonify(
        {
            "jql": jql,
            "date_start": date_start,
            "date_end": date_end,
            "issue_count": len(all_issues),
            "overall_avg_resolution_hours": round(overall_avg, 2) if overall_avg is not None else None,
            "overall_avg_resolution_display": _fmt_hours(overall_avg) if overall_avg is not None else "—",
            "teams": team_summaries,
        }
    )


@app.route("/api/customer_dashboard/data", methods=["POST"])
def customer_dashboard_data():
    payload = request.get_json(silent=True) or {}
    jql = (payload.get("jql") or "").strip()
    if not jql:
        return jsonify({"error": "JQL is required"}), 400

    project_key_str = str(PROJECT_KEY).strip()
    if not project_key_str:
        return jsonify({"error": "Missing project key. Save it in Settings first."}), 400
    headers_dict = _get_explicit_request_jira_headers()
    if "Authorization" not in headers_dict:
        return jsonify({"error": "Add Jira email and API token in local settings first."}), 401

    # customfield_10001 is commonly the Team field (team[team]) in Jira.
    fields = "summary,description,status,priority,issuetype,created,labels,customfield_10077,customfield_10001,assignee"

    all_issues = []
    seen_issue_keys = set()
    start_at = 0
    next_page_token = None
    page_safety = 0
    while page_safety < 80:  # hard stop safety guard
        page_safety += 1
        params = {"jql": jql, "maxResults": 100, "fields": fields}
        if next_page_token:
            params["nextPageToken"] = next_page_token
        else:
            params["startAt"] = start_at

        # Atlassian now recommends /search/jql.
        jira_res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params=params
        )
        if jira_res.status_code != 200:
            return jsonify({"error": f"Jira API error: {jira_res.text}"}), jira_res.status_code
        data = jira_res.json()
        issues = data.get("issues", [])
        if not issues:
            break

        # Deduplicate by issue key as an extra safety guard.
        new_count = 0
        for issue in issues:
            key = issue.get("key")
            if not key or key in seen_issue_keys:
                continue
            seen_issue_keys.add(key)
            all_issues.append(issue)
            new_count += 1

        # If we didn't get any new issues, stop to avoid repeated-page loops.
        if new_count == 0:
            break
        next_page_token = data.get("nextPageToken")
        if next_page_token:
            continue
        if len(issues) < 100:
            break
        start_at += 100

    rows = []
    unique_customers = set()

    for issue in all_issues:
        fields_obj = issue.get("fields") or {}
        customers = _extract_customer_values(fields_obj.get("customfield_10077"))
        for c in customers:
            unique_customers.add(c)
        labels = fields_obj.get("labels") or []
        launch_blocker = _is_launch_blocker(labels)
        priority_name = ((fields_obj.get("priority") or {}).get("name") or "Unspecified").strip()
        status_name = ((fields_obj.get("status") or {}).get("name") or "Unknown").strip()
        created_raw = (fields_obj.get("created") or "")[:10]
        created_for_chart = created_raw if re.match(r"^\d{4}-\d{2}-\d{2}$", created_raw) else "Unknown"
        assignee = fields_obj.get("assignee") or {}
        team_raw = fields_obj.get("customfield_10001")
        if isinstance(team_raw, dict):
            team_name = (team_raw.get("name") or team_raw.get("value") or "").strip()
        elif isinstance(team_raw, list) and team_raw:
            first = team_raw[0]
            if isinstance(first, dict):
                team_name = (first.get("name") or first.get("value") or "").strip()
            else:
                team_name = str(first).strip()
        elif team_raw:
            team_name = str(team_raw).strip()
        else:
            team_name = "Unspecified"

        rows.append({
            "issue_key": issue.get("key"),
            "summary": fields_obj.get("summary") or "",
            "description": (_adf_to_text(fields_obj.get("description")) or "").strip()[:5000],
            "issue_type": ((fields_obj.get("issuetype") or {}).get("name") or "Task").strip(),
            "status": status_name,
            "priority": priority_name,
            "created_date": created_for_chart if created_for_chart != "Unknown" else "",
            "customers": customers,
            "launch_blocker": launch_blocker,
            "labels": labels,
            "assignee_name": assignee.get("displayName") or "Unassigned",
            "team_name": team_name or "Unspecified"
        })

    return jsonify({
        "rows": rows,
        "customers": sorted(unique_customers),
        "count": len(rows)
    })

@app.route("/api/sprint_dashboard/data", methods=["POST"])
def sprint_dashboard_data():
    payload = request.get_json(silent=True) or {}
    jql = (payload.get("jql") or "").strip()
    if not jql:
        return jsonify({"error": "JQL is required"}), 400

    project_key_str = str(PROJECT_KEY).strip()
    if not project_key_str:
        return jsonify({"error": "Missing project key. Save it in Settings first."}), 400
    if "Authorization" not in dict(HEADERS):
        return jsonify({"error": "Missing Jira credentials. Save them in Settings first."}), 401

    fields = (
        "summary,status,priority,assignee,created,resolutiondate,issuetype,labels,"
        "customfield_10020,customfield_10014,customfield_10077"
    )
    headers_dict = dict(HEADERS)
    all_issues = []
    seen_issue_keys = set()
    start_at = 0
    next_page_token = None
    page_safety = 0
    while page_safety < 80:
        page_safety += 1
        params = {"jql": jql, "maxResults": 100, "fields": fields}
        if next_page_token:
            params["nextPageToken"] = next_page_token
        else:
            params["startAt"] = start_at
        jira_res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params=params
        )
        if jira_res.status_code != 200:
            return jsonify({"error": f"Jira API error: {jira_res.text}"}), jira_res.status_code
        data = jira_res.json()
        issues = data.get("issues", [])
        if not issues:
            break
        new_count = 0
        for issue in issues:
            key = issue.get("key")
            if not key or key in seen_issue_keys:
                continue
            seen_issue_keys.add(key)
            all_issues.append(issue)
            new_count += 1
        if new_count == 0:
            break
        next_page_token = data.get("nextPageToken")
        if next_page_token:
            continue
        if len(issues) < 100:
            break
        start_at += 100

    rows = []
    assignee_metrics = {}
    sprint_metrics = {}
    customer_metrics = {}
    priority_distribution = {}
    status_distribution = {}
    type_distribution = {"Bug": 0, "Story": 0, "Task": 0, "Other": 0}
    kpis = {
        "total": 0,
        "done": 0,
        "open": 0,
        "high_priority": 0,
        "customer_tagged": 0,
        "resolved_customer_tagged": 0,
    }
    unique_assignees = set()
    unique_sprint_names = set()
    unique_customers = set()
    unique_labels = set()

    for issue in all_issues:
        fields_obj = issue.get("fields") or {}
        issue_key = issue.get("key") or ""
        summary = fields_obj.get("summary") or ""
        status_obj = fields_obj.get("status") or {}
        status_name = (status_obj.get("name") or "Unknown").strip()
        status_category = ((status_obj.get("statusCategory") or {}).get("key") or "").strip().lower()
        priority_name = ((fields_obj.get("priority") or {}).get("name") or "Unspecified").strip()
        issue_type_name = ((fields_obj.get("issuetype") or {}).get("name") or "Other").strip()
        labels = [str(v).strip() for v in (fields_obj.get("labels") or []) if str(v).strip()]
        assignee = fields_obj.get("assignee") or {}
        assignee_name = (assignee.get("displayName") or "Unassigned").strip() or "Unassigned"
        created_date = (fields_obj.get("created") or "")[:10]
        resolution_date = (fields_obj.get("resolutiondate") or "")[:10]
        customers = _extract_customer_values(fields_obj.get("customfield_10077"))
        sprint_values = _extract_sprint_values(fields_obj.get("customfield_10020"))
        epic_key = (fields_obj.get("customfield_10014") or "").strip()

        is_done = _is_done_like_status(status_name, status_category)
        is_high_priority = _is_high_priority(priority_name)
        has_customer = len(customers) > 0
        type_bucket = _issue_type_bucket(issue_type_name)
        is_bug = _is_bug_type(issue_type_name)

        if not sprint_values:
            sprint_values = [{"id": None, "name": "No Sprint"}]

        rows.append({
            "issue_key": issue_key,
            "summary": summary,
            "status": status_name,
            "status_category": status_category,
            "priority": priority_name,
            "issue_type": issue_type_name,
            "assignee_name": assignee_name,
            "created_date": created_date if re.match(r"^\d{4}-\d{2}-\d{2}$", created_date or "") else "",
            "resolution_date": resolution_date if re.match(r"^\d{4}-\d{2}-\d{2}$", resolution_date or "") else "",
            "epic_key": epic_key,
            "customers": customers,
            "labels": labels,
            "sprints": sprint_values,
            "is_done": is_done,
            "is_high_priority": is_high_priority,
            "type_bucket": type_bucket,
        })
        for lb in labels:
            unique_labels.add(lb)

        kpis["total"] += 1
        if is_done:
            kpis["done"] += 1
        else:
            kpis["open"] += 1
        if is_high_priority:
            kpis["high_priority"] += 1
        if has_customer:
            kpis["customer_tagged"] += 1
            if is_done:
                kpis["resolved_customer_tagged"] += 1

        priority_distribution[priority_name] = priority_distribution.get(priority_name, 0) + 1
        status_distribution[status_name] = status_distribution.get(status_name, 0) + 1
        type_distribution[type_bucket] = type_distribution.get(type_bucket, 0) + 1

        if assignee_name not in assignee_metrics:
            assignee_metrics[assignee_name] = {
                "assignee_name": assignee_name,
                "worked": 0,
                "done": 0,
                "bug_worked": 0,
                "bug_done": 0,
                "high_priority_worked": 0,
                "ticket_keys": []
            }
        assignee_metrics[assignee_name]["worked"] += 1
        assignee_metrics[assignee_name]["ticket_keys"].append(issue_key)
        if is_done:
            assignee_metrics[assignee_name]["done"] += 1
        if is_bug:
            assignee_metrics[assignee_name]["bug_worked"] += 1
            if is_done:
                assignee_metrics[assignee_name]["bug_done"] += 1
        if is_high_priority:
            assignee_metrics[assignee_name]["high_priority_worked"] += 1
        unique_assignees.add(assignee_name)

        for sprint_item in sprint_values:
            sprint_id = sprint_item.get("id")
            sprint_name = sprint_item.get("name") or "Unknown Sprint"
            sprint_key = f"{sprint_id}" if sprint_id is not None else f"name:{sprint_name}"
            if sprint_key not in sprint_metrics:
                sprint_metrics[sprint_key] = {
                    "sprint_id": sprint_id,
                    "sprint_name": sprint_name,
                    "total": 0,
                    "done": 0,
                    "open": 0,
                    "bug_total": 0,
                    "bug_done": 0,
                    "story_total": 0,
                    "story_done": 0,
                    "task_total": 0,
                    "task_done": 0,
                    "other_total": 0,
                    "other_done": 0,
                }
            sprint_metrics[sprint_key]["total"] += 1
            if is_done:
                sprint_metrics[sprint_key]["done"] += 1
            else:
                sprint_metrics[sprint_key]["open"] += 1
            if type_bucket == "Bug":
                sprint_metrics[sprint_key]["bug_total"] += 1
                if is_done:
                    sprint_metrics[sprint_key]["bug_done"] += 1
            elif type_bucket == "Story":
                sprint_metrics[sprint_key]["story_total"] += 1
                if is_done:
                    sprint_metrics[sprint_key]["story_done"] += 1
            elif type_bucket == "Task":
                sprint_metrics[sprint_key]["task_total"] += 1
                if is_done:
                    sprint_metrics[sprint_key]["task_done"] += 1
            else:
                sprint_metrics[sprint_key]["other_total"] += 1
                if is_done:
                    sprint_metrics[sprint_key]["other_done"] += 1
            unique_sprint_names.add(sprint_name)

        if has_customer:
            for customer in customers:
                if customer not in customer_metrics:
                    customer_metrics[customer] = {"customer": customer, "created": 0, "resolved": 0}
                customer_metrics[customer]["created"] += 1
                if is_done:
                    customer_metrics[customer]["resolved"] += 1
                unique_customers.add(customer)

    assignee_list = list(assignee_metrics.values())
    assignee_list.sort(key=lambda x: (x["done"], x["worked"]), reverse=True)
    for a in assignee_list:
        a["ticket_keys"] = sorted(a["ticket_keys"])

    sprint_list = list(sprint_metrics.values())
    sprint_list.sort(key=lambda x: ((x["sprint_id"] is None), x["sprint_name"].lower()))
    for s in sprint_list:
        s["done_ratio"] = round((s["done"] / s["total"]) * 100, 2) if s["total"] else 0.0

    customer_list = list(customer_metrics.values())
    customer_list.sort(key=lambda x: (x["resolved"], x["created"]), reverse=True)

    return jsonify({
        "kpis": kpis,
        "rows": rows,
        "assignee_metrics": assignee_list,
        "sprint_metrics": sprint_list,
        "customer_metrics": customer_list,
        "priority_distribution": priority_distribution,
        "status_distribution": status_distribution,
        "type_distribution": type_distribution,
        "assignees": sorted(unique_assignees),
        "sprints": sorted(unique_sprint_names),
        "customers": sorted(unique_customers),
        "labels": sorted(unique_labels),
        "count": len(rows)
    })

@app.route("/api/sprint_dashboard/plan_actual", methods=["POST"])
def sprint_dashboard_plan_actual():
    payload = request.get_json(silent=True) or {}
    sprint_id = payload.get("sprint_id")
    selected_labels = payload.get("labels") or []
    if not isinstance(selected_labels, list):
        selected_labels = []
    selected_labels = [str(v).strip() for v in selected_labels if str(v).strip()]
    if sprint_id in (None, ""):
        return jsonify({"error": "sprint_id is required"}), 400
    try:
        sprint_id = int(str(sprint_id).strip())
    except Exception:
        return jsonify({"error": "sprint_id must be a number"}), 400

    project_key_str = str(PROJECT_KEY).strip()
    if not project_key_str:
        return jsonify({"error": "Missing project key. Save it in Settings first."}), 400
    if "Authorization" not in dict(HEADERS):
        return jsonify({"error": "Missing Jira credentials. Save them in Settings first."}), 401

    dataset, err, code = _collect_sprint_plan_actual_dataset(sprint_id, selected_labels)
    if err:
        return jsonify({"error": err}), code
    return jsonify(dataset)


def _collect_sprint_plan_actual_dataset(sprint_id, selected_labels):
    """Fetch sprint-only issues and compute plan-vs-actual plus report metrics."""
    project_key_str = str(PROJECT_KEY).strip()
    if not project_key_str:
        return None, "Missing project key. Save it in Settings first.", 400
    if "Authorization" not in dict(HEADERS):
        return None, "Missing Jira credentials. Save them in Settings first.", 401

    jql = f'project = "{project_key_str}" AND sprint = {int(sprint_id)}'
    if selected_labels:
        escaped = [f'"{_escape_jql_value(v)}"' for v in selected_labels]
        jql += f" AND labels IN ({', '.join(escaped)})"

    headers_dict = dict(HEADERS)
    all_issues = []
    seen_issue_keys = set()
    start_at = 0
    next_page_token = None
    page_safety = 0
    while page_safety < 80:
        page_safety += 1
        params = {
            "jql": jql,
            "maxResults": 100,
            "fields": "summary,status,assignee,issuetype,customfield_10014,customfield_10077"
        }
        if next_page_token:
            params["nextPageToken"] = next_page_token
        else:
            params["startAt"] = start_at
        jira_res = requests.get(f"{JIRA_DOMAIN}/rest/api/3/search/jql", headers=headers_dict, params=params)
        if jira_res.status_code != 200:
            return None, f"Jira API error: {jira_res.text}", jira_res.status_code
        data = jira_res.json()
        issues = data.get("issues", [])
        if not issues:
            break
        new_count = 0
        for issue in issues:
            key = issue.get("key")
            if not key or key in seen_issue_keys:
                continue
            seen_issue_keys.add(key)
            all_issues.append(issue)
            new_count += 1
        if new_count == 0:
            break
        next_page_token = data.get("nextPageToken")
        if next_page_token:
            continue
        if len(issues) < 100:
            break
        start_at += 100

    total = len(all_issues)
    done = 0
    done_by_assignee = {}
    planned_tickets, done_tickets, open_tickets = [], [], []
    total_bugs = total_bugs_fixed = 0
    total_story = total_story_done = 0
    total_task = total_task_done = 0
    epic_ticket_total = epic_ticket_done = 0
    epic_map = {}
    customer_metrics = {}
    customer_fixed_total = 0
    for issue in all_issues:
        f = issue.get("fields") or {}
        status = f.get("status") or {}
        status_name = (status.get("name") or "").strip()
        status_category = ((status.get("statusCategory") or {}).get("key") or "").strip().lower()
        is_done = _is_done_like_status(status_name, status_category)
        assignee_name = ((f.get("assignee") or {}).get("displayName") or "Unassigned").strip() or "Unassigned"
        issue_type_name = ((f.get("issuetype") or {}).get("name") or "Other").strip()
        is_bug = _is_bug_type(issue_type_name)
        is_story = _issue_type_bucket(issue_type_name) == "Story"
        epic_key = str(f.get("customfield_10014") or "").strip()
        customers = _extract_customer_values(f.get("customfield_10077"))
        ticket_row = {
            "issue_key": issue.get("key") or "",
            "summary": f.get("summary") or "",
            "status": status_name or "Unknown",
            "assignee_name": assignee_name,
            "issue_type": issue_type_name,
            "is_bug": is_bug,
            "is_story": is_story,
            "epic_key": epic_key,
            "customers": customers,
        }
        planned_tickets.append(ticket_row)
        if is_done:
            done += 1
            done_by_assignee[assignee_name] = done_by_assignee.get(assignee_name, 0) + 1
            done_tickets.append(ticket_row)
        else:
            open_tickets.append(ticket_row)

        if is_bug:
            total_bugs += 1
            if is_done:
                total_bugs_fixed += 1
        if is_story:
            total_story += 1
            if is_done:
                total_story_done += 1
        if _issue_type_bucket(issue_type_name) == "Task":
            total_task += 1
            if is_done:
                total_task_done += 1

        if epic_key:
            epic_ticket_total += 1
            if is_done:
                epic_ticket_done += 1
            if epic_key not in epic_map:
                epic_map[epic_key] = {"total": 0, "done": 0}
            epic_map[epic_key]["total"] += 1
            if is_done:
                epic_map[epic_key]["done"] += 1

        for c in customers:
            if c not in customer_metrics:
                customer_metrics[c] = {"customer": c, "created": 0, "resolved": 0}
            customer_metrics[c]["created"] += 1
            if is_done:
                customer_metrics[c]["resolved"] += 1
                customer_fixed_total += 1

    open_count = max(0, total - done)
    done_ratio = round((done / total) * 100, 2) if total else 0.0
    top_done = [{"assignee_name": k, "done": v} for k, v in sorted(done_by_assignee.items(), key=lambda kv: kv[1], reverse=True)]
    customer_summary = sorted(customer_metrics.values(), key=lambda x: (x["resolved"], x["created"]), reverse=True)
    total_epics = len(epic_map)
    done_epics = len([e for e in epic_map.values() if e["total"] > 0 and e["total"] == e["done"]])

    return {
        "sprint_id": int(sprint_id),
        "selected_labels": selected_labels,
        "total": total,
        "done": done,
        "open": open_count,
        "done_ratio": done_ratio,
        "done_by_assignee": top_done,
        "tickets": {"planned": planned_tickets, "done": done_tickets, "open": open_tickets},
        "report_metrics": {
            "total_bugs": total_bugs,
            "total_bugs_fixed": total_bugs_fixed,
            "total_story": total_story,
            "total_story_done": total_story_done,
            "total_task": total_task,
            "total_task_done": total_task_done,
            "total_epics": total_epics,
            "done_epics": done_epics,
            "epic_ticket_total": epic_ticket_total,
            "epic_ticket_done": epic_ticket_done,
            "customer_issues_fixed": customer_fixed_total,
        },
        "customer_summary": customer_summary,
    }, None, 200


def _sprint_report_bar_chart_b64(labels, values, colors, title):
    fig, ax = plt.subplots(figsize=(7.2, 2.8), dpi=120)
    bars = ax.bar(labels, values, color=colors)
    ax.set_title(title, fontsize=10)
    ax.set_axisbelow(True)
    ax.grid(axis='y', alpha=0.2)
    # Rotate x labels slightly for readability when names are long/crowded.
    if len(labels) > 4:
        ax.tick_params(axis='x', labelrotation=25)
        for tick in ax.get_xticklabels():
            tick.set_horizontalalignment('right')
    for idx, b in enumerate(bars):
        ax.text(b.get_x() + b.get_width() / 2.0, b.get_height() + 0.1, str(values[idx]), ha='center', va='bottom', fontsize=8)
    plt.tight_layout()
    stream = io.BytesIO()
    fig.savefig(stream, format="png")
    plt.close(fig)
    return base64.b64encode(stream.getvalue()).decode()


@app.route("/api/sprint_dashboard/report_pdf", methods=["POST"])
def sprint_dashboard_report_pdf():
    payload = request.get_json(silent=True) or {}
    sprint_id = payload.get("sprint_id")
    selected_labels = payload.get("labels") or []
    if not isinstance(selected_labels, list):
        selected_labels = []
    selected_labels = [str(v).strip() for v in selected_labels if str(v).strip()]
    if sprint_id in (None, ""):
        return jsonify({"error": "sprint_id is required"}), 400
    try:
        sprint_id = int(str(sprint_id).strip())
    except Exception:
        return jsonify({"error": "sprint_id must be a number"}), 400

    dataset, err, code = _collect_sprint_plan_actual_dataset(sprint_id, selected_labels)
    if err:
        return jsonify({"error": err}), code

    labels_text = ", ".join(selected_labels) if selected_labels else "All labels"
    metrics = dataset.get("report_metrics", {})
    customer_summary = dataset.get("customer_summary", [])[:15]
    planned_tickets = (dataset.get("tickets") or {}).get("planned") or []
    done_tickets = (dataset.get("tickets") or {}).get("done") or []
    planned_story = [t for t in planned_tickets if t.get("is_story")]
    done_story = [t for t in done_tickets if t.get("is_story")]
    planned_bugs = [t for t in planned_tickets if t.get("is_bug")]
    done_bugs = [t for t in done_tickets if t.get("is_bug")]
    plan_chart = _sprint_report_bar_chart_b64(
        ["Planned", "Done", "Pending"],
        [dataset.get("total", 0), dataset.get("done", 0), dataset.get("open", 0)],
        ["#5b7ee5", "#43a567", "#e76f51"],
        "Plan vs Actual"
    )
    customer_chart_rows = customer_summary[:8]
    customer_chart = _sprint_report_bar_chart_b64(
        [r["customer"][:14] + ("..." if len(r["customer"]) > 14 else "") for r in customer_chart_rows] or ["No data"],
        [r["resolved"] for r in customer_chart_rows] or [0],
        ["#0ea5e9"] * max(1, len(customer_chart_rows)),
        "Customer Issues Fixed"
    )

    customer_rows_html = "".join(
        f"<tr><td>{(r.get('customer') or '').replace('<','&lt;').replace('>','&gt;')}</td><td>{r.get('created',0)}</td><td>{r.get('resolved',0)}</td></tr>"
        for r in customer_summary
    ) or "<tr><td colspan='3'>No customer-tagged issues in this sprint scope.</td></tr>"

    def _ticket_rows_html(rows):
        if not rows:
            return "<tr><td colspan='2'>No tickets in this category.</td></tr>"
        out = []
        for t in rows:
            key = (t.get("issue_key") or "").replace("<", "&lt;").replace(">", "&gt;")
            summary = (t.get("summary") or "").replace("<", "&lt;").replace(">", "&gt;")
            out.append(f"<tr><td>{key}</td><td>{summary}</td></tr>")
        return "".join(out)

    html_content = f"""
    <html>
    <head>
      <meta charset="utf-8" />
      <style>
        body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11px; color: #1f2937; }}
        h1 {{ font-size: 18px; margin: 0 0 8px; }}
        h2 {{ font-size: 13px; margin: 14px 0 6px; }}
        .meta {{ color: #475569; margin-bottom: 8px; }}
        .grid {{ width: 100%; }}
        .kpi {{ border: 1px solid #dbe3ef; border-radius: 8px; padding: 6px; margin: 4px; }}
        .kpi-title {{ font-size: 10px; color: #64748b; }}
        .kpi-value {{ font-size: 15px; font-weight: bold; margin-top: 2px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 6px; }}
        th, td {{ border: 1px solid #e2e8f0; padding: 5px; text-align: left; }}
        th {{ background: #f8fafc; color: #334155; }}
        .tbl-lite {{ border: 1px solid #e2e8f0; background: #ffffff; }}
        .tbl-lite th {{ background: #f8fafc; }}
        .tbl-plan {{ border: 1px solid #cbd5e1; background: #f8fbff; }}
        .tbl-plan th {{ background: #eff6ff; }}
        .tbl-actual {{ border: 1px solid #bbf7d0; background: #f0fdf4; }}
        .tbl-actual th {{ background: #dcfce7; }}
        .tbl-bug {{ border: 1px solid #fecaca; background: #fef2f2; }}
        .tbl-bug th {{ background: #fee2e2; }}
        .tbl-bugdone {{ border: 1px solid #fed7aa; background: #fff7ed; }}
        .tbl-bugdone th {{ background: #ffedd5; }}
        .chart {{ margin: 8px 0; text-align: center; }}
      </style>
    </head>
    <body>
      <h1>Sprint Plan vs Actual Report</h1>
      <div class="meta">Sprint ID: {sprint_id} | Labels: {labels_text} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}</div>
      <table class="grid">
        <tr>
          <td class="kpi"><div class="kpi-title">Planned Tickets</div><div class="kpi-value">{dataset.get('total',0)}</div></td>
          <td class="kpi"><div class="kpi-title">Total Done</div><div class="kpi-value">{dataset.get('done',0)}</div></td>
          <td class="kpi"><div class="kpi-title">Pending</div><div class="kpi-value">{dataset.get('open',0)}</div></td>
          <td class="kpi"><div class="kpi-title">Completion</div><div class="kpi-value">{dataset.get('done_ratio',0)}%</div></td>
        </tr>
        <tr>
          <td class="kpi"><div class="kpi-title">Total Bugs</div><div class="kpi-value">{metrics.get('total_bugs',0)}</div></td>
          <td class="kpi"><div class="kpi-title">Bugs Fixed</div><div class="kpi-value">{metrics.get('total_bugs_fixed',0)}</div></td>
          <td class="kpi"><div class="kpi-title">Total Stories</div><div class="kpi-value">{metrics.get('total_story',0)}</div></td>
          <td class="kpi"><div class="kpi-title">Stories Done</div><div class="kpi-value">{metrics.get('total_story_done',0)}</div></td>
        </tr>
        <tr>
          <td class="kpi"><div class="kpi-title">Total Tasks</div><div class="kpi-value">{metrics.get('total_task',0)}</div></td>
          <td class="kpi"><div class="kpi-title">Tasks Done</div><div class="kpi-value">{metrics.get('total_task_done',0)}</div></td>
          <td class="kpi"><div class="kpi-title">Customer Issues Fixed</div><div class="kpi-value">{metrics.get('customer_issues_fixed',0)}</div></td>
        </tr>
      </table>

      <h2>Plan vs Actual Graph</h2>
      <div class="chart"><img src="data:image/png;base64,{plan_chart}" width="520" /></div>

      <h2>Customer-Related Graph: Issues Fixed</h2>
      <div class="chart"><img src="data:image/png;base64,{customer_chart}" width="520" /></div>

      <h2>Customer Issue Summary</h2>
      <table class="tbl-lite">
        <thead><tr><th>Customer</th><th>Total Issues</th><th>Fixed Issues</th></tr></thead>
        <tbody>{customer_rows_html}</tbody>
      </table>
      <p>Total customer issues fixed in this sprint scope: <strong>{metrics.get('customer_issues_fixed',0)}</strong></p>

      <h2>Story Plan vs Actual</h2>
      <p>Planned Stories: <strong>{len(planned_story)}</strong> | Done Stories: <strong>{len(done_story)}</strong></p>
      <table class="tbl-plan">
        <thead><tr><th colspan="2">Story Planned Tickets</th></tr><tr><th>Ticket Number</th><th>Ticket Title</th></tr></thead>
        <tbody>{_ticket_rows_html(planned_story[:120])}</tbody>
      </table>
      <table class="tbl-actual">
        <thead><tr><th colspan="2">Story Done Tickets</th></tr><tr><th>Ticket Number</th><th>Ticket Title</th></tr></thead>
        <tbody>{_ticket_rows_html(done_story[:120])}</tbody>
      </table>

      <h2>Bug Plan vs Actual</h2>
      <p>Planned Bugs: <strong>{len(planned_bugs)}</strong> | Fixed Bugs: <strong>{len(done_bugs)}</strong></p>
      <table class="tbl-bug">
        <thead><tr><th colspan="2">Bug Planned Tickets</th></tr><tr><th>Ticket Number</th><th>Ticket Title</th></tr></thead>
        <tbody>{_ticket_rows_html(planned_bugs[:120])}</tbody>
      </table>
      <table class="tbl-bugdone">
        <thead><tr><th colspan="2">Bug Fixed Tickets</th></tr><tr><th>Ticket Number</th><th>Ticket Title</th></tr></thead>
        <tbody>{_ticket_rows_html(done_bugs[:120])}</tbody>
      </table>
    </body>
    </html>
    """

    output = io.BytesIO()
    pdf_status = pisa.CreatePDF(io.BytesIO(html_content.encode("utf-8")), dest=output)
    if pdf_status.err:
        return jsonify({"error": "Failed to generate PDF"}), 500
    output.seek(0)
    filename = f"sprint_{sprint_id}_plan_actual_report.pdf"
    return send_file(output, as_attachment=True, download_name=filename, mimetype="application/pdf")


def _claude_customer_dashboard_insights_json(summary, anthropic_api_key, model="claude-sonnet-4-20250514"):
    """
    Call Claude with a compact dashboard summary; return a list of insight dicts
    {area, finding, suggestion, priority} or None on failure.
    """
    if not summary or not anthropic_api_key:
        return None
    try:
        payload_str = json.dumps(summary, default=str, ensure_ascii=False)
    except (TypeError, ValueError):
        return None
    if len(payload_str) > 100_000:
        payload_str = payload_str[:100_000] + "\n…[truncated]"

    expected = str((summary or {}).get("expected_insights") or "").strip()
    expected_block = (
        f"User expectation for this run:\n{expected}\n\n"
        if expected
        else "User expectation for this run:\nNot provided. Choose the most useful defaults.\n\n"
    )

    prompt = (
        "You are an engineering program analyst. The following JSON is a snapshot of a Jira customer "
        "dashboard: KPIs, distributions by team, customer, status, priority, and a sample of issues.\n"
        "Produce actionable insights a PM or eng lead can use the same day.\n\n"
        "Return ONLY valid JSON (no markdown fences, no surrounding prose) in this exact shape:\n"
        '{"insights":[{"area":"short category label",'
        '"finding":"1-2 sentences; cite numbers from the data when possible",'
        '"suggestion":"concrete next step; reference issue keys only if they appear in the sample",'
        '"priority":"High"}]}\n\n'
        "Rules:\n"
        "- 6 to 10 insight rows.\n"
        "- priority must be one of: High, Medium, Low.\n"
        "- MUST include one row with area exactly 'Overall summary'.\n"
        "- MUST include one row with area exactly 'Current focus areas'.\n"
        "- MUST include one row with area exactly 'Next issues to resolve'. In this row, include 3-6 issue keys from sample_issues and why each should be worked next.\n"
        "- Vary 'area' across rows: e.g. Launch risk, High priority, Team load, Customer concentration, "
        "Status bottlenecks, Unassigned/ownership, Time window / inflow, Data hygiene — only where the data supports it.\n"
        "- If the sample is too small to infer something, say that in the finding and suggest what data to collect.\n"
        "- Do not invent customers, teams, or ticket keys that are not in the JSON.\n\n"
        f"{expected_block}"
        f"JSON snapshot:\n{payload_str}"
    )
    try:
        res = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": anthropic_api_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": model,
                "max_tokens": 4096,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=120,
        )
        if not res.ok:
            print(f"Claude insights HTTP {res.status_code}: {res.text[:500]}")
            return None
        data = res.json()
        text = "".join(b.get("text", "") for b in (data.get("content") or []) if b.get("type") == "text")
        match = re.search(r"\{[\s\S]*\}", text)
        if not match:
            return None
        parsed = json.loads(match.group(0))
        raw = parsed.get("insights")
        if not isinstance(raw, list):
            return None
        out = []
        for item in raw:
            if not isinstance(item, dict):
                continue
            area = str(item.get("area") or "General").strip() or "General"
            finding = str(item.get("finding") or "").strip()
            suggestion = str(item.get("suggestion") or "").strip()
            pr = str(item.get("priority") or "Medium").strip()
            if pr not in ("High", "Medium", "Low"):
                pr = "Medium"
            if not finding and not suggestion:
                continue
            out.append(
                {
                    "area": area[:200],
                    "finding": finding[:1200],
                    "suggestion": suggestion[:1200],
                    "priority": pr,
                }
            )
        return out or None
    except Exception as e:
        print(f"Claude customer dashboard insights error: {e}")
        return None


def _fallback_customer_dashboard_insights(summary, reason=None):
    summary = summary or {}
    total = int(summary.get("total_issues") or 0)
    high = int(summary.get("high_priority_count") or 0)
    lb = int(summary.get("launch_blocker_count") or 0)
    ratio = float(summary.get("launch_blocker_ratio") or 0)
    by_status = summary.get("by_status") or {}
    by_team = summary.get("by_team") or {}
    by_customer = summary.get("by_customer") or {}
    expected = str(summary.get("expected_insights") or "").strip()

    top_status = "Unknown"
    top_status_count = 0
    if isinstance(by_status, dict) and by_status:
        top_status, top_status_count = max(by_status.items(), key=lambda kv: kv[1] or 0)
    top_team = "Unspecified"
    top_team_count = 0
    if isinstance(by_team, dict) and by_team:
        top_team, top_team_count = max(by_team.items(), key=lambda kv: kv[1] or 0)
    top_customer = "Unspecified"
    top_customer_count = 0
    if isinstance(by_customer, dict) and by_customer:
        top_customer, top_customer_count = max(by_customer.items(), key=lambda kv: kv[1] or 0)

    lb_pct = round(ratio * 100, 1) if ratio <= 1 else round(ratio, 1)
    high_pct = round((high / total) * 100, 1) if total else 0

    rows = [
        {
            "area": "Dataset health",
            "finding": f"Current filtered dataset has {total} issues.",
            "suggestion": "Use this run to identify immediate risks and one weekly trend to track.",
            "priority": "Low",
        },
        {
            "area": "High priority load",
            "finding": f"{high} issues are high-priority ({high_pct}% of total).",
            "suggestion": "Create a short burn-down list for P0/P1 issues and assign explicit owners.",
            "priority": "High" if high_pct >= 20 else "Medium",
        },
        {
            "area": "Launch blocker risk",
            "finding": f"{lb} issues are marked as launch blockers ({lb_pct}%).",
            "suggestion": "Review blocker criteria and run a daily unblock checkpoint until ratio trends down.",
            "priority": "High" if lb > 0 else "Medium",
        },
        {
            "area": "Status concentration",
            "finding": f"Top status bucket is '{top_status}' with {top_status_count} issues.",
            "suggestion": "Inspect this status bucket for aging tickets and define exit criteria.",
            "priority": "Medium",
        },
        {
            "area": "Team distribution",
            "finding": f"Top team by ticket share is '{top_team}' with {top_team_count} issues.",
            "suggestion": "Check whether this team has balanced ownership and if parallelization is needed.",
            "priority": "Medium",
        },
        {
            "area": "Customer concentration",
            "finding": f"Top customer in this view is '{top_customer}' with {top_customer_count} issues.",
            "suggestion": "Confirm roadmap alignment with this customer and capture recurrent root causes.",
            "priority": "Medium",
        },
    ]

    if expected:
        rows.insert(
            1,
            {
                "area": "Requested focus",
                "finding": f"Expected insight focus: {expected[:300]}",
                "suggestion": "Use this focus to prioritize which sections of the dashboard you review first.",
                "priority": "Medium",
            },
        )

    if reason:
        rows.insert(
            0,
            {
                "area": "AI availability",
                "finding": f"Claude output unavailable for this request ({reason}). Fallback analytics were generated.",
                "suggestion": "Verify API key/network and regenerate for richer narrative while keeping this table as baseline.",
                "priority": "Low",
            },
        )

    return rows[:10]


@app.route("/api/customer_dashboard/ai_status", methods=["GET"])
def customer_dashboard_ai_status():
    """Whether Anthropic (Claude) key is set in app_config — no secret exposed."""
    key = (_get_app_config_value("anthropic_api_key") or "").strip()
    return jsonify({"claude_configured": bool(key)})


@app.route("/api/customer_dashboard/insights", methods=["POST"])
def customer_dashboard_insights():
    """
    Generate table-style AI insights from a client-built summary. Uses server-side anthropic_api_key.
    """
    payload = request.get_json(silent=True) or {}
    summary = payload.get("summary")
    if not summary or not isinstance(summary, dict):
        fallback = _fallback_customer_dashboard_insights({}, reason="invalid request payload")
        return jsonify({"insights": fallback, "fallback": True})

    anthropic_key = (_get_app_config_value("anthropic_api_key") or "").strip()
    if not anthropic_key:
        fallback = _fallback_customer_dashboard_insights(summary, reason="missing Anthropic API key")
        return jsonify({"insights": fallback, "fallback": True})

    insights = _claude_customer_dashboard_insights_json(summary, anthropic_key)
    if not insights:
        fallback = _fallback_customer_dashboard_insights(summary, reason="Claude response parse/transport error")
        return jsonify({"insights": fallback, "fallback": True})

    return jsonify({"insights": insights})


@app.route("/api/customer_dashboard/wordcloud", methods=["POST"])
def customer_dashboard_wordcloud():
    """
    Given a list of ticket summaries for a company, ask Claude to extract
    the most meaningful keywords/themes and return them with importance weights
    for rendering a word cloud.
    Returns: { words: [{text, weight}, ...] }
    """
    payload = request.get_json(silent=True) or {}
    tickets  = payload.get("tickets") or []
    context  = str(payload.get("context") or "this company").strip()
    model    = payload.get("model") or "claude-sonnet-4-20250514"

    anthropic_key = (_get_app_config_value("anthropic_api_key") or "").strip()
    if not anthropic_key:
        return jsonify({"error": "Anthropic API key not configured."}), 400
    if not tickets:
        return jsonify({"error": "No tickets provided."}), 400

    summaries = [str(t.get("summary") or "").strip() for t in tickets if t.get("summary")]
    summaries = summaries[:150]   # cap to avoid token overflow
    blob = "\n".join(f"- {s}" for s in summaries)

    prompt = (
        f"You are analysing Jira ticket summaries for {context}.\n"
        "Extract the 30-50 most meaningful technical keywords, feature areas, product domains, "
        "or recurring themes from the list of ticket summaries below.\n"
        "Assign each keyword an integer importance weight from 1 (low) to 10 (high) "
        "based on how frequently and prominently it appears.\n"
        "Return ONLY valid JSON — an array of objects with 'text' and 'weight' keys, e.g.:\n"
        '[{"text":"payments","weight":9},{"text":"bug fix","weight":6}]\n\n'
        "Rules:\n"
        "- No markdown fences, no extra prose — raw JSON array only.\n"
        "- Use concise 1-3 word phrases; no full sentences.\n"
        "- Avoid generic words like 'ticket', 'jira', 'issue', 'task', 'please'.\n\n"
        f"Ticket summaries:\n{blob}"
    )

    try:
        res = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": anthropic_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": model,
                "max_tokens": 1024,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=60,
        )
        if not res.ok:
            return jsonify({"error": f"Claude HTTP {res.status_code}"}), 502
        body = res.json()
        raw  = (body.get("content") or [{}])[0].get("text", "").strip()
        raw  = raw.strip("`").strip()
        if raw.startswith("json"):
            raw = raw[4:].strip()
        words = json.loads(raw)
        if not isinstance(words, list):
            raise ValueError("Expected list")
        # Normalise
        words = [{"text": str(w.get("text","")).strip(), "weight": int(w.get("weight",1))}
                 for w in words if w.get("text")]
        return jsonify({"words": words})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# =========================
# FETCH SCOREBOARD DATA
# =========================
@app.route("/api/scoreboard_data", methods=["POST"])
def scoreboard_data():
    print("DEBUG: Entered scoreboard_data")
    data = request.json
    range_type = data.get("range", "day")
    
    # Explicitly cast LocalProxy objects
    project_key_str = str(PROJECT_KEY)
    headers_dict = dict(HEADERS)
    
    # Use quotes for safety
    jql = f'project = "{project_key_str}" AND statusCategory = "Done"'
    
    if range_type == "day":
        jql += " AND resolutiondate >= startOfDay()"
    elif range_type == "yesterday":
        jql += " AND resolutiondate >= startOfDay(-1d) AND resolutiondate < startOfDay()"
    elif range_type == "week":
        jql += " AND resolutiondate >= startOfWeek()"
    elif range_type == "month":
        jql += " AND resolutiondate >= startOfMonth()"
    elif range_type == "year":
        jql += " AND resolutiondate >= startOfYear()"
        
    # If custom date is passed
    if "startDate" in data and data["startDate"]:
        jql += f" AND resolutiondate >= '{data['startDate']}'"
    if "endDate" in data and data["endDate"]:
        jql += f" AND resolutiondate <= '{data['endDate']}'"

    # Fetch ALL completed issues for the period
    all_issues = []
    start_at = 0
    max_results = 100
    
    print(f"DEBUG: Executing JQL: {jql}")
    
    # payload setup
    payload = {
        "jql": jql,
        "maxResults": int(max_results),
        "startAt": int(start_at),
        "fields": ["assignee", "status", "resolutiondate"]
    }
    
    while True:
        # Using GET /rest/api/3/search/jql as required by the error message.
        # GET requests pass parameters in the URL query string.
        params = {
            "jql": jql,
            "maxResults": max_results,
            "startAt": start_at,
            "fields": "assignee,status,resolutiondate" # Comma separated for GET
        }
        
        res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params=params
        ).json()
        
        # print(f"DEBUG: Jira Response Keys: {res.keys()}")
        if "errorMessages" in res:
            print(f"DEBUG: Error from Jira: {res['errorMessages']}")
        
        issues = res.get("issues", [])
        print(f"DEBUG: Found {len(issues)} issues (Total: {res.get('total', 0)})")
        all_issues.extend(issues)
        
        if start_at + len(issues) >= res.get("total", 0):
            break
        start_at += len(issues)

    # Aggregate by assignee
    stats = {}
    for i in all_issues:
        assignee = i["fields"].get("assignee")
        if assignee:
            aid = assignee["accountId"]
            name = assignee["displayName"]
            if aid not in stats:
                stats[aid] = {"id": aid, "name": name, "count": 0, "avatar": assignee["avatarUrls"]["48x48"]}
            stats[aid]["count"] += 1
            
            
    return jsonify(list(stats.values()))

# =========================
# SPRINT DASHBOARD API
# =========================
@app.route("/api/dashboard/sprint_stats", methods=["POST"])
def sprint_stats():
    data = request.json
    team_ids = data.get("team_id")
    sprint_id = data.get("sprint_id") # Optional Jira Sprint ID
    report_id = data.get("report_id")
    production_only = data.get("production_only", False)
    force_refresh = data.get("force_refresh", False)
    custom_jql = data.get("custom_jql") # NEW: Support for Custom Reports
    
    if not team_ids and not sprint_id and not custom_jql and not report_id:
        return jsonify({"error": "Team, Sprint, Report, or Custom JQL is required"}), 400
        
    if team_ids and not isinstance(team_ids, list):
        team_ids = [team_ids]

    members = []
    member_ids_str = ""
    
    # 1. Get Team Members for all selected teams (IF TEAMS SELECTED)
    if team_ids:
        conn, cursor = get_db_connection()
        if not conn: return jsonify({"error": "Database error"}), 500
        
        placeholders = ', '.join(['%s'] * len(team_ids))
        cursor.execute(f"SELECT account_id, display_name, avatar_url FROM team_members WHERE team_id IN ({placeholders})", tuple(team_ids))
        
        # Use a dictionary to ensure members are unique across multiple teams
        members_map = {}
        for r in cursor.fetchall():
            members_map[r[0]] = {"id": r[0], "name": r[1], "avatar": r[2]}
        
        members = list(members_map.values())
        conn.close()
        
        if not members:
            # If teams selected but no members, we might return empty or error?
            # Let's return error as "No members found for this team" implies configuration issue.
            return jsonify({"error": "No members found for this team"}), 400
            
        member_ids = [m["id"] for m in members]
        member_ids_str = ", ".join([f'"{mid}"' for mid in member_ids])
    
    headers_dict = dict(HEADERS)
    project_key_str = str(PROJECT_KEY)

    if report_id and not custom_jql:
        conn, cursor = get_db_connection()
        if not conn:
            return jsonify({"error": "Database error"}), 500
        cursor.execute("SELECT jql FROM custom_reports WHERE id = %s", (report_id,))
        report_row = cursor.fetchone()
        conn.close()
        if not report_row:
            return jsonify({"error": "Report not found"}), 404
        custom_jql = report_row[0]

    # 2. Build JQL
    if custom_jql:
        # If Custom JQL is provided, use it directly (but append Prod filter if needed)
        jql = custom_jql
        if production_only:
            jql = f'({jql}) AND "platform[checkboxes]" = PRODUCTION'
    else:
        # Standard logic
        # We must have project key
        jql = f'project = "{project_key_str}"'
        
        # Optional Sprint filter
        if sprint_id:
            jql += f' AND sprint = {sprint_id}'
        
        # Optional Assignee filter (Only if teams selected)
        if member_ids_str:
            jql += f' AND assignee IN ({member_ids_str})'
        
        # Optional Production filter
        if production_only:
            jql += ' AND "platform[checkboxes]" = PRODUCTION'
    
    try:
        res = get_jira_cached(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params={
                "jql": jql, 
                "maxResults": 1000, 
                "fields": "summary,status,priority,assignee,issuetype"
            },
            force_refresh=force_refresh
        )
        
        if "errorMessages" in res:
            return jsonify({"error": res["errorMessages"]}), 400
            
        all_issues = res.get("issues", [])
        
        # 3. Calculate Stats — completed = Jira statusCategory "done" or named terminal / QA-ready statuses
        done_issues = []
        active_issues = []
        for i in all_issues:
            st = i["fields"]["status"]
            status_name = (st.get("name") or "").strip()
            status_cat = (st.get("statusCategory") or {}).get("key") or ""

            if _is_done_like_status(status_name, status_cat):
                done_issues.append(i)
            else:
                active_issues.append(i)
        
        total_solved = len(done_issues)
        bugs_solved = len([i for i in done_issues if i["fields"]["issuetype"]["name"].lower() == "bug"])
        
        # User Performance
        # If we have members (Team selected), we calculate for them.
        # If NO team selected (Custom JQL or just Sprint), we calculate for ALL assignees found in the tickets.
        
        performance = {}
        
        if members:
            # Initialize with team members
            for m in members:
                performance[m["id"]] = {"name": m["name"], "avatar": m["avatar"], "solved": 0, "active": 0}
        
        for i in all_issues:
            assignee = i["fields"].get("assignee")
            if assignee:
                aid = assignee["accountId"]
                
                # If no specific team selected, add any assignee encountered
                if not members and aid not in performance:
                     performance[aid] = {
                         "name": assignee["displayName"], 
                         "avatar": assignee["avatarUrls"]["48x48"], 
                         "solved": 0, 
                         "active": 0
                     }
                
                if aid in performance:
                    st = i["fields"]["status"]
                    status_name = (st.get("name") or "").strip()
                    status_cat = (st.get("statusCategory") or {}).get("key") or ""

                    if _is_done_like_status(status_name, status_cat):
                        performance[aid]["solved"] += 1
                    else:
                        performance[aid]["active"] += 1
                        
        # Today's Work: Active tickets with their current status
        today_work = []
        for i in active_issues:
            today_work.append({
                "key": i["key"],
                "summary": i["fields"]["summary"],
                "assignee": (i["fields"].get("assignee") or {}).get("displayName", "Unassigned"),
                "status": i["fields"]["status"]["name"],
                "type_icon": i["fields"]["issuetype"].get("iconUrl")
            })
            
        # Group status distribution for chart
        status_dist = {}
        for i in all_issues:
            s = i["fields"]["status"]["name"]
            status_dist[s] = status_dist.get(s, 0) + 1
            
        # Group type distribution
        type_dist = {}
        for i in all_issues:
            t = i["fields"]["issuetype"]["name"]
            type_dist[t] = type_dist.get(t, 0) + 1
            
        # Group priority distribution
        pri_dist = {}
        for i in all_issues:
            p = i["fields"]["priority"]["name"]
            pri_dist[p] = pri_dist.get(p, 0) + 1
            
        # NEW: Fetch local tracking data for the table
        # If we have a sprint_id, we fetch data for that sprint.
        # If we have a custom_jql (no sprint_id), we might want to fetch data based on issue keys?
        # But our DB structure links tickets to sprint_id.
        # So for Custom Reports (which might span multiple sprints or have no sprint), 
        # local tracking data might be tricky if it depends on sprint_id.
        
        # Current DB Schema: sprint_tickets (sprint_id, issue_key, ...)
        # If we don't have a sprint_id, we can't easily fetch local data unless we fetch by issue_key.
        # Let's try to fetch by issue_key for the issues returned by JQL.
        
        tracking_data = {}
        issue_keys = [i["key"] for i in all_issues]
        
        if issue_keys:
            conn, cursor = get_db_connection(dictionary=True)
            if conn:
                # If we have a sprint_id, we can filter by it to be precise.
                # If not (Custom Report), we fetch the LATEST entry for that ticket? 
                # Or just any entry? A ticket might be in multiple sprints over time.
                # Ideally, we should show the data for the current context.
                # If custom report, maybe we just show the latest known status?
                
                if sprint_id:
                    cursor.execute("""
                        SELECT issue_key, pr_raised, pr_merged, deploy_status, qa_assignee, qa_status, bugs_found, requirements_clear, completed, is_flagged, comment
                        FROM sprint_tickets 
                        WHERE sprint_id = %s
                    """, (sprint_id,))
                else:
                    # For Custom Reports: Fetch data for these keys.
                    # Problem: A ticket might appear multiple times in sprint_tickets table (different sprints).
                    # We should probably pick the most recent one (highest sprint_id or created_at? we don't have created_at).
                    # Let's assume fetching by issue_key is "okay" and if multiple exist, we pick one (or latest).
                    # Actually, let's just fetch all matching keys and overwrite (so we get one of them).
                    # Or better: `WHERE issue_key IN (...)`
                    
                    placeholders = ', '.join(['%s'] * len(issue_keys))
                    cursor.execute(f"""
                        SELECT issue_key, pr_raised, pr_merged, deploy_status, qa_assignee, qa_status, bugs_found, requirements_clear, completed, is_flagged, comment
                        FROM sprint_tickets 
                        WHERE issue_key IN ({placeholders})
                        ORDER BY id DESC
                    """, tuple(issue_keys))
                    
                    # Note: This might return multiple rows per key. The dictionary comprehension below 
                    # will overwrite earlier ones with later ones (due to ORDER BY id DESC, later ones come first? No, wait).
                    # If we iterate, the last one processed wins. 
                    # If we ORDER BY id ASC, the last one (newest) wins.
                    
                rows = cursor.fetchall()
                # If using IN clause with multiple rows per key, we want the latest.
                # If we ORDER BY id ASC, the loop will set key -> val, then update key -> newer_val.
                # So existing logic works fine if we order correctly or just accept "a" value.
                
                for r in rows:
                    tracking_data[r["issue_key"]] = {
                        "pr_raised": bool(r["pr_raised"]),
                        "pr_merged": bool(r["pr_merged"]),
                        "deploy_status": r["deploy_status"],
                        "qa_assignee": r["qa_assignee"],
                        "qa_status": r["qa_status"],
                        "bugs_found": r["bugs_found"],
                        "requirements_clear": r["requirements_clear"] or "No",
                        "completed": bool(r["completed"]),
                        "is_flagged": bool(r["is_flagged"]),
                        "comment": r["comment"] or ""
                    }
                conn.close()

        # Merge local data into all issues for the tracking table
        tracking_issues = []
        extra_bug_keys = set()
        for i in all_issues:
            key = i["key"]
            local = tracking_data.get(key, {
                "pr_raised": False,
                "pr_merged": False,
                "deploy_status": "N/A",
                "qa_assignee": "",
                "qa_status": "Pending",
                "bugs_found": "",
                "requirements_clear": "No",
                "completed": False,
                "is_flagged": False,
                "comment": ""
            })
            
            # Collect bug keys to fetch their status later if they are not in all_issues
            if local["bugs_found"]:
                for b_key in local["bugs_found"].split(","):
                    b_key = b_key.strip()
                    if b_key and b_key not in [iss["key"] for iss in all_issues]:
                        extra_bug_keys.add(b_key)

            tracking_issues.append({
                "key": key,
                "summary": i["fields"]["summary"],
                "status": i["fields"]["status"]["name"],
                "status_category": i["fields"]["status"]["statusCategory"]["key"],
                "assignee": (i["fields"].get("assignee") or {}).get("displayName", "Unassigned"),
                "type_icon": i["fields"]["issuetype"].get("iconUrl"),
                "type_name": i["fields"]["issuetype"].get("name"),
                "local": local
            })

        # Fetch status for extra bugs that are not in the main issue list
        if extra_bug_keys:
            try:
                extra_jql = f"key IN ({','.join(extra_bug_keys)})"
                extra_res = get_jira_cached(
                    f"{JIRA_DOMAIN}/rest/api/3/search/jql",
                    headers=headers_dict,
                    params={
                        "jql": extra_jql,
                        "fields": "status"
                    },
                    force_refresh=force_refresh
                )
                
                for i in extra_res.get("issues", []):
                    tracking_issues.append({
                        "key": i["key"],
                        "status": i["fields"]["status"]["name"],
                        "status_category": i["fields"]["status"]["statusCategory"]["key"],
                        "is_extra": True # Mark as extra info for frontend map
                    })
            except Exception as e:
                print(f"Error fetching extra bug statuses: {e}")

        return jsonify({
            "total_solved": total_solved,
            "bugs_solved": bugs_solved,
            "active_load": len(active_issues),
            "performance": list(performance.values()),
            "today_work": today_work,
            "tracking_issues": tracking_issues,
            "charts": {
                "status": status_dist,
                "type": type_dist,
                "priority": pri_dist
            }
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/sprint_tickets/update_field", methods=["POST"])
@login_required
def update_sprint_ticket_field():
    if not current_user.has_permission("edit_tickets") and current_user.role_name != 'Admin':
        return jsonify({"error": "Permission denied"}), 403

    data = request.json
    sprint_id = data.get("sprint_id")
    issue_key = data.get("issue_key")
    field = data.get("field")
    value = data.get("value")
    
    if not sprint_id or not issue_key or not field:
        return jsonify({"error": "Missing required fields"}), 400
        
    # Security: whitelist fields
    allowed_fields = ['pr_raised', 'pr_merged', 'deploy_status', 'qa_assignee', 'qa_status', 'bugs_found', 'requirements_clear', 'completed', 'comment', 'demo_done', 'is_flagged']
    if field not in allowed_fields:
        return jsonify({"error": "Invalid field"}), 400
        
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    # Check if record exists
    cursor.execute(f"SELECT id, {field} FROM sprint_tickets WHERE sprint_id = %s AND issue_key = %s", (sprint_id, issue_key))
    row = cursor.fetchone()
    
    old_value = None
    if row:
        old_value = row[1]
        cursor.execute(f"UPDATE sprint_tickets SET {field} = %s WHERE id = %s", (value, row[0]))
    else:
        cursor.execute(f"INSERT INTO sprint_tickets (sprint_id, issue_key, {field}) VALUES (%s, %s, %s)", (sprint_id, issue_key, value))
        
    conn.commit()
    conn.close()

    # Add Audit Log
    add_audit_log(page="Dashboard", item_key=issue_key, field_name=field, new_value=value, old_value=old_value)

    return jsonify({"success": True})

@app.route("/api/dashboard/recently_added", methods=["POST"])
@login_required
def recently_added_tickets():
    data = request.json
    team_ids = data.get("team_id")
    sprint_id = data.get("sprint_id")
    filter_mode = data.get("filter_mode", "team")
    
    if not team_ids:
        return jsonify({"error": "Team is required"}), 400
        
    if not isinstance(team_ids, list):
        team_ids = [team_ids]

    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    placeholders = ', '.join(['%s'] * len(team_ids))
    cursor.execute(f"SELECT account_id FROM team_members WHERE team_id IN ({placeholders})", tuple(team_ids))
    member_ids = [r[0] for r in cursor.fetchall() if r and r[0]]
    conn.close()

    headers_dict = dict(HEADERS)
    project_key_str = str(PROJECT_KEY)
    parent_team_filter = 'AND "team[team]" = "4da67a24-33ef-42a2-b940-840dd6e450bc"'
    # JQL Construction
    # We rely on 'Sprint CHANGED AFTER -3d' as requested.
    # Note: 'Sprint' field name is standard.
    # We add 'Sprint IS NOT EMPTY' to ensure we only target tickets that actually have a sprint context.
    jql_primary = f'project = "{project_key_str}" {parent_team_filter} AND Sprint IS NOT EMPTY AND Sprint CHANGED AFTER "-3d" ORDER BY updated DESC'
    
    # Fallback to updated time ONLY if the primary Sprint query fails (e.g. Jira instance config issue)
    jql_fallback = f'project = "{project_key_str}" {parent_team_filter} AND updated >= "-3d" ORDER BY updated DESC'
    
    def fetch_jira(jql_query):
        return requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params={
                "jql": jql_query,
                "maxResults": 50,
                "fields": "summary,status,assignee,created,updated,issuetype,customfield_10020",
                "expand": "changelog"
            }
        )

    def parse_jira_datetime(raw):
        if not raw:
            return None
        try:
            return datetime.strptime(raw, "%Y-%m-%dT%H:%M:%S.%f%z")
        except Exception:
            try:
                return datetime.strptime(raw, "%Y-%m-%dT%H:%M:%S%z")
            except Exception:
                return None

    try:
        print(f"DEBUG: Trying Primary JQL: {jql_primary}")
        res = fetch_jira(jql_primary)
        if res.status_code == 400 and jql_primary != jql_fallback:
            print(f"DEBUG: Primary JQL failed ({res.text}). Trying Fallback: {jql_fallback}")
            res = fetch_jira(jql_fallback)
            
        jira_payload = res.json()
        if "errorMessages" in jira_payload:
            return jsonify({"error": jira_payload["errorMessages"]}), 400
            
        issues = jira_payload.get("issues", [])
        cutoff = datetime.now(timezone.utc) - timedelta(days=3)
        formatted = []

        print(f"DEBUG: Processing {len(issues)} issues for recent activity...")
        
        for i in issues:
            f = i["fields"]
            sprints = f.get("customfield_10020") or []
            sprint_ids = []

            if isinstance(sprints, list):
                for s in sprints:
                    if isinstance(s, dict) and s.get("id") is not None:
                        try:
                            sprint_ids.append(int(s.get("id")))
                        except Exception:
                            pass
                    elif isinstance(s, str):
                        m = re.search(r"id=(\d+)", s)
                        if m:
                            sprint_ids.append(int(m.group(1)))

            # Check if ticket was CREATED recently (e.g. within cutoff)
            created_at = parse_jira_datetime(f.get("created"))
            is_newly_created = created_at and created_at >= cutoff
            
            sprint_changed_recently = False
            
            if is_newly_created:
                sprint_changed_recently = True
            else:
                # Check changelog for Sprint changes
                changelog = i.get("changelog") or {}
                histories = changelog.get("histories") or []
                for h in histories:
                    changed_at = parse_jira_datetime(h.get("created"))
                    if not changed_at or changed_at < cutoff:
                        continue
                    for item in h.get("items") or []:
                        field_name = (item.get("field") or "").lower()
                        field_id = (item.get("fieldId") or "").lower()
                        # Some instances use 'Sprint' with capital S
                        if field_name == "sprint" or field_id == "customfield_10020":
                            sprint_changed_recently = True
                            break
                    if sprint_changed_recently:
                        break

            if not sprint_changed_recently:
                continue

            assignee_obj = f.get("assignee") or {}
            formatted.append({
                "key": i["key"],
                "summary": f.get("summary"),
                "status": (f.get("status") or {}).get("name"),
                "assignee": assignee_obj.get("displayName"),
                "assignee_id": assignee_obj.get("accountId"),
                "created": f.get("updated"), # Return Updated time as 'created' for sorting
                "type_icon": (f.get("issuetype") or {}).get("iconUrl"),
                "sprint_ids": sprint_ids,
                "is_new": is_newly_created
            })

        print(f"DEBUG: Found {len(formatted)} matching recent tickets.")

        return jsonify({
            "tickets": formatted,
            "team_member_ids": member_ids,
            "sprint_id": sprint_id,
            "mode": filter_mode
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# =========================
# VELOCITY DATA (TRENDS)
# =========================
@app.route("/api/velocity_data", methods=["POST"])
def velocity_data():
    data = request.json
    range_type = data.get("range", "day")
    
    project_key_str = str(PROJECT_KEY)
    headers_dict = dict(HEADERS)
    
    jql = f'project = "{project_key_str}" AND statusCategory = "Done"'
    
    if range_type == "day":
        jql += " AND resolutiondate >= startOfDay()"
    elif range_type == "yesterday":
        jql += " AND resolutiondate >= startOfDay(-1d) AND resolutiondate < startOfDay()"
    elif range_type == "week":
        jql += " AND resolutiondate >= startOfWeek()"
    elif range_type == "month":
        jql += " AND resolutiondate >= startOfMonth()"
    elif range_type == "year":
        jql += " AND resolutiondate >= startOfYear()"

    # Fetch simple fields
    fields = "resolutiondate"
    
    # We need ALL issues to count properly
    all_issues = []
    start_at = 0
    max_results = 100
    
    while True:
        res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params={
                "jql": jql,
                "maxResults": max_results,
                "startAt": start_at,
                "fields": fields
            }
        ).json()
        
        issues = res.get("issues", [])
        all_issues.extend(issues)
        
        if start_at + len(issues) >= res.get("total", 0):
            break
        start_at += len(issues)

    # Aggregate by Date
    # For Year -> by Month
    # For others -> by Day
    from datetime import datetime
    
    timeline = {}
    
    for i in all_issues:
        date_str = i["fields"]["resolutiondate"]
        # Parse ISO date 2023-10-25T12:00:00.000+0000
        dt = datetime.strptime(date_str, "%Y-%m-%dT%H:%M:%S.%f%z")
        
        if range_type == "year":
            key = dt.strftime("%Y-%m") # Group by month
        else:
            key = dt.strftime("%Y-%m-%d") # Group by day

        timeline[key] = timeline.get(key, 0) + 1
        
    # Sort by date
    sorted_timeline = [{"date": k, "count": v} for k, v in sorted(timeline.items())]
    
    return jsonify(sorted_timeline)

# =========================
# USER DETAILS ROUTE
# =========================
@app.route("/user/<account_id>")
@login_required
@page_permission_required("dashboard") # Assuming user details is part of dashboard access
def user_details(account_id):
    return render_template("user_tickets.html", project=PROJECT_KEY, account_id=account_id)

# =========================
# FETCH USER TICKETS
# =========================
@app.route("/api/user_tickets", methods=["POST"])
def user_tickets():
    data = request.json
    account_id = data.get("accountId")
    range_type = data.get("range", "month")
    specific_date = data.get("specificDate")
    
    project_key_str = str(PROJECT_KEY)
    headers_dict = dict(HEADERS)
    
    if not account_id:
        return jsonify({"error": "Missing accountId"}), 400

    # Strict JQL construction
    jql = f'project = "{project_key_str}" AND statusCategory = "Done" AND assignee = "{account_id}"'
    
    if range_type == "date" and specific_date:
        from datetime import datetime, timedelta
        date_dt = datetime.strptime(specific_date, "%Y-%m-%d")
        next_day = (date_dt + timedelta(days=1)).strftime("%Y-%m-%d")
        jql += f' AND resolutiondate >= "{specific_date}" AND resolutiondate < "{next_day}"'
    elif range_type == "day":
        jql += " AND resolutiondate >= startOfDay()"
    elif range_type == "yesterday":
        jql += " AND resolutiondate >= startOfDay(-1d) AND resolutiondate < startOfDay()"
    elif range_type == "week":
        jql += " AND resolutiondate >= startOfWeek()"
    elif range_type == "month":
        jql += " AND resolutiondate >= startOfMonth()"
    elif range_type == "year":
        jql += " AND resolutiondate >= startOfYear()"
    
    # We want details
    fields = "summary,status,resolutiondate,priority,issuetype,assignee"
    
    # Fetch issues (using GET /jql as established)
    res = requests.get(
        f"{JIRA_DOMAIN}/rest/api/3/search/jql",
        headers=headers_dict,
        params={
            "jql": jql,
            "maxResults": 100, # Limit to last 100 for now
            "fields": fields
        }
    ).json()
    
    return jsonify(res.get("issues", []))

# =========================
# ASSIGNEEE FULL STATS (NEW)
# =========================
@app.route("/api/suggest_tickets", methods=["POST"])
def suggest_tickets():
    query = request.json.get("query", "")
    if not query:
        return jsonify({"results": []})
    
    # Explicitly cast LocalProxy objects
    headers_dict = dict(HEADERS)
    project_key_str = str(PROJECT_KEY)

    # Search for tickets matching key or summary
    jql = f'project = "{project_key_str}" AND (key ~ "{query}*" OR summary ~ "{query}*")'
    
    try:
        res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params={
                "jql": jql,
                "maxResults": 10,
                "fields": "summary,key"
            }
        ).json()
        
        issues = res.get("issues", [])
        results = [
            {"id": i["key"], "text": f'{i["key"]}: {i["fields"]["summary"]}'}
            for i in issues
        ]
        return jsonify({"results": results})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/assignee_full_profile", methods=["POST"])
def assignee_full_profile():
    data = request.json
    account_id = data.get("accountId")
    range_type = data.get("range", "month")
    specific_date = data.get("specificDate") # YYYY-MM-DD
    
    if not account_id:
        return jsonify({"error": "Missing accountId"}), 400

    # Explicitly cast LocalProxy objects to ensure they are passed correctly
    headers_dict = dict(HEADERS)
    project_key_str = str(PROJECT_KEY)

    # 1. FETCH DONE ISSUES (In Range)
    jql_done = f'project = "{project_key_str}" AND statusCategory = "Done" AND assignee = "{account_id}"'
    
    if range_type == "date" and specific_date:
        # Fetch for that specific day
        from datetime import datetime, timedelta
        date_dt = datetime.strptime(specific_date, "%Y-%m-%d")
        next_day = (date_dt + timedelta(days=1)).strftime("%Y-%m-%d")
        jql_done += f' AND resolutiondate >= "{specific_date}" AND resolutiondate < "{next_day}"'
    elif range_type == "day":
        jql_done += " AND resolutiondate >= startOfDay()"
    elif range_type == "yesterday":
        jql_done += " AND resolutiondate >= startOfDay(-1d) AND resolutiondate < startOfDay()"
    elif range_type == "week":
        jql_done += " AND resolutiondate >= startOfWeek()"
    elif range_type == "month":
        jql_done += " AND resolutiondate >= startOfMonth()"
    elif range_type == "year":
        jql_done += " AND resolutiondate >= startOfYear()"
        
    res_done = requests.get(
        f"{JIRA_DOMAIN}/rest/api/3/search/jql",
        headers=headers_dict,
        params={
            "jql": jql_done,
            "maxResults": 100,
            "fields": "priority,issuetype"
        }
    ).json()
    done_issues = res_done.get("issues", [])

    # 2. FETCH OPEN ISSUES (Current Load - No date range)
    jql_open = f'project = "{project_key_str}" AND statusCategory != "Done" AND assignee = "{account_id}"'
    
    res_open = requests.get(
        f"{JIRA_DOMAIN}/rest/api/3/search/jql",
        headers=headers_dict,
        params={
            "jql": jql_open,
            "maxResults": 100,
            "fields": "priority,issuetype,status"
        }
    ).json()
    open_issues = res_open.get("issues", [])

    return jsonify({
        "done": done_issues,
        "open": open_issues
    })


# =========================
# CUSTOM REPORTS
# =========================
@app.route("/reports")
@page_permission_required("custom_reports")
def reports_page():
    return render_template("custom_reports.html", project=PROJECT_KEY)

# =========================
# CUSTOM REPORTS API
# =========================

@app.route("/api/reports", methods=["GET", "POST"])
def manage_reports():
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "POST":
        data = request.json
        name = data.get("name")
        jql = data.get("jql")
        
        if not name or not jql:
            return jsonify({"error": "Name and JQL are required"}), 400
            
        cursor.execute("INSERT INTO custom_reports (name, jql) VALUES (%s, %s)", (name, jql))
        conn.commit()
        report_id = cursor.lastrowid
        conn.close()
        return jsonify({"id": report_id, "name": name, "jql": jql})
    
    else:  # GET
        cursor.execute("SELECT id, name, jql, created_at FROM custom_reports ORDER BY created_at DESC")
        reports = [{"id": r[0], "name": r[1], "jql": r[2], "created": r[3]} for r in cursor.fetchall()]
        conn.close()
        return jsonify(reports)

@app.route("/api/reports/<int:report_id>", methods=["GET", "PUT", "DELETE"])
def report_detail(report_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "GET":
        cursor.execute("SELECT id, name, jql, created_at FROM custom_reports WHERE id = %s", (report_id,))
        row = cursor.fetchone()
        conn.close()
        if row:
            return jsonify({"id": row[0], "name": row[1], "jql": row[2], "created": row[3]})
        return jsonify({"error": "Report not found"}), 404
        
    elif request.method == "PUT":
        data = request.json
        name = data.get("name")
        jql = data.get("jql")
        
        if not name or not jql:
            return jsonify({"error": "Name and JQL are required"}), 400
            
        cursor.execute("UPDATE custom_reports SET name = %s, jql = %s WHERE id = %s", (name, jql, report_id))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
        
    elif request.method == "DELETE":
        cursor.execute("DELETE FROM custom_reports WHERE id = %s", (report_id,))
        conn.commit()
        conn.close()
        return jsonify({"success": True})

# =========================
# 📊 CUSTOM TRACKERS (V2)
# =========================
@app.route("/trackers_v2")
@login_required
def trackers_v2_page():
    _, _, _, domain = _get_jira_config()
    return render_template("trackers_v2.html", jira_domain=domain)

@app.route("/api/trackers_v2", methods=["GET", "POST"])
@login_required
def manage_trackers_v2():
    if request.method == "POST" and current_user.role_name == 'Viewer':
        return jsonify({"error": "Permission denied"}), 403
        
    conn, cursor = get_db_connection(dictionary=True)
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "POST":
        data = request.json
        name = data.get("name")
        jql = data.get("jql")
        columns = data.get("columns", [])
        is_public = data.get("is_public", False)
        
        if not name or not jql:
            return jsonify({"error": "Name and JQL are required"}), 400
            
        try:
            # Create Tracker
            cursor.execute("""
                INSERT INTO trackers_v2 (name, jql, created_by, is_public)
                VALUES (%s, %s, %s, %s)
            """, (name, jql, current_user.id, is_public))
            tracker_id = cursor.lastrowid
            
            # Create Columns
            for idx, col in enumerate(columns):
                cursor.execute("""
                    INSERT INTO tracker_columns (tracker_id, name, column_type, options, order_index)
                    VALUES (%s, %s, %s, %s, %s)
                """, (tracker_id, col["name"], col["type"], json.dumps(col.get("options", [])), idx))
                
            conn.commit()
            return jsonify({"id": tracker_id, "success": True})
        except Exception as e:
            conn.rollback()
            return jsonify({"error": str(e)}), 500
        finally:
            conn.close()
    
    else:  # GET - List all trackers
        cursor.execute("""
            SELECT t.*, u.name as creator_name 
            FROM trackers_v2 t
            LEFT JOIN users u ON t.created_by = u.id
            WHERE t.created_by = %s OR t.is_public = TRUE
            ORDER BY t.created_at DESC
        """, (current_user.id,))
        trackers = cursor.fetchall()
        conn.close()
        return jsonify(trackers)

@app.route("/api/trackers_v2/<int:tracker_id>", methods=["GET", "PUT", "DELETE"])
@login_required
def tracker_v2_detail(tracker_id):
    if request.method in ["PUT", "DELETE"] and current_user.role_name == 'Viewer':
        return jsonify({"error": "Permission denied"}), 403
        
    conn, cursor = get_db_connection(dictionary=True)
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "DELETE":
        # Check ownership
        cursor.execute("SELECT created_by FROM trackers_v2 WHERE id = %s", (tracker_id,))
        tracker = cursor.fetchone()
        if not tracker:
            conn.close()
            return jsonify({"error": "Tracker not found"}), 404
        if tracker["created_by"] != current_user.id and current_user.role_name != 'Admin':
            conn.close()
            return jsonify({"error": "Permission denied"}), 403
            
        cursor.execute("DELETE FROM trackers_v2 WHERE id = %s", (tracker_id,))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
        
    elif request.method == "PUT":
        data = request.json
        try:
            # 1. Update Tracker Metadata
            cursor.execute("""
                UPDATE trackers_v2 
                SET name = %s, jql = %s, is_public = %s
                WHERE id = %s AND (created_by = %s OR %s = 'Admin')
            """, (data["name"], data["jql"], data["is_public"], tracker_id, current_user.id, current_user.role_name))
            
            # 2. Update Columns
            new_columns = data.get("columns", [])
            existing_col_ids = []
            
            for idx, col in enumerate(new_columns):
                col_name = col.get("name")
                col_type = col.get("type")
                col_options = json.dumps(col.get("options", []))
                col_id = col.get("id")
                
                if col_id:
                    # Update existing column
                    cursor.execute("""
                        UPDATE tracker_columns 
                        SET name = %s, column_type = %s, options = %s, order_index = %s
                        WHERE id = %s AND tracker_id = %s
                    """, (col_name, col_type, col_options, idx, col_id, tracker_id))
                    existing_col_ids.append(col_id)
                else:
                    # Insert new column
                    cursor.execute("""
                        INSERT INTO tracker_columns (tracker_id, name, column_type, options, order_index)
                        VALUES (%s, %s, %s, %s, %s)
                    """, (tracker_id, col_name, col_type, col_options, idx))
                    existing_col_ids.append(cursor.lastrowid)
            
            # 3. Delete removed columns
            if existing_col_ids:
                format_strings = ','.join(['%s'] * len(existing_col_ids))
                cursor.execute(f"DELETE FROM tracker_columns WHERE tracker_id = %s AND id NOT IN ({format_strings})", [tracker_id] + existing_col_ids)
            else:
                cursor.execute("DELETE FROM tracker_columns WHERE tracker_id = %s", (tracker_id,))
                
            conn.commit()
            return jsonify({"success": True})
        except Exception as e:
            conn.rollback()
            return jsonify({"error": str(e)}), 500
        finally:
            conn.close()

    else: # GET
        cursor.execute("SELECT * FROM trackers_v2 WHERE id = %s", (tracker_id,))
        tracker = cursor.fetchone()
        if not tracker:
            conn.close()
            return jsonify({"error": "Tracker not found"}), 404
            
        cursor.execute("SELECT * FROM tracker_columns WHERE tracker_id = %s ORDER BY order_index ASC", (tracker_id,))
        columns = cursor.fetchall()
        for col in columns:
            col["options"] = json.loads(col["options"]) if col["options"] else []
            
        conn.close()
        return jsonify({**tracker, "columns": columns})

@app.route("/api/trackers_v2/<int:tracker_id>/data", methods=["GET", "POST"])
@login_required
def tracker_v2_data(tracker_id):
    if request.method == "POST" and current_user.role_name == 'Viewer':
        return jsonify({"error": "Permission denied"}), 403
        
    if request.method == "POST":
        data = request.json
        issue_key = data.get("issue_key")
        column_id = data.get("column_id")
        value = data.get("value")
        
        if not all([issue_key, column_id]):
            return jsonify({"error": "Missing key data"}), 400
            
        conn, cursor = get_db_connection()
        if not conn: return jsonify({"error": "Database error"}), 500
        try:
            # Get old value and column name for audit log
            cursor.execute("SELECT name FROM tracker_columns WHERE id = %s", (column_id,))
            col_row = cursor.fetchone()
            col_name = col_row[0] if col_row else f"Column {column_id}"

            cursor.execute("SELECT value FROM tracker_data_v2 WHERE tracker_id = %s AND issue_key = %s AND column_id = %s", (tracker_id, issue_key, column_id))
            old_val_row = cursor.fetchone()
            old_value = old_val_row[0] if old_val_row else None

            cursor.execute("""
                INSERT INTO tracker_data_v2 (tracker_id, issue_key, column_id, value)
                VALUES (%s, %s, %s, %s)
                ON DUPLICATE KEY UPDATE value = VALUES(value)
            """, (tracker_id, issue_key, column_id, value))
            conn.commit()

            # Add Audit Log
            add_audit_log(page=f"Tracker #{tracker_id}", item_key=issue_key, field_name=col_name, new_value=value, old_value=old_value)

            return jsonify({"success": True})
        except Exception as e:
            conn.rollback()
            return jsonify({"error": str(e)}), 500
        finally:
            conn.close()
            
    else: # GET - Fetch Jira Data + Stored Data
        conn, cursor = get_db_connection(dictionary=True)
        if not conn: return jsonify({"error": "Database error"}), 500
        
        try:
            cursor.execute("SELECT jql FROM trackers_v2 WHERE id = %s", (tracker_id,))
            tracker = cursor.fetchone()
            if not tracker:
                return jsonify({"error": "Tracker not found"}), 404
                
            # Fetch Local Data
            cursor.execute("SELECT issue_key, column_id, value FROM tracker_data_v2 WHERE tracker_id = %s", (tracker_id,))
            stored_rows = cursor.fetchall()
            stored_data = {}
            for row in stored_rows:
                if row["issue_key"] not in stored_data:
                    stored_data[row["issue_key"]] = {}
                stored_data[row["issue_key"]][row["column_id"]] = row["value"]

            # Fetch RCA status
            cursor.execute("SELECT issue_key, submitted_at FROM tracker_rca WHERE tracker_id = %s", (tracker_id,))
            rca_rows = cursor.fetchall()
            rca_status = {row["issue_key"]: (row["submitted_at"] is not None) for row in rca_rows}
                
        except Exception as e:
            return jsonify({"error": f"Database Error: {str(e)}"}), 500
        finally:
            conn.close()

        # Fetch Jira Data using JQL
        email, token, _, domain = _get_jira_config()
        if not email or not token:
            return jsonify({"error": "Jira credentials not configured"}), 400
            
        auth_str = f"{email}:{token}"
        auth_b64 = base64.b64encode(auth_str.encode()).decode()
        headers = {
            "Authorization": f"Basic {auth_b64}",
            "Content-Type": "application/json"
        }
        
        jira_issues = []
        try:
            params = {
                "jql": tracker["jql"],
                "maxResults": 100,
                "fields": "summary,status,assignee,priority,issuetype"
            }
            # Use /search/jql as required by latest Jira Cloud API
            res = requests.get(f"{domain.rstrip('/')}/rest/api/3/search/jql", headers=headers, params=params)
            
            if res.status_code == 200:
                issues_data = res.json().get("issues", [])
                for issue in issues_data:
                    key = issue["key"]
                    fields = issue.get("fields", {})
                    jira_issues.append({
                        "key": key,
                        "summary": fields.get("summary", "No Summary"),
                        "status": fields.get("status", {}).get("name", "Unknown"),
                        "status_category": fields.get("status", {}).get("statusCategory", {}).get("key", "new"),
                        "priority": fields.get("priority", {}).get("name", "Medium"),
                        "type": fields.get("issuetype", {}).get("name", "Task"),
                        "type_icon": fields.get("issuetype", {}).get("iconUrl"),
                        "assignee": fields.get("assignee", {}).get("displayName", "Unassigned") if fields.get("assignee") else "Unassigned",
                        "custom_data": stored_data.get(key, {}),
                        "rca_filled": rca_status.get(key, False)
                    })
            else:
                # Jira error (e.g. 400 Bad JQL)
                error_msg = res.json().get("errorMessages", [res.text])[0]
                return jsonify({"error": f"Jira API Error: {error_msg}"}), res.status_code
                
        except Exception as e:
            return jsonify({"error": f"Jira Connection Error: {str(e)}"}), 500
            
        return jsonify({
            "issues": jira_issues
        })

@app.route("/api/trackers_v2/<int:tracker_id>/rca/<issue_key>", methods=["GET", "POST"])
@login_required
def tracker_v2_rca(tracker_id, issue_key):
    conn, cursor = get_db_connection(dictionary=True)
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "POST":
        if current_user.role_name == 'Viewer':
            conn.close()
            return jsonify({"error": "Permission denied"}), 403
            
        data = request.json
        rca_text = data.get("rca_text")
        fix_text = data.get("fix_text")
        prevention_text = data.get("prevention_text")
        issue_details = data.get("issue_details") # JSON
        
        try:
            cursor.execute("""
                INSERT INTO tracker_rca (tracker_id, issue_key, issue_details, rca_text, fix_text, prevention_text, submitted_at)
                VALUES (%s, %s, %s, %s, %s, %s, NOW())
                ON DUPLICATE KEY UPDATE 
                    rca_text = VALUES(rca_text), 
                    fix_text = VALUES(fix_text), 
                    prevention_text = VALUES(prevention_text),
                    submitted_at = NOW()
            """, (tracker_id, issue_key, json.dumps(issue_details) if issue_details else None, rca_text, fix_text, prevention_text))
            conn.commit()
            return jsonify({"success": True})
        except Exception as e:
            return jsonify({"error": str(e)}), 500
        finally:
            conn.close()
    
    else: # GET
        cursor.execute("SELECT * FROM tracker_rca WHERE tracker_id = %s AND issue_key = %s", (tracker_id, issue_key))
        row = cursor.fetchone()
        conn.close()
        return jsonify(row if row else {})

@app.route("/api/trackers_v2/<int:tracker_id>/rca/link", methods=["POST"])
@login_required
def tracker_v2_rca_link(tracker_id, issue_key=None):
    if current_user.role_name == 'Viewer':
        return jsonify({"error": "Permission denied"}), 403
        
    data = request.json
    issue_key = data.get("issue_key")
    issue_details = data.get("issue_details") # Dict
    
    if not issue_key:
        return jsonify({"error": "Issue key is required"}), 400
        
    token = secrets.token_urlsafe(32)
    
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    try:
        cursor.execute("""
            INSERT INTO tracker_rca (tracker_id, issue_key, issue_details, token)
            VALUES (%s, %s, %s, %s)
            ON DUPLICATE KEY UPDATE token = VALUES(token), issue_details = VALUES(issue_details)
        """, (tracker_id, issue_key, json.dumps(issue_details), token))
        conn.commit()
        
        base_url = request.url_root.rstrip('/')
        link = f"{base_url}/rca/form/{token}"
        return jsonify({"link": link})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()

# Public RCA Form Route
@app.route("/rca/form/<token>")
def rca_public_form(token):
    conn, cursor = get_db_connection(dictionary=True)
    if not conn: return "Database Error", 500
    
    cursor.execute("SELECT * FROM tracker_rca WHERE token = %s", (token,))
    row = cursor.fetchone()
    conn.close()
    
    if not row:
        return "Invalid or expired link", 404
        
    issue = json.loads(row["issue_details"]) if row["issue_details"] else {"key": row["issue_key"]}
    return render_template("rca_form.html", token=token, issue=issue, rca=row, jira_domain=JIRA_DOMAIN)

# Public API for form submission
@app.route("/api/rca/submit/<token>", methods=["POST"])
def rca_public_submit(token):
    data = request.json
    rca_text = data.get("rca_text")
    fix_text = data.get("fix_text")
    prevention_text = data.get("prevention_text")
    
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    try:
        cursor.execute("""
            UPDATE tracker_rca 
            SET rca_text = %s, fix_text = %s, prevention_text = %s, submitted_at = NOW()
            WHERE token = %s
        """, (rca_text, fix_text, prevention_text, token))
        conn.commit()
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()

@app.route("/tracker")
@page_permission_required("tracker")
def tracker():
    return render_template("tracker.html", project=PROJECT_KEY)

@app.route("/report_view")
@login_required
@page_permission_required("custom_reports")
def report_view():
    return render_template("report_view.html", project=PROJECT_KEY)

@app.route("/api/execute_jql", methods=["POST"])
def execute_jql():
    data = request.json
    jql = data.get("jql")
    
    if not jql:
        return jsonify({"error": "No JQL provided"}), 400

    print(f"DEBUG: Executing Custom JQL: {jql}")
    
    # Explicitly cast LocalProxy objects
    headers_dict = dict(HEADERS)

    fields = "summary,status,assignee,priority,issuetype,created,updated,resolutiondate,project,reporter,resolution,duedate,customfield_10016,customfield_10077"
    
    all_issues = []
    start_at = 0
    max_results = 100
    
    while True:
        try:
            jira_response = requests.get(
                f"{str(JIRA_DOMAIN)}/rest/api/3/search/jql",
                headers=headers_dict,
                params={
                    "jql": jql,
                    "maxResults": max_results,
                    "startAt": start_at,
                    "fields": fields
                }
            )
            
            if jira_response.status_code != 200:
                print(f"DEBUG: Jira Error {jira_response.status_code}: {jira_response.text}")
                return jsonify({"error": f"Jira API Error {jira_response.status_code}", "details": jira_response.text}), jira_response.status_code

            res = jira_response.json()
        except Exception as e:
            print(f"DEBUG: Exception during Jira Request: {e}")
            return jsonify({"error": "Failed to connect to Jira or parse response", "details": str(e)}), 500
        
        if "errorMessages" in res:
             return jsonify({"error": res["errorMessages"]}), 400

        issues = res.get("issues", [])
        all_issues.extend(issues)
        
        if start_at + len(issues) >= res.get("total", 0):
            break
        start_at += len(issues)
    
    print(f"DEBUG: Found {len(all_issues)} issues for tracker/reports")
    return jsonify(all_issues)

# =========================
# MULTI-TRACKER API
# =========================

@app.route("/api/trackers", methods=["GET", "POST"])
def manage_trackers():
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "POST":
        title = request.json.get("title")
        if not title:
            return jsonify({"error": "Title is required"}), 400
        cursor.execute("INSERT INTO trackers (title) VALUES (%s)", (title,))
        conn.commit()
        tracker_id = cursor.lastrowid
        conn.close()
        return jsonify({"id": tracker_id, "title": title})
    
    else:
        cursor.execute("SELECT id, title, created_at FROM trackers ORDER BY created_at DESC")
        trackers = [{"id": r[0], "title": r[1], "created_at": r[2]} for r in cursor.fetchall()]
        conn.close()
        return jsonify(trackers)

@app.route("/api/trackers/<int:tracker_id>", methods=["DELETE"])
@permission_required("manage_settings")
def delete_tracker(tracker_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    cursor.execute("DELETE FROM trackers WHERE id = %s", (tracker_id,))
    cursor.execute("DELETE FROM tracker_tickets WHERE tracker_id = %s", (tracker_id,))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route("/api/trackers/<int:tracker_id>/tickets", methods=["GET", "POST"])
def tracker_tickets(tracker_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "POST":
        issue_key = request.json.get("issueKey")
        if not issue_key:
            return jsonify({"error": "Issue key is required"}), 400
        cursor.execute("INSERT INTO tracker_tickets (tracker_id, issue_key) VALUES (%s, %s)", (tracker_id, issue_key))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
    
    else:
        cursor.execute("SELECT issue_key, comment FROM tracker_tickets WHERE tracker_id = %s", (tracker_id,))
        tickets = [{"issue_key": r[0], "comment": r[1]} for r in cursor.fetchall()]
        conn.close()
        return jsonify(tickets)

@app.route("/api/trackers/<int:tracker_id>/tickets/<string:issue_key>", methods=["DELETE"])
def delete_tracker_ticket(tracker_id, issue_key):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    cursor.execute("DELETE FROM tracker_tickets WHERE tracker_id = %s AND issue_key = %s", (tracker_id, issue_key))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route("/api/trackers/<int:tracker_id>/tickets/comment", methods=["PUT"])
def update_ticket_comment(tracker_id):
    data = request.json
    issue_key = data.get("issueKey")
    comment = data.get("comment")
    
    if not issue_key:
        return jsonify({"error": "Issue key is required"}), 400
        
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    cursor.execute("UPDATE tracker_tickets SET comment = %s WHERE tracker_id = %s AND issue_key = %s", (comment, tracker_id, issue_key))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

# =========================
# TODO API
# =========================

@app.route("/todo")
@page_permission_required("todo")
def todo_page():
    return render_template("todo.html", project=PROJECT_KEY)

@app.route("/api/todos", methods=["GET", "POST"])
@login_required
def manage_todos():
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    if request.method == "POST":
        data = request.json
        title = data.get("title")
        description = data.get("description", "")
        priority = data.get("priority", "Low")
        due_date = data.get("due_date") # YYYY-MM-DD
        tags = data.get("tags", "[]")
        if isinstance(tags, list):
            tags = json.dumps(tags)
        
        if not title or not due_date:
            return jsonify({"error": "Title and due date are required"}), 400
            
        cursor.execute('''
            INSERT INTO todos (user_id, title, description, priority, due_date, tags)
            VALUES (%s, %s, %s, %s, %s, %s)
        ''', (current_user.id, title, description, priority, due_date, tags))
        conn.commit()
        todo_id = cursor.lastrowid
        conn.close()
        return jsonify({"id": todo_id, "title": title, "status": "Pending"})
    
    else:
        date_filter = request.args.get("date")
        if date_filter:
            cursor.execute("SELECT id, title, description, priority, due_date, status, tags FROM todos WHERE user_id = %s AND due_date = %s ORDER BY created_at DESC", (current_user.id, date_filter,))
        else:
            cursor.execute("SELECT id, title, description, priority, due_date, status, tags FROM todos WHERE user_id = %s ORDER BY due_date ASC, created_at DESC", (current_user.id,))
            
        todos = []
        for r in cursor.fetchall():
            todos.append({
                "id": r[0],
                "title": r[1],
                "description": r[2],
                "priority": r[3],
                "due_date": r[4],
                "status": r[5],
                "tags": r[6] or "[]"
            })
        conn.close()
        return jsonify(todos)

@app.route("/api/todos/<int:todo_id>", methods=["PUT", "DELETE"])
@login_required
def update_delete_todo(todo_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    # Check if the todo exists and belongs to the user (unless they are Admin)
    cursor.execute("SELECT user_id FROM todos WHERE id = %s", (todo_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        return jsonify({"error": "Todo not found"}), 404
    
    todo_user_id = row[0]
    is_admin = current_user.role_name == 'Admin'
    
    if todo_user_id != current_user.id and not is_admin:
        conn.close()
        return jsonify({"error": "Permission denied"}), 403

    if request.method == "DELETE":
        if not current_user.has_permission("delete_tickets") and not is_admin and todo_user_id != current_user.id:
            conn.close()
            return jsonify({"error": "Permission denied"}), 403
            
        cursor.execute("DELETE FROM todos WHERE id = %s", (todo_id,))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
    
    else:
        data = request.json
        status = data.get("status")
        title = data.get("title")
        description = data.get("description")
        priority = data.get("priority")
        due_date = data.get("due_date")
        tags = data.get("tags")
        if isinstance(tags, list):
            tags = json.dumps(tags)
        
        update_fields = []
        params = []
        if status:
            update_fields.append("status = %s")
            params.append(status)
        if title:
            update_fields.append("title = %s")
            params.append(title)
        if description is not None:
            update_fields.append("description = %s")
            params.append(description)
        if priority:
            update_fields.append("priority = %s")
            params.append(priority)
        if due_date:
            update_fields.append("due_date = %s")
            params.append(due_date)
        if tags is not None:
            update_fields.append("tags = %s")
            params.append(tags)
            
        if update_fields:
            params.append(todo_id)
            cursor.execute(f"UPDATE todos SET {', '.join(update_fields)} WHERE id = %s", params)
            conn.commit()
            
        conn.close()
        return jsonify({"success": True})

# =========================
# EXPLORER API
# =========================

@app.route("/explorer")
@page_permission_required("explorer")
def explorer_page():
    return render_template("explorer.html", project=PROJECT_KEY)

@app.route("/api/explorer", methods=["POST"])
def explorer_data():
    # Explicitly cast LocalProxy objects
    headers_dict = dict(HEADERS)
    project_key_str = str(PROJECT_KEY)

    data = request.json
    assignee = data.get("assignee")
    priority = data.get("priority")
    timeline = data.get("timeline", "all")
    start_date = data.get("startDate")
    end_date = data.get("endDate")
    is_production = data.get("production", False)
    customer = data.get("customer")
    
    # If raw JQL is provided, use it instead of building one
    raw_jql = data.get("jql")
    if raw_jql:
        jql = raw_jql
    else:
        jql = f'project = "{project_key_str}"'
        
        if assignee:
            jql += f' AND assignee = "{assignee}"'
        
        if customer:
            jql += f' AND "Customer" = "{customer}"'
        
        if priority and priority != "All":
            jql += f' AND priority = "{priority}"'
            
        # Timeline Logic
        if timeline == "today":
            jql += " AND created >= startOfDay()"
        elif timeline == "week":
            jql += " AND created >= startOfWeek()"
        elif timeline == "2weeks":
            jql += " AND created >= '-2w'"
        elif timeline == "month":
            jql += " AND created >= startOfMonth()"
        elif timeline == "year":
            jql += " AND created >= startOfYear()"
        elif timeline == "range":
            if start_date:
                jql += f" AND created >= '{start_date} 00:00'"
            if end_date:
                jql += f" AND created <= '{end_date} 23:59'"

        query_type = data.get("queryType", "normal")

        if query_type == "stale":
            # Logic: Created more than 1 week ago AND no update since 5 days
            jql += ' AND created <= "-1w" AND updated <= "-5d"'
        elif query_type == "critical":
            # Logic: Priority is Highest AND status is NOT in any closed/resolved category
            jql += ' AND priority = "Highest" AND status NOT IN ("Deployed", "Done", "Not A Bug", "Ready for Staging", "Resolved", "Staged", "Unable to Reproduce", "Ready for QA")'

        if is_production:
            jql += ' AND "platform[checkboxes]" = PRODUCTION'

        # Filter out common "Closed" or "Resolved" statuses for normal view if no specific status is requested
        if query_type == "normal":
            jql += ' AND status NOT IN ("Deployed", "Done", "Not A Bug", "Ready for Staging", "Resolved", "Staged", "Unable to Reproduce", "Ready for QA")'

        jql += " ORDER BY created DESC, assignee ASC, updated DESC"
    
    print(f"DEBUG Explorer JQL: {jql}")
    
    # Fetch issues
    start_at = 0
    max_results = 50
    
    params = {
        "jql": jql,
        "maxResults": max_results,
        "startAt": start_at,
        "fields": "summary,status,assignee,priority,created,updated,issuetype,customfield_10077"
    }
    
    try:
        res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params=params
        ).json()
        
        if "errorMessages" in res:
            return jsonify({"error": res["errorMessages"]}), 400
            
        issues = res.get("issues", [])
        if issues:
            print(f"DEBUG: First issue fields: {list(issues[0]['fields'].keys())}")
            print(f"DEBUG: Created: {issues[0]['fields'].get('created')}")
        formatted = []
        for i in issues:
            f = i["fields"]
            formatted.append({
                "key": i["key"],
                "summary": f.get("summary"),
                "status": f.get("status", {}).get("name"),
                "assignee": f.get("assignee", {}).get("displayName") if f.get("assignee") else "Unassigned",
                "priority": f.get("priority", {}).get("name"),
                "customer": ", ".join([c.get("value") for c in f.get("customfield_10077")]) if isinstance(f.get("customfield_10077"), list) else (f.get("customfield_10077", {}).get("value") if isinstance(f.get("customfield_10077"), dict) else f.get("customfield_10077")),
                "created": f.get("created"),
                "updated": f.get("updated"),
                "type": f.get("issuetype", {}).get("name")
            })
            
        return jsonify({
            "issues": formatted,
            "total": res.get("total", 0)
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/query_builder")
@page_permission_required("query_builder")
def query_builder_page():
    return render_template("query_builder.html", project=PROJECT_KEY)

@app.route("/api/query_builder", methods=["POST"])
def query_builder_data():
    # Explicitly cast LocalProxy objects
    headers_dict = dict(HEADERS)
    project_key_str = str(PROJECT_KEY)

    data = request.json
    assignees = data.get("assignees", [])
    priorities = data.get("priorities", [])
    statuses = data.get("statuses", [])
    start_date = data.get("startDate")
    end_date = data.get("endDate")
    platform_mode = data.get("platform", "all")
    
    jql = f'project = "{project_key_str}"'
    
    if assignees:
        assignee_list = ", ".join([f'"{a}"' for a in assignees])
        jql += f' AND assignee IN ({assignee_list})'
    
    if platform_mode == "production":
        jql += ' AND "platform[checkboxes]" = PRODUCTION'
    elif platform_mode == "non-production":
        jql += ' AND ("platform[checkboxes]" != PRODUCTION OR "platform[checkboxes]" is EMPTY)'
    
    if priorities:
        priority_list = ", ".join([f'"{p}"' for p in priorities])
        jql += f' AND priority IN ({priority_list})'
        
    if statuses:
        status_list = ", ".join([f'"{s}"' for s in statuses])
        jql += f' AND status IN ({status_list})'
        
    if start_date:
        jql += f" AND updated >= '{start_date} 00:00'"
    if end_date:
        jql += f" AND updated <= '{end_date} 23:59'"
        
    jql += " ORDER BY updated DESC"
    
    print(f"DEBUG Query Builder JQL: {jql}")
    
    # Fetch issues using existing explorer_data logic but with the new JQL
    # Since I can't easily call explorer_data from here without refactoring,
    # I'll implement a common fetcher or just duplicate for now as is typical in this script.
    
    fields = "summary,status,assignee,priority,created,updated,issuetype,customfield_10077"
    params = {
        "jql": jql,
        "maxResults": 100,
        "fields": fields
    }
    
    try:
        res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params=params
        ).json()
        
        if "errorMessages" in res:
            return jsonify({"error": res["errorMessages"]}), 400
            
        issues = res.get("issues", [])
        formatted = []
        for i in issues:
            f = i["fields"]
            formatted.append({
                "key": i["key"],
                "summary": f.get("summary"),
                "status": f.get("status", {}).get("name"),
                "assignee": f.get("assignee", {}).get("displayName") if f.get("assignee") else "Unassigned",
                "priority": f.get("priority", {}).get("name"),
                "customer": ", ".join([c.get("value") for c in f.get("customfield_10077")]) if isinstance(f.get("customfield_10077"), list) else (f.get("customfield_10077", {}).get("value") if isinstance(f.get("customfield_10077"), dict) else f.get("customfield_10077")),
                "created": f.get("created"),
                "updated": f.get("updated"),
                "type": f.get("issuetype", {}).get("name")
            })
            
        return jsonify({
            "issues": formatted,
            "total": res.get("total", 0)
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/bulk_update")
@page_permission_required("bulk_update")
def bulk_update_page():
    return render_template("bulk_update.html", project=PROJECT_KEY)

@app.route("/api/bulk_update", methods=["POST"])
def bulk_update():
    # Explicitly cast LocalProxy objects
    headers_dict = dict(HEADERS)

    data = request.json
    issue_keys = data.get("issueKeys", [])
    field = data.get("field")
    value = data.get("value", "")
    if isinstance(value, str):
        value = value.strip()

    if not issue_keys or not field:
        return jsonify({"error": "Issue keys and field are required"}), 400

    results = []
    
    # Prepare Jira Payload
    jira_field = field
    if field == "sprint":
        jira_field = "customfield_10020"
    elif field == "labels":
        jira_field = "labels"
    elif field == "priority":
        jira_field = "priority"
    elif field == "assignee":
        jira_field = "assignee"

    for key in issue_keys:
        key = key.strip()
        if not key: continue
        
        url = f"{JIRA_DOMAIN}/rest/api/3/issue/{key}"
        
        update_data = {}
        if field == "labels":
            # Jira labels are an array. We usually append or replace. 
            # For simplicity in "bulk update", we append if it's a single label or replace if desired.
            # User said "add a value to update it". 
            update_data = {"update": {"labels": [{"add": value}]}}
        elif field == "assignee":
            update_data = {"fields": {"assignee": {"accountId": value if value != "unassigned" else None}}}
        elif field == "priority":
            update_data = {"fields": {"priority": {"name": value}}}
        elif field == "sprint":
            # Sprint is a custom field, strictly expects an integer ID or None to clear
            try:
                val_str = str(value).strip().lower()
                if not val_str or val_str == "none" or val_str == "null":
                    update_data = {"fields": {"customfield_10020": None}}
                else:
                    sprint_id = int(str(value).strip())
                    update_data = {"fields": {"customfield_10020": sprint_id}}
            except ValueError:
                results.append({"key": key, "status": "error", "message": f"Invalid Sprint ID: '{value}'. A numeric ID or 'None' is required."})
                continue
        else:
            # Generic field update
            update_data = {"fields": {jira_field: value}}

        try:
            jira_res = requests.put(url, headers=headers_dict, json=update_data)
            if jira_res.status_code == 204:
                results.append({"key": key, "status": "success"})
            else:
                results.append({"key": key, "status": "error", "message": jira_res.text})
        except Exception as e:
            results.append({"key": key, "status": "error", "message": str(e)})

    return jsonify(results)

@app.route("/api/jira_metadata", methods=["GET"])
def jira_metadata():
    # Explicitly cast LocalProxy objects
    headers_dict = dict(HEADERS)
    
    # Fetch priorities and statuses for the project to populate dropdowns
    try:
        pri_res = requests.get(f"{JIRA_DOMAIN}/rest/api/3/priority", headers=headers_dict)
        priorities = [{"id": p["id"], "name": p["name"]} for p in pri_res.json()]
        
        # We could also fetch sprints if we had the board ID, but that's complex.
        # For now, just priorities and maybe projects?
        
        return jsonify({
            "priorities": priorities
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/sprints", methods=["GET"])
def get_sprints():
    # Explicitly cast LocalProxy objects
    headers_dict = dict(HEADERS)
    project_key_str = str(PROJECT_KEY)
    
    query = request.args.get("q", "").lower()
    
    # NEW APPROACH (Optimized & Comprehensive):
    # 1. We prioritize "Scrum" boards and check ALL of them (not just top 2/5).
    # 2. We fetch Active/Future sprints FIRST.
    # 3. If query is present, we ALSO check closed sprints (limited history).
    # 4. We use aggressive caching to make it fast.

    try:
        # 1. Find ALL Boards
        # Cache this for a long time (1 hour) as boards rarely change
        board_res = get_jira_cached(
            f"{JIRA_DOMAIN}/rest/agile/1.0/board",
            headers=headers_dict,
            params={"projectKeyOrId": project_key_str},
            ttl=3600 
        )
        
        boards = board_res.get("values", [])
        if not boards:
             return get_sprints_fallback(query, headers_dict, project_key_str)

        # Filter: Only Scrum boards are relevant for sprints usually
        # But some projects might mix types, so let's keep others but deprioritize them
        boards.sort(key=lambda b: 0 if b.get("type") == "scrum" else 1)
        
        # User requested: "dont search for payroll board" (or specifically avoid hardcoding it?)
        # Wait, previous request was "only need sprint from payroll board".
        # Current request: "dont search for payroll board show app now" -> "don't search for payroll board SPECIFICALLY, show ALL now"?
        # OR "don't search for payroll board, show APP now"?
        # Re-reading: "once user login make a call a cache record for sprint and dont search for payroll board show app now"
        # It sounds like: "Pre-fetch/cache sprints on login so it's fast, and stop the specific 'payroll board' filter I asked for earlier."
        
        # So I will REMOVE the payroll-specific filter and search ALL boards.
        
        target_boards = boards # Search ALL boards (or maybe top 10 to be safe from timeout)
        if len(target_boards) > 10:
            target_boards = target_boards[:10]

        all_sprints_map = {}
        
        for board in target_boards:
            board_id = board.get("id")
            sprints_url = f"{JIRA_DOMAIN}/rest/agile/1.0/board/{board_id}/sprint"
            
            # Fetch ALL Active and Future sprints
            # We page through results to ensure we get every single one
            try:
                start_at_sprint = 0
                max_results_sprint = 50
                
                while True:
                    # Cache Key needs to include board_id and startAt
                    # TTL 5 mins for active/future lists is reasonable
                    res = get_jira_cached(
                        sprints_url,
                        headers=headers_dict,
                        params={
                            "state": "active,future",
                            "startAt": start_at_sprint,
                            "maxResults": max_results_sprint
                        },
                        ttl=300 
                    )
                    
                    sprints_page = res.get("values", [])
                    if not sprints_page:
                        break
                        
                    for s in sprints_page:
                        s_id = s.get("id")
                        if s_id in all_sprints_map: continue
                        s_name = s.get("name", "")
                        
                        if query and query not in s_name.lower() and query not in str(s_id):
                            continue
                            
                        all_sprints_map[s_id] = {"id": s_id, "name": s_name, "state": s.get("state")}
                    
                    if res.get("isLast"):
                        break
                    start_at_sprint += len(sprints_page)
                    
            except Exception as e:
                print(f"DEBUG: Failed to fetch sprints for board {board_id}: {e}")

            # If searching for specific sprint (query present) and found nothing yet, check CLOSED sprints
            # Limit to last 50 closed sprints per board to avoid slowness
            if query and len(all_sprints_map) == 0:
                 try:
                     res_closed = get_jira_cached(
                        sprints_url,
                        headers=headers_dict,
                        params={
                            "state": "closed",
                            "maxResults": 50 
                        },
                        ttl=600 # Cache closed sprints for 10 mins
                    )
                     for s in res_closed.get("values", []):
                        s_id = s.get("id")
                        if s_id in all_sprints_map: continue
                        s_name = s.get("name", "")
                        if query in s_name.lower() or query in str(s_id):
                            all_sprints_map[s_id] = {"id": s_id, "name": s_name, "state": s.get("state")}
                 except Exception:
                     pass

        all_sprints = list(all_sprints_map.values())
        
        # Sort: Active first, then Future, then Closed (descending ID)
        state_order = {"active": 0, "future": 1, "closed": 2}
        all_sprints.sort(key=lambda x: (state_order.get(x["state"], 3), -x["id"]))
        
        return jsonify(all_sprints)
        
    except Exception as e:
        print(f"Error fetching sprints via Board API: {e}")
        return get_sprints_fallback(query, headers_dict, project_key_str)

def get_sprints_fallback(query, headers_dict, project_key_str):
    # Use JQL and paginate through results to collect all sprint references
    jql = f'project = "{project_key_str}"'
    url = f"{JIRA_DOMAIN}/rest/api/3/search/jql"
    
    try:
        all_sprints_map = {}
        start_at = 0
        max_results = 100

        while True:
            # Limit scan to recent 500 issues to avoid timeouts
            if start_at > 500: break
            
            params = {
                "jql": jql,
                "maxResults": max_results,
                "startAt": start_at,
                "fields": "customfield_10020"
            }
            res = requests.get(url, headers=headers_dict, params=params)
            if res.status_code != 200:
                return jsonify({"error": f"Jira error: {res.text}"}), res.status_code

            data = res.json()
            issues = data.get("issues", [])
            if not issues:
                break

            for issue in issues:
                sprints = issue["fields"].get("customfield_10020")
                if sprints and isinstance(sprints, list):
                    for s in sprints:
                        s_id = s.get("id")
                        s_name = s.get("name", "")
                        s_state = s.get("state", "unknown")
                        if not s_id:
                            continue
                        if query and query not in s_name.lower() and query not in str(s_id):
                            continue
                        all_sprints_map[s_id] = {
                            "id": s_id,
                            "name": s_name,
                            "state": s_state
                        }

            total = data.get("total", 0)
            start_at += len(issues)
            if start_at >= total:
                break
        
        final_sprints = list(all_sprints_map.values())
        # Sort by ID descending (newest first)
        final_sprints.sort(key=lambda x: x["id"], reverse=True)
        return jsonify(final_sprints)
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# =========================
# TEAM MANAGEMENT API
# =========================

@app.route("/teams")
@page_permission_required("teams")
def teams_page():
    return render_template("teams.html", project=PROJECT_KEY)

@app.route("/api/teams", methods=["GET", "POST"])
@login_required
def manage_teams():
    if request.method == "POST" and not current_user.has_permission("manage_teams") and current_user.role_name != 'Admin':
        return jsonify({"error": "Permission denied"}), 403
        
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    if request.method == "POST":
        name = request.json.get("name")
        if not name:
            return jsonify({"error": "Team name is required"}), 400
        cursor.execute("INSERT INTO teams (name) VALUES (%s)", (name,))
        conn.commit()
        team_id = cursor.lastrowid
        conn.close()
        return jsonify({"id": team_id, "name": name})
    else:
        cursor.execute("SELECT id, name, created_at FROM teams ORDER BY name ASC")
        teams = [{"id": r[0], "name": r[1], "created_at": r[2]} for r in cursor.fetchall()]
        conn.close()
        return jsonify(teams)

@app.route("/api/teams/<int:team_id>", methods=["DELETE"])
@permission_required("manage_teams")
def delete_team(team_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    cursor.execute("DELETE FROM teams WHERE id = %s", (team_id,))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

@app.route("/api/teams/<int:team_id>/members", methods=["GET", "POST"])
def team_members(team_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    if request.method == "POST":
        data = request.json
        account_id = data.get("accountId")
        display_name = data.get("displayName")
        avatar_url = data.get("avatarUrl")
        if not account_id or not display_name:
            return jsonify({"error": "Account ID and Display Name are required"}), 400
        cursor.execute("INSERT INTO team_members (team_id, account_id, display_name, avatar_url) VALUES (%s, %s, %s, %s)",
                       (team_id, account_id, display_name, avatar_url))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
    else:
        cursor.execute("SELECT id, account_id, display_name, avatar_url FROM team_members WHERE team_id = %s", (team_id,))
        members = [
            {"id": r[0], "accountId": r[1], "account_id": r[1], "displayName": r[2], "avatarUrl": r[3]}
            for r in cursor.fetchall()
        ]
        conn.close()
        return jsonify(members)

@app.route("/api/teams/<int:team_id>/members/<int:member_id>", methods=["DELETE"])
def delete_team_member(team_id, member_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    cursor.execute("DELETE FROM team_members WHERE id = %s AND team_id = %s", (member_id, team_id))
    conn.commit()
    conn.close()
    return jsonify({"success": True})

# =========================
# SPRINT PLANNING API
# =========================

@app.route("/planning/<int:team_id>")
@page_permission_required("planning")
def planning_page(team_id):
    return render_template("planning.html", project=PROJECT_KEY, team_id=team_id)

@app.route("/api/teams/<int:team_id>/sprints", methods=["GET", "POST"])
def team_sprints(team_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    if request.method == "POST":
        name = request.json.get("name")
        if not name:
            return jsonify({"error": "Sprint name is required"}), 400
        cursor.execute("INSERT INTO sprints (team_id, name) VALUES (%s, %s)", (team_id, name))
        conn.commit()
        sprint_id = cursor.lastrowid
        conn.close()
        return jsonify({"id": sprint_id, "name": name})
    else:
        cursor.execute("SELECT id, name, state, created_at FROM sprints WHERE team_id = %s ORDER BY created_at DESC", (team_id,))
        sprints = [{"id": r[0], "name": r[1], "state": r[2], "created_at": r[3]} for r in cursor.fetchall()]
        conn.close()
        return jsonify(sprints)

@app.route("/api/sprints/<int:sprint_id>/weeks", methods=["GET", "POST"])
def sprint_weeks(sprint_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    if request.method == "POST":
        data = request.json
        week_number = data.get("weekNumber")
        goal = data.get("goal", "")
        cursor.execute("INSERT INTO sprint_weeks (sprint_id, week_number, goal) VALUES (%s, %s, %s)",
                       (sprint_id, week_number, goal))
        conn.commit()
        week_id = cursor.lastrowid
        conn.close()
        return jsonify({"id": week_id, "weekNumber": week_number})
    else:
        cursor.execute("SELECT id, week_number, goal FROM sprint_weeks WHERE sprint_id = %s ORDER BY week_number ASC", (sprint_id,))
        weeks = [{"id": r[0], "weekNumber": r[1], "goal": r[2]} for r in cursor.fetchall()]
        conn.close()
        return jsonify(weeks)

@app.route("/api/sprint_weeks/<int:week_id>", methods=["PUT", "DELETE"])
def manage_sprint_week(week_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    if request.method == "DELETE":
        cursor.execute("DELETE FROM sprint_weeks WHERE id = %s", (week_id,))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
    elif request.method == "PUT":
        data = request.json
        goal = data.get("goal")
        cursor.execute("UPDATE sprint_weeks SET goal = %s WHERE id = %s", (goal, week_id))
        conn.commit()
        conn.close()
        return jsonify({"success": True})

@app.route("/api/sprints/<int:sprint_id>/tickets", methods=["GET", "POST"])
def sprint_tickets_api(sprint_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    if request.method == "POST":
        data = request.json
        week_id = data.get("weekId")
        issue_key = data.get("issueKey")
        if not issue_key:
            return jsonify({"error": "Issue key is required"}), 400
        cursor.execute("INSERT INTO sprint_tickets (sprint_id, week_id, issue_key) VALUES (%s, %s, %s)",
                       (sprint_id, week_id, issue_key))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
    else:
        cursor.execute("SELECT id, week_id, issue_key, comment, pr_raised, demo_done, pr_merged FROM sprint_tickets WHERE sprint_id = %s", (sprint_id,))
        tickets = [{"id": r[0], "weekId": r[1], "issueKey": r[2], "comment": r[3], "prRaised": bool(r[4]), "demoDone": bool(r[5]), "prMerged": bool(r[6])} for r in cursor.fetchall()]
        conn.close()
        return jsonify(tickets)

@app.route("/api/sprint_tickets/<int:ticket_id>", methods=["PUT", "DELETE"])
def manage_sprint_ticket(ticket_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    if request.method == "DELETE":
        cursor.execute("DELETE FROM sprint_tickets WHERE id = %s", (ticket_id,))
        conn.commit()
        conn.close()
        return jsonify({"success": True})
    elif request.method == "PUT":
        data = request.json
        comment = data.get("comment")
        pr_raised = 1 if data.get("prRaised") else 0
        demo_done = 1 if data.get("demoDone") else 0
        pr_merged = 1 if data.get("prMerged") else 0
        week_id = data.get("weekId") # New field for moving tickets
        
        # Build dynamic query
        fields = ["comment = %s", "pr_raised = %s", "demo_done = %s", "pr_merged = %s"]
        params = [comment, pr_raised, demo_done, pr_merged]
        
        if week_id is not None:
            fields.append("week_id = %s")
            params.append(week_id)
            
        params.append(ticket_id)
        
        cursor.execute(f"""
            UPDATE sprint_tickets 
            SET {', '.join(fields)}
            WHERE id = %s
        """, params)
        
        conn.commit()
        conn.close()
        return jsonify({"success": True})

@app.route("/status_tracker")
@page_permission_required("status_tracker")
def status_tracker():
    return render_template("status_tracker.html", project=PROJECT_KEY)

@app.route("/api/status_history", methods=["GET"])
def status_history():
    # Explicitly cast LocalProxy objects
    headers_dict = dict(HEADERS)
    project_key_str = str(PROJECT_KEY)

    keys_input = request.args.get("issue_keys", "")
    target_date = request.args.get("target_date", "") # Expected format YYYY-MM-DD
    
    # Process keys
    keys = re.split(r'[\s,]+', keys_input)
    keys = [k.strip().upper() for k in keys if k.strip()]
    
    # User's timezone offset (IST: +5:30)
    user_tz = timezone(timedelta(hours=5, minutes=30))

    # JQL to fetch these issues
    if keys:
        keys_str = ", ".join([f'"{k}"' for k in keys])
        jql = f"key IN ({keys_str})"
    elif target_date:
        # Fetch status changes for the target date and adjacent days to account for TZ shifts
        # JQL search is inclusive. Searching around the date helps capture all relevant issues.
        target_dt = datetime.strptime(target_date, "%Y-%m-%d")
        prev_day = (target_dt - timedelta(days=1)).strftime("%Y-%m-%d")
        next_day = (target_dt + timedelta(days=1)).strftime("%Y-%m-%d")
        jql = f'project = "{project_key_str}" AND status changed DURING ("{prev_day}", "{next_day}")'
    else:
        return jsonify([])
    
    url = f"{JIRA_DOMAIN}/rest/api/3/search/jql"
    print(f"DEBUG: HITTING STATUS_HISTORY. JQL: {jql}")
    
    params = {
        "jql": jql,
        "expand": "changelog",
        "fields": "summary,assignee,status,avatarUrls",
        "maxResults": 100
    }
    
    try:
        res = requests.get(url, headers=headers_dict, params=params)
        if res.status_code != 200:
            return jsonify({"error": f"Jira error: {res.text}"}), res.status_code
        
        issues = res.json().get("issues", [])
        updates = []
        
        for issue in issues:
            key = issue["key"]
            summary = issue["fields"]["summary"]
            changelog = issue.get("changelog", {}).get("histories", [])
            
            for history in changelog:
                # Jira timestamp example: 2026-02-15T16:05:05.123+0530
                # We need to handle the format which might have +HHMM instead of +HH:MM
                ts_str = history["created"]
                # Clean up timezone offset if it doesn't have a colon (common in Jira REST API)
                if "+" in ts_str and ts_str[-3] != ":":
                    ts_str = ts_str[:-2] + ":" + ts_str[-2:]
                elif "-" in ts_str and ts_str[-3] != ":":
                    ts_str = ts_str[:-2] + ":" + ts_str[-2:]
                
                try:
                    # fromisoformat handles +HH:MM in Python 3.7+
                    created_dt_utc = datetime.fromisoformat(ts_str)
                    # Convert to user's timezone (IST)
                    created_dt_local = created_dt_utc.astimezone(user_tz)
                    created_date_local = created_dt_local.strftime("%Y-%m-%d")
                except Exception as e:
                    print(f"DEBUG: Parsing error for {ts_str}: {e}")
                    created_date_local = history["created"][:10]
                
                # If target_date is provided, only show updates for that date in local time
                if target_date and created_date_local != target_date:
                    continue
                
                for item in history["items"]:
                    if item["field"] == "status":
                        updates.append({
                            "key": key,
                            "summary": summary,
                            "author": {
                                "name": history["author"]["displayName"],
                                "avatar": history["author"]["avatarUrls"]["24x24"]
                            },
                            "created": history["created"],
                            "from": item["fromString"],
                            "to": item["toString"]
                        })
        
        # Sort by creation time descending
        updates.sort(key=lambda x: x["created"], reverse=True)
        return jsonify(updates)
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/ticket_history/<key>", methods=["GET"])
def ticket_history(key):
    # Explicitly cast LocalProxy objects
    headers_dict = dict(HEADERS)

    url = f"{JIRA_DOMAIN}/rest/api/3/issue/{key}?expand=changelog"
    try:
        res = requests.get(url, headers=headers_dict)
        if res.status_code != 200:
            return jsonify({"error": f"Jira error: {res.text}"}), res.status_code
        
        data = res.json()
        changelog = data.get("changelog", {}).get("histories", [])
        
        formatted_history = []
        for history in changelog:
            for item in history["items"]:
                formatted_history.append({
                    "author": {
                        "name": history["author"]["displayName"],
                        "avatar": history["author"]["avatarUrls"]["24x24"]
                    },
                    "created": history["created"],
                    "field": item["field"],
                    "from": item["fromString"],
                    "to": item["toString"]
                })
        
        # Sort by creation time descending
        formatted_history.sort(key=lambda x: x["created"], reverse=True)
        return jsonify({
            "key": key,
            "summary": data["fields"]["summary"],
            "history": formatted_history
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# =========================
# SCRUM NOTES
# =========================

@app.route("/scrum_notes")
@page_permission_required("scrum_notes")
def scrum_notes_page():
    resp = make_response(render_template("scrum_notes.html", project=PROJECT_KEY, jira_domain=JIRA_DOMAIN))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp

@app.route("/api/jira/ticket/<key>", methods=["GET"])
def fetch_jira_ticket(key):
    """Fetch a single Jira ticket's details (summary, status, priority, assignee)."""
    # Explicitly cast LocalProxy objects
    headers_dict = dict(HEADERS)
    print(f"DEBUG fetch_jira_ticket: Headers are {dict(headers_dict)}")

    try:
        res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/issue/{key.upper()}",
            headers=headers_dict,
            params={"fields": "summary,status,priority,assignee,issuetype,description,reporter,created,updated,customfield_10016,customfield_10077,customfield_10020,labels,components,comment"}
        )
        print(f"DEBUG fetch_jira_ticket: Jira returned {res.status_code} {res.text}")
        if res.status_code == 404:
            return jsonify({"error": f"Ticket {key} not found"}), 404
        if res.status_code != 200:
            return jsonify({"error": res.text}), res.status_code
        data = res.json()
        f = data["fields"]
        
        # Extract Customer
        customer = "N/A"
        cust_val = f.get("customfield_10077")
        if cust_val:
            if isinstance(cust_val, list) and len(cust_val) > 0:
                v = cust_val[0]
                customer = v.get("value") or v if isinstance(v, dict) else str(v)
            elif isinstance(cust_val, dict):
                customer = cust_val.get("value") or str(cust_val)
            else:
                customer = str(cust_val)

        # Extract Sprint
        sprint = "N/A"
        sprint_val = f.get("customfield_10020")
        if sprint_val and isinstance(sprint_val, list) and len(sprint_val) > 0:
            sp = sprint_val[-1]
            sprint = sp.get("name") or str(sp)

        # Extract Story Points
        story_points = f.get("customfield_10016")

        # Extract Description (Jira v3 uses Doc format)
        description = f.get("description")

        return jsonify({
            "key": data["key"],
            "summary": f.get("summary"),
            "status": f.get("status", {}).get("name"),
            "priority": f.get("priority", {}).get("name"),
            "assignee": f.get("assignee", {}).get("displayName") if f.get("assignee") else "Unassigned",
            "assignee_avatar": f.get("assignee", {}).get("avatarUrls", {}).get("48x48") if f.get("assignee") else None,
            "type": f.get("issuetype", {}).get("name"),
            "type_icon": f.get("issuetype", {}).get("iconUrl"),
            "reporter": f.get("reporter", {}).get("displayName") if f.get("reporter") else "Unknown",
            "created": f.get("created"),
            "updated": f.get("updated"),
            "customer": customer,
            "sprint": sprint,
            "story_points": story_points,
            "description": description,
            "labels": f.get("labels", []),
            "components": [c.get("name") for c in f.get("components", [])],
            "comments": [
                {
                    "author": c.get("author", {}).get("displayName"),
                    "body": c.get("body"),
                    "created": c.get("created")
                } for c in f.get("comment", {}).get("comments", [])
            ] if f.get("comment") else []
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/scrum_notes", methods=["GET", "POST"])
def scrum_notes():
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500

    if request.method == "POST":
        data = request.json
        date = data.get("date")
        team_id = data.get("team_id")
        member_id = data.get("member_id")
        member_name = data.get("member_name")
        ticket_key = (data.get("ticket_key") or "").strip().upper()
        comment = data.get("comment", "")
        deadline = data.get("deadline") or None

        if isinstance(member_id, str):
            member_id = member_id.strip()

        invalid_member_ids = {None, "", "undefined", "null"}
        
        # If member_id is explicitly invalid, try to resolve it
        if (member_id in invalid_member_ids or member_id is None) and member_name and team_id:
            cursor.execute(
                """
                SELECT account_id
                FROM team_members
                WHERE team_id = %s
                  AND lower(trim(display_name)) = lower(trim(%s))
                LIMIT 1
                """,
                (team_id, member_name),
            )
            row = cursor.fetchone()
            if row:
                member_id = row[0]
            else:
                # If resolution fails, ensure we don't save "undefined"
                conn.close()
                return jsonify({"error": f"Could not resolve member ID for '{member_name}'. Please refresh the page."}), 400

        if not all([date, team_id, member_id, member_name, ticket_key]) or member_id in invalid_member_ids:
            conn.close()
            return jsonify({"error": "Missing required fields (member_id)"}), 400

        cursor.execute("""
            INSERT INTO scrum_notes (date, team_id, member_id, member_name, ticket_key, comment, deadline, tags)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
        """, (date, team_id, member_id, member_name, ticket_key, comment, deadline, data.get("tags", "")))
        conn.commit()
        new_id = cursor.lastrowid
        conn.close()
        return jsonify({"id": new_id, "success": True})

    else:  # GET
        date = request.args.get("date")
        team_id = request.args.get("team_id")
        if not date or not team_id:
            conn.close()
            return jsonify({"error": "date and team_id are required"}), 400
        cursor.execute("""
            SELECT id, date, team_id, member_id, member_name, ticket_key, comment, deadline, status, tags
            FROM scrum_notes
            WHERE date = %s AND team_id = %s
            ORDER BY member_name, created_at
        """, (date, team_id))
        rows = [{
            "id": r[0], "date": r[1], "team_id": r[2],
            "member_id": r[3], "member_name": r[4],
            "ticket_key": r[5], "comment": r[6],
            "deadline": r[7], "status": r[8], "tags": r[9]
        } for r in cursor.fetchall()]
        conn.close()
        return jsonify(rows)

@app.route("/api/scrum_notes/<int:note_id>", methods=["PUT", "DELETE"])
def scrum_note_item(note_id):
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500

    if request.method == "DELETE":
        cursor.execute("DELETE FROM scrum_notes WHERE id = %s", (note_id,))
        conn.commit()
        conn.close()
        return jsonify({"success": True})

    else:  # PUT
        data = request.json
        fields, params = [], []
        if "comment" in data:
            fields.append("comment = %s")
            params.append(data["comment"])
        if "deadline" in data:
            fields.append("deadline = %s")
            params.append(data["deadline"] or None)
        if "status" in data:
            fields.append("status = %s")
            params.append(data["status"])
        if "tags" in data:
            fields.append("tags = %s")
            params.append(data["tags"])
        if not fields:
            conn.close()
            return jsonify({"error": "No fields to update"}), 400
        params.append(note_id)
        cursor.execute(f"UPDATE scrum_notes SET {', '.join(fields)} WHERE id = %s", params)
        conn.commit()
        conn.close()
        return jsonify({"success": True})

@app.route("/api/scrum_notes/summary", methods=["GET"])
def scrum_notes_summary():
    """Fetch unique ticket keys worked on by a team/project in a date range."""
    start_date = request.args.get("start")
    end_date = request.args.get("end")
    team_id = request.args.get("team_id")
    
    if not start_date or not end_date:
        return jsonify({"error": "start and end dates are required"}), 400
        
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    # Logic: Get unique tickets from scrum_notes for the range, 
    # Aggregate member names, and pick the latest non-empty comment/tags
    query = """
        SELECT ticket_key, 
               GROUP_CONCAT(DISTINCT member_name) as members,
               GROUP_CONCAT(DISTINCT member_id) as member_ids
        FROM scrum_notes
        WHERE date >= %s AND date <= %s
    """
    params = [start_date, end_date]
    
    if team_id:
        query += " AND team_id = %s"
        params.append(team_id)
        
    query += " GROUP BY ticket_key"
    
    try:
        cursor.execute(query, tuple(params))
        rows = cursor.fetchall()
        
        results = []
        for r in rows:
            # For each ticket, fetch the latest non-empty comment and tags
            # MySQL doesn't support the subquery in the same way SQLite did in the SELECT list easily without performance issues or specific syntax, 
            # but we can do it separately or use a join. For simplicity, we'll do it separately or keep the subquery if MySQL likes it.
            # MySQL supports subqueries in SELECT.
            
            cursor.execute("SELECT comment FROM scrum_notes WHERE ticket_key = %s AND comment != '' ORDER BY date DESC, created_at DESC LIMIT 1", (r[0],))
            c_row = cursor.fetchone()
            latest_comment = c_row[0] if c_row else ""
            
            cursor.execute("SELECT tags FROM scrum_notes WHERE ticket_key = %s AND tags != '' ORDER BY date DESC, created_at DESC LIMIT 1", (r[0],))
            t_row = cursor.fetchone()
            latest_tags = t_row[0] if t_row else ""

            results.append({
                "ticket_key": r[0],
                "members": (r[1].split(",") if r[1] else []),
                "member_ids": (r[2].split(",") if r[2] else []),
                "comment": latest_comment,
                "tags": latest_tags
            })
            
        conn.close()
        return jsonify(results)
    except Exception as e:
        if conn: conn.close()
        return jsonify({"error": str(e)}), 500

@app.route("/api/scrum_notes/ticket/<key>", methods=["PUT"])
def scrum_note_by_ticket(key):
    """Update tags or comment for all notes of a specific ticket."""
    data = request.json
    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    
    fields, params = [], []
    if "comment" in data:
        fields.append("comment = %s")
        params.append(data["comment"])
    if "tags" in data:
        fields.append("tags = %s")
        params.append(data["tags"])
        
    if not fields:
        conn.close()
        return jsonify({"error": "No fields to update"}), 400
        
    try:
        # Update the LATEST note for this ticket
        # MySQL syntax for UPDATE with subquery on same table is tricky, using a temporary table or joining.
        # Simplest: Get the ID first.
        cursor.execute("SELECT id FROM scrum_notes WHERE ticket_key = %s ORDER BY date DESC, created_at DESC LIMIT 1", (key,))
        row = cursor.fetchone()
        if row:
            note_id = row[0]
            params.append(note_id)
            cursor.execute(f"UPDATE scrum_notes SET {', '.join(fields)} WHERE id = %s", params)
            conn.commit()
        
        conn.close()
        return jsonify({"success": True})
    except Exception as e:
        if conn: conn.close()
        return jsonify({"error": str(e)}), 500

@app.route("/api/scrum_notes/report", methods=["GET"])
def scrum_notes_report():
    start = request.args.get("start")
    end = request.args.get("end")
    team_id = request.args.get("team_id")
    member_id = request.args.get("member_id")

    if not start or not end:
        return jsonify({"error": "start and end dates are required"}), 400

    query = """
        SELECT id, date, team_id, member_id, member_name, ticket_key, comment, deadline, status
        FROM scrum_notes
        WHERE date BETWEEN %s AND %s
    """
    params = [start, end]

    if team_id:
        query += " AND team_id = %s"
        params.append(team_id)
    if member_id:
        query += " AND member_id = %s"
        params.append(member_id)

    query += " ORDER BY date DESC, member_name, ticket_key"

    conn, cursor = get_db_connection()
    if not conn: return jsonify({"error": "Database error"}), 500
    cursor.execute(query, params)
    rows = [{
        "id": r[0], "date": r[1], "team_id": r[2],
        "member_id": r[3], "member_name": r[4],
        "ticket_key": r[5], "comment": r[6],
        "deadline": r[7], "status": r[8]
    } for r in cursor.fetchall()]
    conn.close()
    return jsonify(rows)


@app.route("/work_report")
@page_permission_required("work_report")
def work_report():
    return render_template("work_report.html", project=PROJECT_KEY)


@app.route("/team_report")
def team_report_page():
    return render_template("team_report.html", project=PROJECT_KEY, show_navbar=False)


@app.route("/api/team_report/data", methods=["POST"])
def team_report_data():
    """Fetch Jira tickets for an arbitrary JQL query, returning rows with team info."""
    payload = request.get_json(silent=True) or {}
    jql = (payload.get("jql") or "").strip()
    if not jql:
        return jsonify({"error": "JQL is required"}), 400

    headers_dict = _get_explicit_request_jira_headers()
    if "Authorization" not in headers_dict:
        return jsonify({"error": "Add Jira email and API token in local settings first."}), 401

    fields = "summary,status,priority,issuetype,created,updated,assignee,labels,customfield_10001,customfield_10020"
    _plat_cf = _get_platform_checkboxes_field_id()
    if _plat_cf:
        fields += f",{_plat_cf}"

    all_issues = []
    seen_keys = set()
    start_at = 0
    next_page_token = None
    page_safety = 0
    while page_safety < 80:
        page_safety += 1
        params = {"jql": jql, "maxResults": 100, "fields": fields}
        if next_page_token:
            params["nextPageToken"] = next_page_token
        else:
            params["startAt"] = start_at
        jira_res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params=params,
        )
        if jira_res.status_code != 200:
            return jsonify({"error": f"Jira API error: {jira_res.text}"}), jira_res.status_code
        data = jira_res.json()
        issues = data.get("issues", [])
        if not issues:
            break
        new_count = 0
        for issue in issues:
            key = issue.get("key")
            if not key or key in seen_keys:
                continue
            seen_keys.add(key)
            all_issues.append(issue)
            new_count += 1
        if new_count == 0:
            break
        next_page_token = data.get("nextPageToken")
        if next_page_token:
            continue
        if len(issues) < 100:
            break
        start_at += 100

    rows = []
    for issue in all_issues:
        f = issue.get("fields") or {}
        status_name = ((f.get("status") or {}).get("name") or "Unknown").strip()
        priority_name = ((f.get("priority") or {}).get("name") or "Unspecified").strip()
        issue_type = ((f.get("issuetype") or {}).get("name") or "Task").strip()
        assignee = (f.get("assignee") or {}).get("displayName") or "Unassigned"
        created = (f.get("created") or "")[:10]
        updated = (f.get("updated") or "")[:10]
        labels = f.get("labels") or []
        # Team
        team_raw = f.get("customfield_10001")
        if isinstance(team_raw, dict):
            team_name = (team_raw.get("name") or team_raw.get("value") or "").strip()
        elif isinstance(team_raw, list) and team_raw:
            first = team_raw[0]
            team_name = (first.get("name") or first.get("value") or str(first)).strip() if isinstance(first, dict) else str(first).strip()
        elif team_raw:
            team_name = str(team_raw).strip()
        else:
            team_name = "Unspecified"
        # Sprints
        raw_sprints = f.get("customfield_10020") or []
        sprint_names = []
        if isinstance(raw_sprints, list):
            for s in raw_sprints:
                if isinstance(s, dict):
                    nm = s.get("name")
                    if nm:
                        sprint_names.append(nm)
                elif isinstance(s, str):
                    import re as _re
                    m = _re.search(r"name=([^,\]]+)", s)
                    if m:
                        sprint_names.append(m.group(1))
        # Platform / production
        _plat_raw = f.get(_plat_cf) if _plat_cf else None
        is_production = _jira_platform_includes_production(_plat_raw)

        rows.append({
            "issue_key": issue.get("key"),
            "summary": f.get("summary") or "",
            "status": status_name,
            "priority": priority_name,
            "issue_type": issue_type,
            "assignee": assignee,
            "team_name": team_name or "Unspecified",
            "sprints": sprint_names,
            "labels": labels,
            "created": created,
            "updated": updated,
            "is_production": is_production,
        })

    return jsonify({"rows": rows, "total": len(rows), "jql": jql})


@app.route("/api/team_report/ai_summary", methods=["POST"])
def team_report_ai_summary():
    """Generate a per-team AI summary using Claude."""
    payload = request.get_json(silent=True) or {}
    team_name = str(payload.get("team_name") or "").strip()
    tickets = payload.get("tickets") or []
    report_title = str(payload.get("report_title") or "").strip()
    model = payload.get("model") or "claude-sonnet-4-20250514"

    anthropic_key = (_get_app_config_value("anthropic_api_key") or "").strip()
    if not anthropic_key:
        return jsonify({"error": "Anthropic API key not configured."}), 400
    if not tickets:
        return jsonify({"error": "No tickets provided."}), 400

    _open_statuses = {"open", "to do", "new", "reopened"}
    compact_all = [
        {
            "key": t.get("issue_key"),
            "type": t.get("issue_type"),
            "status": t.get("status"),
            "priority": t.get("priority"),
            "summary": str(t.get("summary") or "")[:140],
            "assignee": t.get("assignee"),
        }
        for t in tickets[:100]
    ]
    # AI summary should focus on client pain points and avoid open-ticket narration.
    compact = [t for t in compact_all if str(t.get("status") or "").strip().lower() not in _open_statuses]
    if not compact:
        compact = compact_all[:40]
    tickets_json = json.dumps(compact, ensure_ascii=False)
    context = f" for the report '{report_title}'" if report_title else ""

    prompt = (
        f"You are a customer-impact analyst. Below is a JSON list of Jira tickets for the '{team_name}' team{context}.\n"
        "Write concise bullets ONLY about customer-facing problems, pain points, and unresolved needs.\n"
        "Rules:\n"
        "- DO NOT describe internal implementation, team activity, or engineering work.\n"
        "- DO NOT use future-tense action language (e.g., 'will fix', 'working on', 'we are fixing').\n"
        "- DO NOT mention ticket workflow/state words like 'open tickets', 'in progress', 'closed'.\n"
        "- Keep each bullet factual, neutral, and client-impact focused.\n"
        "Return ONLY a JSON array of 4-8 bullet strings (plain English, no markdown), e.g.:\n"
        '[\"Payroll export mismatch creates reconciliation delays for client finance teams\"]\n\n'
        f"Tickets:\n{tickets_json}"
    )

    try:
        res = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": anthropic_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": model,
                "max_tokens": 800,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=60,
        )
        if not res.ok:
            return jsonify({"error": f"Claude HTTP {res.status_code}"}), 502
        body = res.json()
        raw = (body.get("content") or [{}])[0].get("text", "").strip()
        raw = raw.strip("`").strip()
        if raw.startswith("json"):
            raw = raw[4:].strip()
        bullets = json.loads(raw)
        if not isinstance(bullets, list):
            bullets = [str(bullets)]
        return jsonify({"team": team_name, "bullets": bullets})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ── Sprint Delivery Dashboard ─────────────────────────────────────────────────

@app.route("/sprint_delivery")
@page_permission_required("sprint_delivery")
def sprint_delivery_page():
    return render_template("sprint_delivery.html", project=PROJECT_KEY, show_navbar=False)


@app.route("/payroll_dashboard")
@page_permission_required("payroll_dashboard")
def payroll_dashboard_page():
    return render_template("payroll_dashboard.html", project=PROJECT_KEY)


def _first_done_transition_at(issue, fields_obj):
    histories = (issue.get("changelog") or {}).get("histories") or []
    hits = []
    for history in histories:
        ts = _parse_jira_datetime(history.get("created"))
        if not ts:
            continue
        for item in history.get("items") or []:
            if item.get("field") != "status":
                continue
            to_s = item.get("toString") or ""
            if _is_done_like_status(to_s, None):
                hits.append(ts)
    if hits:
        return min(hits)
    for k in ("resolutiondate", "updated"):
        fallback = _parse_jira_datetime((fields_obj or {}).get(k))
        if fallback:
            return fallback
    return None


def _safe_dt_days(start_dt, end_dt):
    if not start_dt or not end_dt:
        return None
    delta = end_dt - start_dt
    return round(max(delta.total_seconds(), 0) / 86400.0, 2)


def _linked_bug_count(fields_obj):
    links = (fields_obj or {}).get("issuelinks") or []
    bug_count = 0
    for link in links:
        linked = link.get("inwardIssue") or link.get("outwardIssue") or {}
        l_fields = linked.get("fields") or {}
        l_type = ((l_fields.get("issuetype") or {}).get("name") or "").strip().lower()
        if l_type == "bug":
            bug_count += 1
    return bug_count


def _load_payroll_meta_map(issue_keys):
    if not issue_keys:
        return {}
    placeholders = ",".join(["%s"] * len(issue_keys))
    conn, cursor = get_db_connection(dictionary=True)
    if not conn:
        return {}
    try:
        cursor.execute(
            f"""SELECT issue_key, manual_bug_count, manual_bug_links, test_cases_count, notes
                FROM payroll_ticket_meta
                WHERE issue_key IN ({placeholders})""",
            tuple(issue_keys),
        )
        rows = cursor.fetchall() or []
        return {r["issue_key"]: r for r in rows}
    except Exception:
        return {}
    finally:
        conn.close()


@app.route("/api/payroll_dashboard/data", methods=["POST"])
@login_required
def payroll_dashboard_data():
    if not current_user.can_view_page("payroll_dashboard"):
        return jsonify({"error": "Access denied."}), 403
    project_key_str = str(PROJECT_KEY).strip()
    if not project_key_str:
        return jsonify({"error": "Missing project key. Save it in Settings first."}), 400
    headers_dict = dict(HEADERS)
    if "Authorization" not in headers_dict:
        return jsonify({"error": "Missing Jira credentials. Save them in Settings first."}), 401

    payload = request.get_json(silent=True) or {}
    sprint_ids = payload.get("sprint_ids") or []
    if not isinstance(sprint_ids, list):
        sprint_ids = []
    sprint_ids = [int(s) for s in sprint_ids if str(s).strip().isdigit()]
    if not sprint_ids:
        return jsonify({"error": "At least one sprint_id is required."}), 400

    sprint_clause = ", ".join(str(s) for s in sprint_ids)
    jql = f'project = "{project_key_str}" AND sprint in ({sprint_clause}) ORDER BY created DESC'
    fields = (
        "summary,status,priority,issuetype,assignee,labels,customfield_10020,"
        "issuelinks,created,updated,resolutiondate"
    )
    _plat_cf = _get_platform_checkboxes_field_id()
    if _plat_cf:
        fields += f",{_plat_cf}"

    all_issues = []
    start_at = 0
    for _ in range(80):
        params = {
            "jql": jql,
            "maxResults": 100,
            "fields": fields,
            "expand": "changelog",
            "startAt": start_at,
        }
        try:
            res = requests.get(
                f"{JIRA_DOMAIN}/rest/api/3/search/jql",
                headers=headers_dict,
                params=params,
                timeout=45,
            )
            if not res.ok:
                return jsonify({"error": f"Jira API error: {res.text}"}), res.status_code
            body = res.json()
            page = body.get("issues") or []
            all_issues.extend(page)
            if len(page) < 100:
                break
            start_at += len(page)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500

    seen = set()
    unique = []
    for issue in all_issues:
        k = issue.get("key")
        if k and k not in seen:
            seen.add(k)
            unique.append(issue)

    issue_keys = [i.get("key") for i in unique if i.get("key")]
    meta_map = _load_payroll_meta_map(issue_keys)

    by_sprint = {}
    all_rows = []
    for issue in unique:
        fields_obj = issue.get("fields") or {}
        raw_sprints = _extract_sprint_values(fields_obj.get("customfield_10020") or [])
        sprint_ids_on_issue = [s["id"] for s in raw_sprints if s.get("id")]
        matched = [sid for sid in sprint_ids_on_issue if sid in sprint_ids]

        status_obj = fields_obj.get("status") or {}
        status_name = (status_obj.get("name") or "Unknown").strip()
        status_category = ((status_obj.get("statusCategory") or {}).get("key") or "").strip()
        is_delivered = _is_done_like_status(status_name, status_category)
        created_at = _parse_jira_datetime(fields_obj.get("created"))
        done_at = _first_done_transition_at(issue, fields_obj) if is_delivered else None
        cycle_days = _safe_dt_days(created_at, done_at)
        jira_bug_count = _linked_bug_count(fields_obj)
        meta = meta_map.get(issue.get("key"), {}) or {}
        manual_bug_count = int(meta.get("manual_bug_count") or 0)
        total_bug_count = jira_bug_count + manual_bug_count
        _plat_raw = fields_obj.get(_plat_cf) if _plat_cf else None

        row = {
            "issue_key": issue.get("key"),
            "summary": fields_obj.get("summary") or "",
            "status": status_name,
            "priority": ((fields_obj.get("priority") or {}).get("name") or "Unspecified").strip(),
            "issue_type": ((fields_obj.get("issuetype") or {}).get("name") or "Task").strip(),
            "assignee": ((fields_obj.get("assignee") or {}).get("displayName") or "Unassigned"),
            "labels": fields_obj.get("labels") or [],
            "sprints": raw_sprints,
            "matched_sprint_ids": matched,
            "is_delivered": is_delivered,
            "created": fields_obj.get("created"),
            "done_at": done_at.isoformat() if done_at else None,
            "cycle_days": cycle_days,
            "jira_bug_count": jira_bug_count,
            "manual_bug_count": manual_bug_count,
            "total_bug_count": total_bug_count,
            "is_production": _jira_platform_includes_production(_plat_raw),
            "manual_bug_links": meta.get("manual_bug_links") or "",
            "test_cases_count": int(meta.get("test_cases_count") or 0),
            "notes": meta.get("notes") or "",
        }
        all_rows.append(row)
        for sid in matched:
            by_sprint.setdefault(sid, []).append(row)

    sprint_summaries = []
    for sid in sprint_ids:
        rows = by_sprint.get(sid, [])
        delivered = [r for r in rows if r["is_delivered"]]
        spilled = [r for r in rows if not r["is_delivered"]]
        cycle_vals = [r["cycle_days"] for r in delivered if isinstance(r.get("cycle_days"), (int, float))]
        avg_cycle = round(sum(cycle_vals) / len(cycle_vals), 2) if cycle_vals else None
        avg_bugs = round(sum((r.get("total_bug_count") or 0) for r in rows) / len(rows), 2) if rows else 0
        bug_density_pct = round((sum((r.get("total_bug_count") or 0) for r in rows) / len(rows)) * 100, 2) if rows else 0
        avg_test_cases = round(sum((r.get("test_cases_count") or 0) for r in rows) / len(rows), 2) if rows else 0
        sprint_summaries.append({
            "sprint_id": sid,
            "total": len(rows),
            "delivered_count": len(delivered),
            "spilled_count": len(spilled),
            "avg_cycle_days": avg_cycle,
            "avg_bugs_per_ticket": avg_bugs,
            "avg_bug_percentage_score": bug_density_pct,
            "avg_test_cases_per_ticket": avg_test_cases,
            "tickets": rows,
            "delivered": delivered,
            "spilled": spilled,
        })

    total_rows = len(all_rows)
    all_delivered = [r for r in all_rows if r["is_delivered"]]
    all_cycle_vals = [r["cycle_days"] for r in all_delivered if isinstance(r.get("cycle_days"), (int, float))]
    return jsonify({
        "sprints": sprint_summaries,
        "overall": {
            "total": total_rows,
            "delivered_count": len(all_delivered),
            "spilled_count": total_rows - len(all_delivered),
            "avg_cycle_days": round(sum(all_cycle_vals) / len(all_cycle_vals), 2) if all_cycle_vals else None,
            "avg_bugs_per_ticket": round(sum((r.get("total_bug_count") or 0) for r in all_rows) / total_rows, 2) if total_rows else 0,
            "avg_test_cases_per_ticket": round(sum((r.get("test_cases_count") or 0) for r in all_rows) / total_rows, 2) if total_rows else 0,
        },
        "jql": jql,
    })


@app.route("/api/payroll_dashboard/ticket_meta/<issue_key>", methods=["PUT"])
@login_required
def payroll_dashboard_upsert_ticket_meta(issue_key):
    if not current_user.can_view_page("payroll_dashboard"):
        return jsonify({"error": "Access denied."}), 403
    issue_key = (issue_key or "").strip().upper()
    if not issue_key:
        return jsonify({"error": "issue_key is required."}), 400
    payload = request.get_json(silent=True) or {}
    manual_bug_count = payload.get("manual_bug_count", 0)
    test_cases_count = payload.get("test_cases_count", 0)
    manual_bug_links = (payload.get("manual_bug_links") or "").strip()
    notes = (payload.get("notes") or "").strip()
    try:
        manual_bug_count = max(int(manual_bug_count), 0)
    except Exception:
        manual_bug_count = 0
    try:
        test_cases_count = max(int(test_cases_count), 0)
    except Exception:
        test_cases_count = 0

    conn, cursor = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        cursor.execute(
            """INSERT INTO payroll_ticket_meta
               (issue_key, manual_bug_count, manual_bug_links, test_cases_count, notes, updated_by)
               VALUES (%s, %s, %s, %s, %s, %s)
               ON DUPLICATE KEY UPDATE
               manual_bug_count = VALUES(manual_bug_count),
               manual_bug_links = VALUES(manual_bug_links),
               test_cases_count = VALUES(test_cases_count),
               notes = VALUES(notes),
               updated_by = VALUES(updated_by)""",
            (issue_key, manual_bug_count, manual_bug_links, test_cases_count, notes, getattr(current_user, "email", "system")),
        )
        conn.commit()
        return jsonify({"ok": True})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500
    finally:
        conn.close()


@app.route("/api/payroll_dashboard/ai_summary", methods=["POST"])
@login_required
def payroll_dashboard_ai_summary():
    if not current_user.can_view_page("payroll_dashboard"):
        return jsonify({"error": "Access denied."}), 403
    payload = request.get_json(silent=True) or {}
    sprint_name = str(payload.get("sprint_name") or "").strip()
    delivered = payload.get("delivered") or []
    spilled = payload.get("spilled") or []
    model = payload.get("model") or "claude-sonnet-4-20250514"

    anthropic_key = (_get_app_config_value("anthropic_api_key") or "").strip()
    if not anthropic_key:
        return jsonify({"error": "Anthropic API key not configured."}), 400
    if not delivered and not spilled:
        return jsonify({"error": "No tickets provided."}), 400

    def compact(rows, limit=70):
        out = []
        for r in rows[:limit]:
            out.append({
                "key": r.get("issue_key"),
                "type": r.get("issue_type"),
                "status": r.get("status"),
                "summary": str(r.get("summary") or "")[:140],
                "cycle_days": r.get("cycle_days"),
                "bugs": r.get("total_bug_count"),
                "test_cases": r.get("test_cases_count"),
            })
        return out

    prompt = (
        f"You are an engineering delivery analyst. Sprint: '{sprint_name or 'Selected sprint'}'.\n"
        f"Delivered tickets: {json.dumps(compact(delivered), ensure_ascii=False)}\n"
        f"Spilled tickets: {json.dumps(compact(spilled), ensure_ascii=False)}\n\n"
        "Write a short dashboard summary in JSON with keys:\n"
        "- delivered_summary: 2-4 bullet strings\n"
        "- spilled_summary: 1-3 bullet strings\n"
        "- quality_summary: 1-3 bullet strings mentioning bug trend and test-case trend\n"
        "Return only a valid JSON object."
    )
    try:
        res = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": anthropic_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": model,
                "max_tokens": 800,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=60,
        )
        if not res.ok:
            return jsonify({"error": f"Claude HTTP {res.status_code}"}), 502
        body = res.json()
        raw = (body.get("content") or [{}])[0].get("text", "").strip().strip("`").strip()
        if raw.startswith("json"):
            raw = raw[4:].strip()
        result = json.loads(raw)
        return jsonify({"sprint": sprint_name, "summary": result})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/sprint_delivery/teams", methods=["GET"])
def sprint_delivery_teams():
    """Return unique team names from recent Jira issues."""
    project_key_str = str(PROJECT_KEY).strip()
    if not project_key_str:
        return jsonify({"error": "Missing project key. Save it in Settings first."}), 400
    headers_dict = dict(HEADERS)
    if "Authorization" not in headers_dict:
        return jsonify({"error": "Missing Jira credentials. Save them in Settings first."}), 401

    jql = f'project = "{project_key_str}" ORDER BY created DESC'
    params = {"jql": jql, "maxResults": 500, "fields": "customfield_10001"}
    try:
        res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params=params,
            timeout=30,
        )
        if not res.ok:
            return jsonify({"error": f"Jira API error: {res.text}"}), res.status_code
        teams = set()
        for issue in (res.json().get("issues") or []):
            t = _team_name_from_jira_fields(issue.get("fields") or {})
            if t and t.lower() != "unspecified":
                teams.add(t)
        return jsonify({"teams": sorted(teams, key=lambda x: x.lower())})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/sprint_delivery/data", methods=["POST"])
def sprint_delivery_data():
    """Fetch tickets for selected team + sprint(s), classify as delivered or spilled."""
    project_key_str = str(PROJECT_KEY).strip()
    if not project_key_str:
        return jsonify({"error": "Missing project key. Save it in Settings first."}), 400
    headers_dict = dict(HEADERS)
    if "Authorization" not in headers_dict:
        return jsonify({"error": "Missing Jira credentials. Save them in Settings first."}), 401

    payload = request.get_json(silent=True) or {}
    team_name = (payload.get("team_name") or "").strip()
    sprint_ids = payload.get("sprint_ids") or []
    if not isinstance(sprint_ids, list):
        sprint_ids = []
    sprint_ids = [int(s) for s in sprint_ids if str(s).strip().isdigit()]

    if not sprint_ids:
        return jsonify({"error": "At least one sprint_id is required."}), 400

    sprint_clause = ", ".join(str(s) for s in sprint_ids)
    jql = f'project = "{project_key_str}" AND sprint in ({sprint_clause})'
    if team_name:
        jql += f' AND team = "{team_name}"'
    jql += " ORDER BY created DESC"

    fields = "summary,status,priority,issuetype,assignee,labels,customfield_10001,customfield_10020"

    all_issues = []
    start_at = 0
    for _ in range(80):
        params = {"jql": jql, "maxResults": 100, "fields": fields, "startAt": start_at}
        try:
            res = requests.get(
                f"{JIRA_DOMAIN}/rest/api/3/search/jql",
                headers=headers_dict,
                params=params,
                timeout=30,
            )
            if not res.ok:
                break
            body = res.json()
            page = body.get("issues") or []
            all_issues.extend(page)
            if len(page) < 100:
                break
            npt = body.get("nextPageToken")
            if npt:
                start_at = npt
            else:
                start_at += len(page)
        except Exception:
            break

    # Deduplicate
    seen = set()
    unique = []
    for issue in all_issues:
        k = issue.get("key")
        if k and k not in seen:
            seen.add(k)
            unique.append(issue)

    rows = []
    for issue in unique:
        f = issue.get("fields") or {}
        status_name = ((f.get("status") or {}).get("name") or "Unknown").strip()
        status_category = ((f.get("status") or {}).get("statusCategory") or {}).get("key", "")
        priority_name = ((f.get("priority") or {}).get("name") or "Unspecified").strip()
        issue_type = ((f.get("issuetype") or {}).get("name") or "Task").strip()
        assignee = ((f.get("assignee") or {}).get("displayName") or "Unassigned")
        labels = f.get("labels") or []
        team = _team_name_from_jira_fields(f)

        raw_sprints = f.get("customfield_10020") or []
        sprint_names = []
        sprint_ids_on_issue = []
        if isinstance(raw_sprints, list):
            for s in raw_sprints:
                if isinstance(s, dict):
                    nm = s.get("name")
                    sid = s.get("id")
                    if nm:
                        sprint_names.append(nm)
                    if sid:
                        sprint_ids_on_issue.append(int(sid))
                elif isinstance(s, str):
                    import re as _re
                    m_name = _re.search(r"name=([^,\]]+)", s)
                    m_id = _re.search(r"id=(\d+)", s)
                    if m_name:
                        sprint_names.append(m_name.group(1))
                    if m_id:
                        sprint_ids_on_issue.append(int(m_id.group(1)))

        # Match which of the requested sprints this ticket belongs to
        matched_sprint_ids = [sid for sid in sprint_ids_on_issue if sid in sprint_ids]

        is_delivered = _is_done_like_status(status_name, status_category)

        rows.append({
            "issue_key": issue.get("key"),
            "summary": f.get("summary") or "",
            "status": status_name,
            "priority": priority_name,
            "issue_type": issue_type,
            "assignee": assignee,
            "team_name": team,
            "sprint_names": sprint_names,
            "sprint_ids": sprint_ids_on_issue,
            "matched_sprint_ids": matched_sprint_ids,
            "labels": labels,
            "is_delivered": is_delivered,
        })

    delivered = [r for r in rows if r["is_delivered"]]
    spilled = [r for r in rows if not r["is_delivered"]]

    return jsonify({
        "delivered": delivered,
        "spilled": spilled,
        "total": len(rows),
        "delivered_count": len(delivered),
        "spilled_count": len(spilled),
    })


@app.route("/api/sprint_delivery/ai_summary", methods=["POST"])
def sprint_delivery_ai_summary():
    """Generate an AI summary for a sprint delivery using Claude."""
    payload = request.get_json(silent=True) or {}
    sprint_name = str(payload.get("sprint_name") or "").strip()
    delivered = payload.get("delivered") or []
    spilled = payload.get("spilled") or []
    model = payload.get("model") or "claude-sonnet-4-20250514"

    anthropic_key = (_get_app_config_value("anthropic_api_key") or "").strip()
    if not anthropic_key:
        return jsonify({"error": "Anthropic API key not configured."}), 400
    if not delivered and not spilled:
        return jsonify({"error": "No tickets provided."}), 400

    def compact(tickets, limit=60):
        return [
            {
                "key": t.get("issue_key"),
                "type": t.get("issue_type"),
                "status": t.get("status"),
                "priority": t.get("priority"),
                "summary": str(t.get("summary") or "")[:140],
                "assignee": t.get("assignee"),
            }
            for t in tickets[:limit]
        ]

    context_label = f" for sprint '{sprint_name}'" if sprint_name else ""
    prompt = (
        f"You are a delivery analyst. Below are Jira tickets{context_label}.\n\n"
        f"DELIVERED ({len(delivered)} tickets): {json.dumps(compact(delivered), ensure_ascii=False)}\n\n"
        f"SPILLED OVER ({len(spilled)} tickets): {json.dumps(compact(spilled), ensure_ascii=False)}\n\n"
        "Write a concise sprint delivery summary with these sections:\n"
        "1. What was delivered (2-4 bullets about key outcomes and customer value)\n"
        "2. What spilled over (1-3 bullets about unfinished work and potential impact)\n"
        "3. Overall health (1 bullet: delivery rate assessment)\n\n"
        "Rules:\n"
        "- Focus on customer-facing impact, not internal engineering details.\n"
        "- Keep each bullet short and factual.\n"
        "- Return ONLY a JSON object with keys 'delivered_summary', 'spilled_summary', 'health' — each being an array of bullet strings.\n"
        "Example: {\"delivered_summary\": [\"...\"], \"spilled_summary\": [\"...\"], \"health\": [\"...\"]}"
    )

    try:
        res = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": anthropic_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": model,
                "max_tokens": 800,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=60,
        )
        if not res.ok:
            return jsonify({"error": f"Claude HTTP {res.status_code}"}), 502
        body = res.json()
        raw = (body.get("content") or [{}])[0].get("text", "").strip()
        raw = raw.strip("`").strip()
        if raw.startswith("json"):
            raw = raw[4:].strip()
        result = json.loads(raw)
        return jsonify({"sprint": sprint_name, "summary": result})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ── Assignee Work ─────────────────────────────────────────────────────────────

@app.route("/assignee_work")
@page_permission_required("assignee_work")
def assignee_work_page():
    return render_template("assignee_work.html", project=PROJECT_KEY)


@app.route("/api/assignee_work", methods=["POST"])
def assignee_work_search():
    """Fetch Jira tickets filtered by date range (updated), optional labels,
    sprint id, and assignee(s). Returns a flat JSON list suitable for the
    Assignee Work page table."""
    headers_dict = dict(HEADERS)
    project_key_str = str(PROJECT_KEY)

    if not project_key_str:
        return jsonify({"error": "Missing project key. Save it in Settings first."}), 400
    if "Authorization" not in headers_dict:
        return jsonify({"error": "Missing Jira credentials. Save them in Settings first."}), 401

    data = request.get_json(silent=True) or {}
    from_date = (data.get("from_date") or "").strip()
    to_date = (data.get("to_date") or "").strip()
    labels = data.get("labels") or []
    sprint_id = data.get("sprint_id")
    assignee_ids = data.get("assignee_ids") or []
    if not isinstance(assignee_ids, list):
        assignee_ids = []
    assignee_ids = [str(a).strip() for a in assignee_ids if str(a).strip()]
    # Backward compatibility for older client payload shape
    if not assignee_ids:
        single_assignee = (data.get("assignee_id") or "").strip() if data.get("assignee_id") else ""
        if single_assignee:
            assignee_ids = [single_assignee]

    if not from_date or not to_date:
        return jsonify({"error": "from_date and to_date are required."}), 400

    # Build JQL
    jql_parts = [f'project = "{project_key_str}"']
    jql_parts.append(f'updated >= "{from_date} 00:00"')
    jql_parts.append(f'updated <= "{to_date} 23:59"')

    if assignee_ids:
        quoted_assignees = ", ".join(f'"{a}"' for a in assignee_ids)
        jql_parts.append(f'assignee in ({quoted_assignees})')

    if sprint_id:
        jql_parts.append(f'sprint = {int(sprint_id)}')

    if isinstance(labels, list) and labels:
        cleaned = [str(l).strip() for l in labels if str(l).strip()]
        if cleaned:
            quoted = ", ".join(f'"{l}"' for l in cleaned)
            jql_parts.append(f'labels in ({quoted})')

    jql = " AND ".join(jql_parts) + " ORDER BY updated DESC"

    try:
        _base_aw_fields = "summary,status,assignee,priority,created,updated,issuetype,labels,customfield_10020"
        _plat_cf = _get_platform_checkboxes_field_id()
        _fields = f"{_base_aw_fields},{_plat_cf}" if _plat_cf else _base_aw_fields
        params = {
            "jql": jql,
            "maxResults": 100,
            "startAt": 0,
            "fields": _fields
        }
        res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params=params,
            timeout=30
        )

        if res.status_code != 200:
            try:
                err_body = res.json()
            except Exception:
                err_body = {"error": res.text}
            msg = err_body.get("errorMessages") or err_body.get("error") or f"Jira API error ({res.status_code})"
            return jsonify({"error": msg, "jql": jql}), res.status_code

        payload = res.json()
        issues = payload.get("issues", [])
        formatted = []
        for i in issues:
            f = i.get("fields", {}) or {}
            raw_sprints = f.get("customfield_10020") or []
            sprint_names = []
            if isinstance(raw_sprints, list):
                for s in raw_sprints:
                    if isinstance(s, dict):
                        nm = s.get("name")
                        if nm:
                            sprint_names.append(nm)
                    elif isinstance(s, str):
                        # Legacy Jira returns a string blob; try to extract name=
                        import re as _re
                        m = _re.search(r"name=([^,\]]+)", s)
                        if m:
                            sprint_names.append(m.group(1))

            _plat_raw = f.get(_plat_cf) if _plat_cf else None
            formatted.append({
                "key": i.get("key"),
                "summary": f.get("summary"),
                "status": (f.get("status") or {}).get("name"),
                "assignee": (f.get("assignee") or {}).get("displayName") if f.get("assignee") else "Unassigned",
                "priority": (f.get("priority") or {}).get("name"),
                "type": (f.get("issuetype") or {}).get("name"),
                "labels": f.get("labels") or [],
                "sprints": sprint_names,
                "created": f.get("created"),
                "updated": f.get("updated"),
                "platform_production": _jira_platform_includes_production(_plat_raw),
            })

        return jsonify({
            "issues": formatted,
            "total": payload.get("total", len(formatted)),
            "jql": jql
        })
    except requests.Timeout:
        return jsonify({"error": "Jira request timed out.", "jql": jql}), 504
    except Exception as e:
        return jsonify({"error": str(e), "jql": jql}), 500


@app.route("/api/assignee_work_by_sprint", methods=["POST"])
def assignee_work_by_sprint():
    """Fetch ALL tickets for a named sprint (no date filter) using
    JQL: project = X AND sprint = "sprint_name" ORDER BY updated DESC.
    Returns the same flat issue format as /api/assignee_work."""
    headers_dict = dict(HEADERS)
    project_key_str = str(PROJECT_KEY)

    if not project_key_str:
        return jsonify({"error": "Missing project key."}), 400
    if "Authorization" not in headers_dict:
        return jsonify({"error": "Missing Jira credentials."}), 401

    data = request.get_json(silent=True) or {}
    sprint_name = (data.get("sprint_name") or "").strip()
    if not sprint_name:
        return jsonify({"error": "sprint_name is required."}), 400

    # Escape double quotes inside the sprint name
    safe_sprint = sprint_name.replace('"', '\\"')
    jql = (
        f'project = "{project_key_str}" AND sprint = "{safe_sprint}"'
        ' ORDER BY updated DESC'
    )

    try:
        _base_aw_fields = "summary,status,assignee,priority,created,updated,issuetype,labels,customfield_10020"
        _plat_cf = _get_platform_checkboxes_field_id()
        _fields = f"{_base_aw_fields},{_plat_cf}" if _plat_cf else _base_aw_fields
        params = {"jql": jql, "maxResults": 200, "startAt": 0, "fields": _fields}
        res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict, params=params, timeout=30
        )
        if res.status_code != 200:
            try:
                err_body = res.json()
            except Exception:
                err_body = {}
            msg = (err_body.get("errorMessages") or
                   err_body.get("error") or
                   f"Jira API error ({res.status_code})")
            return jsonify({"error": msg, "jql": jql}), res.status_code

        payload = res.json()
        issues  = payload.get("issues", [])
        formatted = []
        import re as _re
        for i in issues:
            f = i.get("fields") or {}
            sprint_names = []
            raw_sprints = f.get("customfield_10020") or []
            if isinstance(raw_sprints, list):
                for sp in raw_sprints:
                    if isinstance(sp, dict):
                        sprint_names.append(sp.get("name", ""))
                    elif isinstance(sp, str):
                        m = _re.search(r'name=([^,\]]+)', sp)
                        if m:
                            sprint_names.append(m.group(1))
            _plat_raw = f.get(_plat_cf) if _plat_cf else None
            formatted.append({
                "key":      i.get("key"),
                "summary":  f.get("summary"),
                "status":   (f.get("status") or {}).get("name"),
                "assignee": (f.get("assignee") or {}).get("displayName") if f.get("assignee") else "Unassigned",
                "priority": (f.get("priority") or {}).get("name"),
                "type":     (f.get("issuetype") or {}).get("name"),
                "labels":   f.get("labels") or [],
                "sprints":  sprint_names,
                "created":  f.get("created"),
                "updated":  f.get("updated"),
                "platform_production": _jira_platform_includes_production(_plat_raw),
            })

        return jsonify({"issues": formatted, "total": len(formatted), "jql": jql})

    except requests.Timeout:
        return jsonify({"error": "Jira request timed out.", "jql": jql}), 504
    except Exception as e:
        return jsonify({"error": str(e), "jql": jql}), 500


@app.route("/api/assignee_work_ai_summary", methods=["POST"])
def assignee_work_ai_summary():
    """
    Accept a list of Jira tickets for a chosen sprint and return two AI summaries:
      - overall_bullets : what the team accomplished this sprint
      - production_bullets : production issue fixes summary
    Uses the server-side anthropic_api_key from app config.
    """
    data = request.get_json(silent=True) or {}
    sprint_name = str(data.get("sprint") or "").strip()
    tickets = data.get("tickets") or []
    model = data.get("model") or "claude-sonnet-4-20250514"

    anthropic_key = (_get_app_config_value("anthropic_api_key") or "").strip()
    if not anthropic_key:
        return jsonify({"error": "Anthropic API key not configured. Save it in Settings."}), 400
    if not tickets:
        return jsonify({"error": "No tickets provided."}), 400

    def _compact(t):
        return {
            "key": t.get("key"),
            "type": t.get("type"),
            "status": t.get("status"),
            "summary": str(t.get("summary") or "")[:120],
            "assignee": t.get("assignee"),
            "priority": t.get("priority"),
            "production": t.get("platform_production", False),
        }

    compact_tickets = [_compact(t) for t in tickets[:120]]  # cap to avoid token overflow
    prod_tickets    = [t for t in compact_tickets if t["production"]]
    tickets_json    = json.dumps(compact_tickets, ensure_ascii=False)
    prod_json       = json.dumps(prod_tickets, ensure_ascii=False) if prod_tickets else "[]"
    sprint_label    = sprint_name or "the selected sprint"

    overall_prompt = (
        f"You are an engineering team analyst. Below is a JSON list of Jira tickets worked on during '{sprint_label}'.\n"
        "Write a concise sprint accomplishment summary for this team — what was built, fixed, or shipped.\n"
        "Return ONLY a JSON array of bullet-point strings (plain English, no markdown, 6-10 items), e.g.:\n"
        '[\"Delivered X feature\",\"Fixed Y bug\"]\n\n'
        f"Tickets:\n{tickets_json}"
    )

    production_prompt = (
        f"You are an engineering team analyst. Below is a JSON list of production-related Jira tickets from '{sprint_label}'.\n"
        "Summarise what production issues were fixed or worked on — be specific and brief.\n"
        "Return ONLY a JSON array of bullet-point strings (plain English, no markdown, 4-8 items), e.g.:\n"
        '[\"Fixed critical payment crash on iOS\",\"Resolved DB timeout in reports API\"]\n\n'
        "If the list is empty, return: [\"No production tickets in this sprint.\"]\n\n"
        f"Production tickets:\n{prod_json}"
    )

    def _call_claude(prompt_text):
        try:
            res = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": anthropic_key,
                    "anthropic-version": "2023-06-01",
                },
                json={
                    "model": model,
                    "max_tokens": 1024,
                    "messages": [{"role": "user", "content": prompt_text}],
                },
                timeout=90,
            )
            if not res.ok:
                return None, f"Claude HTTP {res.status_code}"
            body = res.json()
            raw = (body.get("content") or [{}])[0].get("text", "").strip()
            # Strip possible markdown fences
            raw = raw.strip("`").strip()
            if raw.startswith("json"):
                raw = raw[4:].strip()
            bullets = json.loads(raw)
            if not isinstance(bullets, list):
                bullets = [str(bullets)]
            return bullets, None
        except Exception as exc:
            return None, str(exc)

    overall_bullets, err1 = _call_claude(overall_prompt)
    prod_bullets,    err2 = _call_claude(production_prompt)

    return jsonify({
        "sprint": sprint_label,
        "overall_bullets":    overall_bullets or [err1 or "Failed to generate summary."],
        "production_bullets": prod_bullets    or [err2 or "Failed to generate production summary."],
    })


@app.route("/api/reports/generate", methods=["POST"])
def generate_report_api():
    data = request.json
    team_name = data.get("team_name")
    sprint_name = data.get("sprint_name") or "N/A"
    from_date = data.get("from_date")
    to_date = data.get("to_date")
    team_members = data.get("team_members", "")
    ticket_text = data.get("tickets", "")
    report_format = data.get("format", "word")
    selected_cols = data.get("columns", ["key", "type", "summary", "status", "priority"])
    exclude_open = data.get("exclude_open", False)

    # Column Mapping for name and base widths
    col_config = {
        "key": {"name": "Key", "width": 0.8},
        "type": {"name": "Type", "width": 0.8},
        "summary": {"name": "Summary", "width": 2.0}, # Base/Min width
        "status": {"name": "Status", "width": 1.1},
        "priority": {"name": "Priority", "width": 0.8},
        "assignee": {"name": "Assignee", "width": 1.1},
        "customer": {"name": "Customer", "width": 1.2},
        "sprint": {"name": "Sprint", "width": 1.5}
    }

    # Filter config to only selected ones
    active_cols = [c for c in ["key", "type", "summary", "status", "priority", "assignee", "customer", "sprint"] if c in selected_cols]

    # Status priority for ordering
    top_statuses = [
        "DEPLOYED", "DONE", "NOT A BUG", "READY FOR STAGING", 
        "RESOLVED", "STAGED", "UNABLE TO REPRODUCE", 
        "READY FOR QA", "READY FOR INTERNAL DEMO"
    ]

    def get_status_rank(status_name):
        if not status_name: return 1
        name = status_name.upper()
        if name in top_statuses:
            return 0 # Higher priority (on top)
        return 1 # Lower priority (bottom)

    # Parse tickets
    raw_keys = re.findall(r'[A-Z]+-\d+', ticket_text.upper())
    ticket_keys = list(dict.fromkeys(raw_keys)) # Remove duplicates, preserve order
    if not ticket_keys:
        return jsonify({"error": "No valid ticket keys found"}), 400

    # Fetch Ticket Details from Jira
    bugs = []
    others = []
    
    bug_count = 0
    task_count = 0
    story_count = 0
    epic_count = 0
    other_count = 0

    headers_dict = dict(HEADERS)
    
    # Use JQL Search for efficiency and to filter by Team
    keys_str = ", ".join([f'"{k}"' for k in ticket_keys])
    # Matching user example: "team[team]" = UUID (no quotes)
    jql = f'key in ({keys_str}) AND "team[team]" = 4da67a24-33ef-42a2-b940-840dd6e450bc'
    
    try:
        # Jira Cloud v3 migration: use /search/jql endpoint
        params = {
            "jql": jql,
            "maxResults": len(ticket_keys),
            "fields": "issuetype,summary,priority,status,assignee,customfield_10077,customfield_10020"
        }
        search_res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/search/jql",
            headers=headers_dict,
            params=params
        )
        
        if search_res.status_code == 200:
            found_issues = search_res.json().get("issues", [])
            for issue in found_issues:
                key = issue.get("key")
                fields = issue.get("fields", {})
                itype = fields.get("issuetype", {}).get("name", "Task")
                summary = fields.get("summary", "No Summary")
                priority = fields.get("priority", {}).get("name", "Medium")
                status = fields.get("status", {}).get("name", "Unknown")
                assignee_field = fields.get("assignee")
                assignee = assignee_field.get("displayName", "Unassigned") if assignee_field else "Unassigned"

                # Extract Customer
                customer = "N/A"
                cust_val = fields.get("customfield_10077")
                if cust_val:
                    if isinstance(cust_val, list) and len(cust_val) > 0:
                        v = cust_val[0]
                        customer = v.get("value") or v if isinstance(v, dict) else str(v)
                    elif isinstance(cust_val, dict):
                        customer = cust_val.get("value") or str(cust_val)
                    else:
                        customer = str(cust_val)

                # Extract Sprint
                sprint = "N/A"
                sprint_val = fields.get("customfield_10020")
                if sprint_val and isinstance(sprint_val, list) and len(sprint_val) > 0:
                    # Sprints are often returned as list of objects
                    sp = sprint_val[-1] # Take the latest one
                    sprint = sp.get("name") or str(sp)

                if exclude_open and status.lower() == "open":
                    continue

                ticket_obj = {
                    "key": key,
                    "type": itype,
                    "summary": summary,
                    "priority": priority,
                    "status": status,
                    "assignee": assignee,
                    "customer": customer,
                    "sprint": sprint,
                    "rank": get_status_rank(status) 
                }

                if itype.lower() == "bug":
                    bugs.append(ticket_obj)
                    bug_count += 1
                else:
                    others.append(ticket_obj)
                    if itype.lower() == "task": task_count += 1
                    elif itype.lower() == "story": story_count += 1
                    elif itype.lower() == "epic": epic_count += 1
                    else: other_count += 1
        else:
            print(f"Jira Search Error: {search_res.status_code} - {search_res.text}")
    except Exception as e:
        print(f"Error fetching issues via JQL: {e}")

    # Sort: Status rank (0 first), then by Key
    bugs.sort(key=lambda x: (x['rank'], x['key']))
    others.sort(key=lambda x: (x['rank'], x['key']))

    # Generate Chart
    plt.figure(figsize=(6, 4))
    labels = []
    sizes = []
    colors = []
    if bug_count: labels.append('Bugs'); sizes.append(bug_count); colors.append('#ef4444')
    if task_count: labels.append('Tasks'); sizes.append(task_count); colors.append('#3b82f6')
    if story_count: labels.append('Stories'); sizes.append(story_count); colors.append('#10b981')
    if epic_count: labels.append('Epics'); sizes.append(epic_count); colors.append('#8b5cf6')
    if other_count: labels.append('Others'); sizes.append(other_count); colors.append('#6b7280')

    chart_b64 = ""
    chart_stream = io.BytesIO()
    if sizes:
        plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140, colors=colors)
        plt.axis('equal')
        plt.title("Issue Distribution")
        plt.savefig(chart_stream, format='png', bbox_inches='tight')
        chart_b64 = base64.b64encode(chart_stream.getvalue()).decode()
    else:
        # Create a small blank white pixel so the PDF img tag doesn't break
        chart_b64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+ip1sAAAAASUVORK5CYII="
    
    plt.close()

    if report_format == "word":
        doc = Document()
        
        # Header
        title = doc.add_heading(f"Executive Work Report: {team_name}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Sprint: {sprint_name} | Period: {from_date} to {to_date}")
        run.italic = True
        
        doc.add_heading("Team Members", level=1)
        doc.add_paragraph(team_members)

        if sizes:
            doc.add_heading("Issue Distribution", level=1)
            doc.add_picture(chart_stream, width=Inches(4.5))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Function to add a table to doc
        def add_styled_table(data_list, title_text, cols):
            if not data_list:
                return
            doc.add_heading(title_text, level=1)
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = 'Table Grid'
            
            # Header
            hdr_cells = table.rows[0].cells
            total_fixed_width = 0
            summary_idx = -1
            
            for i, col_id in enumerate(cols):
                hdr_cells[i].text = col_config[col_id]["name"]
                if col_id != "summary":
                    total_fixed_width += col_config[col_id]["width"]
                else:
                    summary_idx = i
            
            # Total width for standard Portrait is ~6.5 inches
            available_width = 6.4
            summary_width = max(2.0, available_width - total_fixed_width)
            
            for i, col_id in enumerate(cols):
                if col_id == "summary":
                    table.columns[i].width = Inches(summary_width)
                else:
                    table.columns[i].width = Inches(col_config[col_id]["width"])

            # Data
            for t in data_list:
                row_cells = table.add_row().cells
                for i, col_id in enumerate(cols):
                    row_cells[i].text = str(t.get(col_id, ""))

        # Tasks First
        add_styled_table(others, "Tasks, Stories & Epics", active_cols)
        # Bugs Second
        add_styled_table(bugs, "Bugs Summary", active_cols)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name=f"Work_Report_{team_name.replace(' ', '_')}.docx")

    else: # PDF
        def g_html(data_list, cols):
            if not data_list: return f"<tr><td colspan='{len(cols)}' style='text-align:center'>No items.</td></tr>"
            html_rows = ""
            for t in data_list:
                html_rows += "<tr>"
                for c in cols:
                    val = t.get(c, "")
                    extra = " class='summary-col'" if c == "summary" else ""
                    html_rows += f"<td{extra}>{val}</td>"
                html_rows += "</tr>"
            return html_rows

        others_html = g_html(others, active_cols)
        bugs_html = g_html(bugs, active_cols)
        
        table_headers = "".join([f"<th>{col_config[c]['name']}</th>" for c in active_cols])
        
        # Calculate summary width purely for CSS
        fixed_count = len(active_cols) - (1 if "summary" in active_cols else 0)
        s_width = "40%" if fixed_count <= 4 else "30%"

        team_members_html = team_members.replace("\n", "<br>")

        html_content = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Helvetica, Arial, sans-serif; color: #333; }}
                .header {{ text-align: center; border-bottom: 2px solid #3b82f6; padding-bottom: 10px; }}
                h1 {{ color: #1e3a8a; }}
                .meta {{ color: #666; font-style: italic; }}
                .section-title {{ background: #f3f4f6; padding: 5px 10px; border-left: 4px solid #3b82f6; margin-top: 20px; }}
                table {{ width: 100%; border-collapse: collapse; margin-top: 10px; table-layout: fixed; }}
                th, td {{ border: 1px solid #ddd; padding: 6px; text-align: left; font-size: 8pt; word-wrap: break-word; }}
                th {{ background-color: #f9fafb; font-weight: bold; }}
                .summary-col {{ width: {s_width}; }}
                td {{ width: auto; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>Work Report: {team_name}</h1>
                <p class="meta">Sprint: {sprint_name} | {from_date} to {to_date}</p>
            </div>
            
            <div class="section-title">Team Members</div>
            <p>{team_members_html}</p>
            
            <div class="section-title">Issue Distribution</div>
            <div style="text-align: center;">
                <img src="data:image/png;base64,{chart_b64}" width="400" />
            </div>
            
            <div class="section-title">Tasks, Stories & Epics</div>
            <table>
                <tr>{table_headers}</tr>
                {others_html}
            </table>
            
            <div class="section-title">Bugs Summary</div>
            <table>
                <tr>{table_headers}</tr>
                {bugs_html}
            </table>
        </body>
        </html>
        """
        
        output = io.BytesIO()
        pisa.CreatePDF(io.BytesIO(html_content.encode("utf-8")), dest=output)
        output.seek(0)
        return send_file(output, as_attachment=True, download_name=f"Work_Report_{team_name.replace(' ', '_')}.pdf")

@app.route("/merge_pdf")
@page_permission_required("merge_pdf")
def merge_pdf_page():
    return render_template("merge_pdf.html")

@app.route("/api/pdf/merge", methods=["POST"])
def pdf_merge_api():
    if 'files' not in request.files:
        return jsonify({"error": "No files provided"}), 400
    
    files = request.files.getlist('files')
    output_name = request.form.get('output_name', 'merged_document')
    
    if len(files) < 2:
        return jsonify({"error": "At least 2 PDF files are required to merge"}), 400
    
    try:
        merger = PdfWriter()
        for pdf in files:
            merger.append(pdf)
        
        output = io.BytesIO()
        merger.write(output)
        merger.close()
        output.seek(0)
        
        return send_file(
            output, 
            as_attachment=True, 
            download_name=f"{output_name}.pdf",
            mimetype='application/pdf'
        )
    except Exception as e:
        return jsonify({"error": f"Merge failed: {str(e)}"}), 500

@app.route("/api/audit_logs", methods=["GET"])
@login_required
def get_audit_logs():
    if current_user.role_name != 'Admin':
        return jsonify({"error": "Permission denied"}), 403
        
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 20))
    offset = (page - 1) * per_page

    conn, cursor = get_db_connection(dictionary=True)
    if not conn: return jsonify({"error": "Database error"}), 500
    
    try:
        # Get total count
        cursor.execute("SELECT COUNT(*) as total FROM audit_logs")
        total = cursor.fetchone()['total']

        # Get paginated logs
        cursor.execute("SELECT * FROM audit_logs ORDER BY created_at DESC LIMIT %s OFFSET %s", (per_page, offset))
        logs = cursor.fetchall()
        
        return jsonify({
            "logs": logs,
            "total": total,
            "page": page,
            "per_page": per_page,
            "total_pages": (total + per_page - 1) // per_page
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    finally:
        conn.close()

# =========================
# SPRINT TRACKER
# =========================

def _json_load(v, default):
    if v is None:
        return default
    if isinstance(v, (list, dict)):
        return v
    try:
        return json.loads(v)
    except Exception:
        return default

def _fetch_sprint_tracker_tree():
    conn, cursor = get_db_connection(dictionary=True)
    if not conn:
        return []
    try:
        cursor.execute("""
            SELECT id, name, sprint_goal, goal_edited, divider_index, sort_order
            FROM sprint_tracker_sprints ORDER BY sort_order ASC, id ASC
        """)
        sprints = cursor.fetchall() or []

        cursor.execute("""
            SELECT id, sprint_id, theme_key, epic_name, sentence, bullets,
                   lb_override, notes, notes_updated_by, notes_updated_at, sort_order
            FROM sprint_tracker_themes ORDER BY sort_order ASC, id ASC
        """)
        themes = cursor.fetchall() or []

        cursor.execute("""
            SELECT id, theme_id, ticket_key, summary, status, customers, lb,
                   description_bullets, last_synced_at, sort_order
            FROM sprint_tracker_tickets ORDER BY sort_order ASC, id ASC
        """)
        tickets = cursor.fetchall() or []

        tickets_by_theme = {}
        for t in tickets:
            tickets_by_theme.setdefault(t["theme_id"], []).append({
                "id": t["id"],
                "ticket_key": t["ticket_key"],
                "summary": t["summary"] or "",
                "status": t["status"] or "Open",
                "customers": _json_load(t["customers"], []),
                "lb": bool(t["lb"]),
                "description_bullets": _json_load(t["description_bullets"], []),
                "last_synced_at": t["last_synced_at"].isoformat() if t["last_synced_at"] else None,
                "sort_order": t["sort_order"],
            })

        themes_by_sprint = {}
        for th in themes:
            themes_by_sprint.setdefault(th["sprint_id"], []).append({
                "id": th["id"],
                "theme_key": th["theme_key"],
                "epic_name": th["epic_name"] or "",
                "sentence": th["sentence"] or "",
                "bullets": _json_load(th["bullets"], []),
                "lb_override": (None if th["lb_override"] is None else bool(th["lb_override"])),
                "notes": th["notes"] or "",
                "notes_updated_by": th.get("notes_updated_by") or "",
                "notes_updated_at": th["notes_updated_at"].isoformat() if th.get("notes_updated_at") else None,
                "sort_order": th["sort_order"],
                "tickets": tickets_by_theme.get(th["id"], []),
            })

        result = []
        for s in sprints:
            result.append({
                "id": s["id"],
                "name": s["name"],
                "sprint_goal": s["sprint_goal"] or "",
                "goal_edited": bool(s["goal_edited"]),
                "divider_index": s["divider_index"] or 0,
                "sort_order": s["sort_order"] or 0,
                "themes": themes_by_sprint.get(s["id"], []),
            })
        return result
    finally:
        conn.close()


@app.route("/sprint_tracker")
@login_required
@page_permission_required("sprint_tracker")
def sprint_tracker_page():
    return render_template(
        "sprint_tracker.html",
        project=PROJECT_KEY,
        can_manage_sprint_tracker=_can_manage_sprint_tracker(),
    )


@app.route("/api/sprint_tracker/data", methods=["GET"])
@login_required
@page_permission_required("sprint_tracker")
def sprint_tracker_data():
    return jsonify({"sprints": _fetch_sprint_tracker_tree()})


# ---- Sprints CRUD ----
@app.route("/api/sprint_tracker/sprints", methods=["POST"])
@login_required
@page_permission_required("sprint_tracker")
@sprint_tracker_write_required("create")
def sprint_tracker_create_sprint():
    data = request.json or {}
    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Name required"}), 400
    conn, cursor = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        cursor.execute("SELECT COALESCE(MAX(sort_order), -1) + 1 FROM sprint_tracker_sprints")
        next_order = cursor.fetchone()[0] or 0
        cursor.execute(
            """INSERT INTO sprint_tracker_sprints (name, sprint_goal, goal_edited, divider_index, sort_order)
               VALUES (%s, %s, 0, 0, %s)""",
            (name, data.get("sprint_goal", ""), next_order),
        )
        conn.commit()
        return jsonify({"id": cursor.lastrowid, "success": True})
    except Error as e:
        return jsonify({"error": str(e)}), 400
    finally:
        conn.close()


@app.route("/api/sprint_tracker/sprints/<int:sprint_id>", methods=["PUT", "DELETE"])
@login_required
@page_permission_required("sprint_tracker")
def sprint_tracker_sprint_detail(sprint_id):
    if request.method == "DELETE" and not _sprint_tracker_action_allowed("delete"):
        return jsonify({"error": "Read-only access. You can view the sprint tracker but cannot modify it."}), 403
    if request.method == "PUT" and not _sprint_tracker_action_allowed("edit"):
        return jsonify({"error": "Read-only access. You can view the sprint tracker but cannot modify it."}), 403
    conn, cursor = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        if request.method == "DELETE":
            cursor.execute("DELETE FROM sprint_tracker_sprints WHERE id = %s", (sprint_id,))
            conn.commit()
            return jsonify({"success": True})

        data = request.json or {}
        fields = []
        values = []
        for col in ("name", "sprint_goal"):
            if col in data:
                fields.append(f"{col} = %s")
                values.append(data[col])
        if "goal_edited" in data:
            fields.append("goal_edited = %s")
            values.append(1 if data["goal_edited"] else 0)
        if "divider_index" in data:
            fields.append("divider_index = %s")
            values.append(int(data["divider_index"]))
        if "sort_order" in data:
            fields.append("sort_order = %s")
            values.append(int(data["sort_order"]))
        if not fields:
            return jsonify({"success": True})
        values.append(sprint_id)
        cursor.execute(f"UPDATE sprint_tracker_sprints SET {', '.join(fields)} WHERE id = %s", tuple(values))
        conn.commit()
        return jsonify({"success": True})
    except Error as e:
        return jsonify({"error": str(e)}), 400
    finally:
        conn.close()


# ---- Themes CRUD ----
@app.route("/api/sprint_tracker/themes", methods=["POST"])
@login_required
@page_permission_required("sprint_tracker")
@sprint_tracker_write_required("create")
def sprint_tracker_create_theme():
    data = request.json or {}
    sprint_id = data.get("sprint_id")
    theme_key = (data.get("theme_key") or "").strip()
    if not sprint_id or not theme_key:
        return jsonify({"error": "sprint_id and theme_key required"}), 400
    conn, cursor = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        cursor.execute(
            "SELECT COALESCE(MAX(sort_order), -1) + 1 FROM sprint_tracker_themes WHERE sprint_id = %s",
            (sprint_id,),
        )
        next_order = cursor.fetchone()[0] or 0
        cursor.execute(
            """INSERT INTO sprint_tracker_themes
               (sprint_id, theme_key, epic_name, sentence, bullets, notes, sort_order)
               VALUES (%s, %s, %s, %s, %s, %s, %s)""",
            (
                int(sprint_id),
                theme_key,
                data.get("epic_name", ""),
                data.get("sentence", ""),
                json.dumps(data.get("bullets", [])),
                data.get("notes", ""),
                next_order,
            ),
        )
        conn.commit()
        return jsonify({"id": cursor.lastrowid, "success": True})
    except Error as e:
        return jsonify({"error": str(e)}), 400
    finally:
        conn.close()


@app.route("/api/sprint_tracker/themes/<int:theme_id>", methods=["PUT", "DELETE"])
@login_required
@page_permission_required("sprint_tracker")
def sprint_tracker_theme_detail(theme_id):
    if request.method == "DELETE" and not _sprint_tracker_action_allowed("delete"):
        return jsonify({"error": "Read-only access. You can view the sprint tracker but cannot modify it."}), 403
    if request.method == "PUT" and not _sprint_tracker_action_allowed("edit"):
        return jsonify({"error": "Read-only access. You can view the sprint tracker but cannot modify it."}), 403
    conn, cursor = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        if request.method == "DELETE":
            cursor.execute("DELETE FROM sprint_tracker_themes WHERE id = %s", (theme_id,))
            conn.commit()
            return jsonify({"success": True})

        data = request.json or {}
        fields = []
        values = []
        for col in ("theme_key", "epic_name", "sentence", "notes"):
            if col in data:
                fields.append(f"{col} = %s")
                values.append(data[col])
        notes_updated = False
        if "bullets" in data:
            fields.append("bullets = %s")
            values.append(json.dumps(data["bullets"] or []))
        if "lb_override" in data:
            lbv = data["lb_override"]
            fields.append("lb_override = %s")
            values.append(None if lbv is None else (1 if lbv else 0))
        if "notes" in data:
            notes_updated = True
            note_text = str(data.get("notes") or "").strip()
            if note_text:
                fields.append("notes_updated_by = %s")
                values.append((current_user.name or current_user.email or "Unknown").strip())
                fields.append("notes_updated_at = %s")
                values.append(datetime.utcnow())
            else:
                fields.append("notes_updated_by = NULL")
                fields.append("notes_updated_at = NULL")
        if "sort_order" in data:
            fields.append("sort_order = %s")
            values.append(int(data["sort_order"]))
        if not fields:
            return jsonify({"success": True})
        values.append(theme_id)
        cursor.execute(f"UPDATE sprint_tracker_themes SET {', '.join(fields)} WHERE id = %s", tuple(values))
        conn.commit()
        response = {"success": True}
        if notes_updated:
            note_text = str(data.get("notes") or "").strip()
            if note_text:
                response["notes_updated_by"] = (current_user.name or current_user.email or "Unknown").strip()
                response["notes_updated_at"] = datetime.utcnow().isoformat()
            else:
                response["notes_updated_by"] = ""
                response["notes_updated_at"] = None
        return jsonify(response)
    except Error as e:
        return jsonify({"error": str(e)}), 400
    finally:
        conn.close()


# ---- Tickets CRUD ----
@app.route("/api/sprint_tracker/tickets", methods=["POST"])
@login_required
@page_permission_required("sprint_tracker")
@sprint_tracker_write_required("create")
def sprint_tracker_create_ticket():
    data = request.json or {}
    theme_id = data.get("theme_id")
    ticket_key = (data.get("ticket_key") or "").strip()
    if not theme_id or not ticket_key:
        return jsonify({"error": "theme_id and ticket_key required"}), 400
    conn, cursor = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        cursor.execute(
            "SELECT COALESCE(MAX(sort_order), -1) + 1 FROM sprint_tracker_tickets WHERE theme_id = %s",
            (theme_id,),
        )
        next_order = cursor.fetchone()[0] or 0
        cursor.execute(
            """INSERT INTO sprint_tracker_tickets
               (theme_id, ticket_key, summary, status, customers, lb, sort_order)
               VALUES (%s, %s, %s, %s, %s, %s, %s)""",
            (
                int(theme_id),
                ticket_key,
                data.get("summary", ""),
                data.get("status", "Open"),
                json.dumps(data.get("customers", [])),
                1 if data.get("lb") else 0,
                next_order,
            ),
        )
        conn.commit()
        return jsonify({"id": cursor.lastrowid, "success": True})
    except Error as e:
        return jsonify({"error": str(e)}), 400
    finally:
        conn.close()


@app.route("/api/sprint_tracker/tickets/<int:ticket_id>", methods=["PUT", "DELETE"])
@login_required
@page_permission_required("sprint_tracker")
def sprint_tracker_ticket_detail(ticket_id):
    if request.method == "DELETE" and not _sprint_tracker_action_allowed("delete"):
        return jsonify({"error": "Read-only access. You can view the sprint tracker but cannot modify it."}), 403
    if request.method == "PUT" and not _sprint_tracker_action_allowed("edit"):
        return jsonify({"error": "Read-only access. You can view the sprint tracker but cannot modify it."}), 403
    conn, cursor = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        if request.method == "DELETE":
            cursor.execute("DELETE FROM sprint_tracker_tickets WHERE id = %s", (ticket_id,))
            conn.commit()
            return jsonify({"success": True})

        data = request.json or {}
        fields = []
        values = []
        for col in ("ticket_key", "summary", "status"):
            if col in data:
                fields.append(f"{col} = %s")
                values.append(data[col])
        if "customers" in data:
            fields.append("customers = %s")
            values.append(json.dumps(data["customers"] or []))
        if "description_bullets" in data:
            fields.append("description_bullets = %s")
            values.append(json.dumps(data["description_bullets"] or []))
        if "lb" in data:
            fields.append("lb = %s")
            values.append(1 if data["lb"] else 0)
        if "sort_order" in data:
            fields.append("sort_order = %s")
            values.append(int(data["sort_order"]))
        if not fields:
            return jsonify({"success": True})
        values.append(ticket_id)
        cursor.execute(f"UPDATE sprint_tracker_tickets SET {', '.join(fields)} WHERE id = %s", tuple(values))
        conn.commit()
        return jsonify({"success": True})
    except Error as e:
        return jsonify({"error": str(e)}), 400
    finally:
        conn.close()


# ---- Reorder themes + divider within a sprint ----
@app.route("/api/sprint_tracker/sprints/<int:sprint_id>/reorder", methods=["POST"])
@login_required
@page_permission_required("sprint_tracker")
@sprint_tracker_write_required("edit")
def sprint_tracker_reorder(sprint_id):
    data = request.json or {}
    theme_ids = data.get("theme_ids") or []
    divider_index = int(data.get("divider_index", len(theme_ids)))
    conn, cursor = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        for idx, tid in enumerate(theme_ids):
            cursor.execute(
                "UPDATE sprint_tracker_themes SET sort_order = %s WHERE id = %s AND sprint_id = %s",
                (idx, int(tid), sprint_id),
            )
        cursor.execute(
            "UPDATE sprint_tracker_sprints SET divider_index = %s WHERE id = %s",
            (divider_index, sprint_id),
        )
        conn.commit()
        return jsonify({"success": True})
    except Error as e:
        return jsonify({"error": str(e)}), 400
    finally:
        conn.close()


# ---- Jira + Claude sync ----
def _adf_to_text(node):
    """Flatten Atlassian Document Format to plain text."""
    if node is None:
        return ""
    if isinstance(node, str):
        return node
    if isinstance(node, list):
        return "".join(_adf_to_text(n) for n in node)
    if isinstance(node, dict):
        ntype = node.get("type")
        if ntype == "text":
            return node.get("text", "")
        if ntype == "hardBreak":
            return "\n"
        content = node.get("content") or []
        inner = _adf_to_text(content)
        if ntype in ("paragraph", "heading", "listItem"):
            return inner + "\n"
        if ntype in ("bulletList", "orderedList"):
            return inner
        return inner
    return ""


def _extract_jira_customers(cust_val):
    """Normalize Jira Customer field (customfield_10077) into a string list."""
    if not cust_val:
        return []
    values = []
    if isinstance(cust_val, list):
        for item in cust_val:
            if isinstance(item, dict):
                raw = item.get("value") or item.get("name") or item.get("displayName")
                if raw:
                    values.append(str(raw).strip())
            elif isinstance(item, str):
                s = item.strip()
                if s:
                    values.append(s)
    elif isinstance(cust_val, dict):
        raw = cust_val.get("value") or cust_val.get("name") or cust_val.get("displayName")
        if raw:
            values.append(str(raw).strip())
    elif isinstance(cust_val, str):
        s = cust_val.strip()
        if s:
            values.append(s)

    seen = set()
    result = []
    for c in values:
        if c and c not in seen:
            seen.add(c)
            result.append(c)
    return result


# Jira "platform[checkboxes]" (checkbox multi-select). Override if field discovery fails.
_PLATFORM_CHECKBOXES_FIELD_ID = None
_PLATFORM_CHECKBOXES_FIELD_DOMAIN = None


def _jira_platform_includes_production(raw):
    """True if the platform[checkboxes] value list includes the PRODUCTION option."""
    if raw is None:
        return False
    if isinstance(raw, list):
        for item in raw:
            if isinstance(item, dict):
                v = item.get("value") or item.get("name") or ""
                if str(v).strip().upper() == "PRODUCTION":
                    return True
            elif isinstance(item, str) and str(item).strip().upper() == "PRODUCTION":
                return True
    elif isinstance(raw, dict):
        v = raw.get("value") or raw.get("name") or ""
        if str(v).strip().upper() == "PRODUCTION":
            return True
    elif isinstance(raw, str) and raw.strip().upper() == "PRODUCTION":
        return True
    return False


def _get_platform_checkboxes_field_id():
    """Resolve custom field id for JQL field name platform[checkboxes]."""
    global _PLATFORM_CHECKBOXES_FIELD_ID, _PLATFORM_CHECKBOXES_FIELD_DOMAIN
    manual = (os.environ.get("JIRA_PLATFORM_CHECKBOXES_FIELD") or "").strip()
    if manual:
        return manual
    dom = str(JIRA_DOMAIN)
    if _PLATFORM_CHECKBOXES_FIELD_ID and _PLATFORM_CHECKBOXES_FIELD_DOMAIN == dom:
        return _PLATFORM_CHECKBOXES_FIELD_ID
    headers_dict = dict(HEADERS)
    if "Authorization" not in headers_dict:
        return None
    try:
        res = requests.get(
            f"{JIRA_DOMAIN}/rest/api/3/field",
            headers=headers_dict,
            timeout=20
        )
        if res.status_code != 200:
            return None
        for fld in res.json() or []:
            for cl in fld.get("clauseNames") or []:
                if str(cl).strip().lower() == "platform[checkboxes]":
                    fid = fld.get("id")
                    if fid and str(fid).startswith("customfield_"):
                        _PLATFORM_CHECKBOXES_FIELD_ID = fid
                        _PLATFORM_CHECKBOXES_FIELD_DOMAIN = dom
                        return fid
    except Exception:
        return None
    return None


NON_CUSTOMER_LABELS = {
    "launch_blocker", "reports-focusteam", "gl-dependent", "ft-req-p0", "fast_follow",
    "prevailing", "notifications", "quickwins", "payroll", "releasetarget", "canada",
    "vacation-pay", "console-ui", "off-cycle", "termination", "reporting",
    "payopsoptimization", "ft-support-p0", "bug_0401",
}

LABEL_TO_CUSTOMER = {
    "thestategroup": "The State Group",
    "bryanconstruction": "Bryan Construction",
    "stratussystems": "Stratus Systems",
    "miinc": "MIINC",
    "precisiongroup": "Precision Group",
    "wiredhq": "WiredHQ",
    "pro1electric": "Pro 1 Electric",
    "qualitywallsandceilings": "Quality Walls & Ceilings",
    "a-core-concrete": "A-Core Concrete",
    "bondedlightning": "Bonded Lightning",
    "yard1": "Yard 1",
    "wagner": "Wagner Roofing",
    "sievert": "Sievert",
    "jmconstruction": "JM Construction",
    "directtrafficcontrol": "Direct Traffic Control",
    "evanscontractinggroupllc": "Evans Contracting",
    "hdconstruction": "HD Construction",
    "evco": "EVCO",
    "k2construction": "K2 Construction",
    "m2roofing": "M2 Roofing",
    "lynco": "Lynco",
    "galindoboyd": "Galindo Boyd",
}


def _normalize_label_key(label):
    return re.sub(r"[^a-z0-9_-]+", "", str(label or "").strip().lower())


def _titlecase_label(label):
    raw = re.sub(r"[_-]+", " ", str(label or "").strip())
    raw = re.sub(r"\s+", " ", raw).strip()
    return " ".join(part.capitalize() for part in raw.split(" ")) if raw else ""


def _customers_from_labels(labels):
    out = []
    for label in (labels or []):
        original = str(label or "").strip()
        if not original:
            continue
        normalized = _normalize_label_key(original)
        if not normalized or normalized in NON_CUSTOMER_LABELS or normalized == "launch_blocker":
            continue
        mapped = LABEL_TO_CUSTOMER.get(normalized)
        if mapped:
            out.append(mapped)
        else:
            guessed = _titlecase_label(original)
            if guessed:
                out.append(guessed)
    seen = set()
    unique = []
    for c in out:
        if c not in seen:
            seen.add(c)
            unique.append(c)
    return unique


def _merge_customers(field_customers, label_customers):
    seen = set()
    merged = []
    for c in (field_customers or []) + (label_customers or []):
        cc = str(c or "").strip()
        if cc and cc not in seen:
            seen.add(cc)
            merged.append(cc)
    return merged


def _get_app_config_value(key):
    conn, cursor = get_db_connection()
    if not conn:
        return None
    try:
        cursor.execute("SELECT config_value FROM app_config WHERE config_key = %s", (key,))
        row = cursor.fetchone()
        return row[0] if row else None
    finally:
        conn.close()


def _claude_rewrite_description(description_text, anthropic_api_key, model="claude-sonnet-4-20250514"):
    """Rewrite a Jira description into max 3 concise bullets using Claude."""
    if not description_text or not anthropic_api_key:
        return []
    trimmed = description_text.strip()
    if len(trimmed) > 8000:
        trimmed = trimmed[:8000]
    prompt = (
        "Rewrite the following Jira ticket description into at most 3 concise, action-oriented "
        "bullet points that capture what's being built and why. Return ONLY a valid JSON object "
        'of the form {"bullets": ["...", "..."]}. No markdown, no prose.\n\n'
        f"Description:\n{trimmed}"
    )
    try:
        res = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": anthropic_api_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": model,
                "max_tokens": 600,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=45,
        )
        if not res.ok:
            return []
        data = res.json()
        text = "".join(b.get("text", "") for b in (data.get("content") or []) if b.get("type") == "text")
        match = re.search(r"\{[\s\S]*\}", text)
        if not match:
            return []
        parsed = json.loads(match.group(0))
        bullets = parsed.get("bullets") or []
        return [str(b).strip() for b in bullets if str(b).strip()][:3]
    except Exception as e:
        print(f"Claude rewrite error: {e}")
        return []


@app.route("/api/sprint_tracker/sprints/<int:sprint_id>/sync", methods=["POST"])
@login_required
@page_permission_required("sprint_tracker")
@sprint_tracker_write_required("edit")
def sprint_tracker_sync(sprint_id):
    # Fast sync path: bulk status/summary/customer update only (no Claude rewrite).
    conn, cursor = get_db_connection(dictionary=True)
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        cursor.execute(
            """SELECT t.id, t.ticket_key
               FROM sprint_tracker_tickets t
               JOIN sprint_tracker_themes th ON th.id = t.theme_id
               WHERE th.sprint_id = %s""",
            (sprint_id,),
        )
        rows = cursor.fetchall() or []
    finally:
        conn.close()

    if not rows:
        return jsonify({"synced": 0, "results": []})

    jira_headers = dict(HEADERS)
    if "Authorization" not in jira_headers:
        return jsonify({"error": "Jira credentials are not configured. Please set them in Settings."}), 400
    jira_domain = str(JIRA_DOMAIN)

    results = []
    status_map = {}
    now_ts = datetime.utcnow()
    keys = [str(r["ticket_key"]).strip().upper() for r in rows if r.get("ticket_key")]
    jira_by_key = {}
    for i in range(0, len(keys), 80):
        chunk = keys[i:i+80]
        quoted = ", ".join([f'"{_escape_jql_value(k)}"' for k in chunk])
        jql = f"key IN ({quoted})"
        r = requests.get(
            f"{jira_domain}/rest/api/3/search/jql",
            headers=jira_headers,
            params={"jql": jql, "maxResults": 100, "fields": "status,summary,customfield_10077,labels"},
            timeout=35,
        )
        if not r.ok:
            continue
        for issue in (r.json().get("issues") or []):
            key = (issue.get("key") or "").strip().upper()
            if key:
                jira_by_key[key] = issue

    conn, cursor = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        for row in rows:
            tid = row["id"]
            key = str(row["ticket_key"]).strip().upper()
            try:
                issue = jira_by_key.get(key)
                if not issue:
                    results.append({"key": key, "ok": False, "error": "Missing in Jira search response"})
                    continue
                fields = issue.get("fields") or {}
                status_name = ((fields.get("status") or {}).get("name")) or ""
                summary = fields.get("summary") or ""
                customers = _extract_jira_customers(fields.get("customfield_10077"))
                labels = fields.get("labels") or []
                lb = 1 if _is_launch_blocker(labels) else 0
                updates = ["status = %s", "summary = %s", "customers = %s", "lb = %s", "last_synced_at = %s"]
                vals = [status_name or "Open", summary, json.dumps(customers), lb, now_ts]
                vals.append(tid)
                cursor.execute(
                    f"UPDATE sprint_tracker_tickets SET {', '.join(updates)} WHERE id = %s",
                    tuple(vals),
                )
                status_map[key] = status_name
                results.append({
                    "key": key,
                    "ok": True,
                    "status": status_name,
                    "customers": customers,
                })
            except Exception as e:
                results.append({"key": key, "ok": False, "error": str(e)})
        conn.commit()
    finally:
        conn.close()

    return jsonify({
        "synced": sum(1 for r in results if r.get("ok")),
        "total": len(results),
        "results": results,
        "status_map": status_map,
    })


@app.route("/api/sprint_tracker/sprints/<int:sprint_id>/sync_jql", methods=["POST"])
@login_required
@page_permission_required("sprint_tracker")
@sprint_tracker_write_required("edit")
def sprint_tracker_sync_jql(sprint_id):
    payload = request.json or {}
    jql = (payload.get("jql") or "").strip()
    if not jql:
        return jsonify({"error": "JQL is required"}), 400

    jira_headers = dict(HEADERS)
    if "Authorization" not in jira_headers:
        return jsonify({"error": "Jira credentials are not configured. Please set them in Settings."}), 400
    jira_domain = str(JIRA_DOMAIN)

    conn, cursor = get_db_connection(dictionary=True)
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        cursor.execute(
            "SELECT id, theme_key FROM sprint_tracker_themes WHERE sprint_id = %s ORDER BY sort_order ASC, id ASC",
            (sprint_id,),
        )
        themes = cursor.fetchall() or []
        theme_ids = [t["id"] for t in themes]
        if not theme_ids:
            return jsonify({"error": "No themes found in this sprint. Create at least one row first."}), 400

        sync_theme = next((t for t in themes if (t.get("theme_key") or "").strip().lower() == "jql_sync"), None)
        if not sync_theme:
            cursor.execute("SELECT COALESCE(MAX(sort_order), -1) + 1 AS next_order FROM sprint_tracker_themes WHERE sprint_id = %s", (sprint_id,))
            next_order = (cursor.fetchone() or {}).get("next_order", 0) or 0
            cursor.execute(
                """INSERT INTO sprint_tracker_themes
                   (sprint_id, theme_key, epic_name, sentence, bullets, notes, sort_order)
                   VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                (sprint_id, "jql_sync", "JQL Sync", "Tickets synced from JQL", json.dumps([]), "", int(next_order)),
            )
            sync_theme_id = cursor.lastrowid
            conn.commit()
        else:
            sync_theme_id = sync_theme["id"]

        placeholders = ", ".join(["%s"] * len(theme_ids))
        cursor.execute(
            f"""SELECT t.id, t.theme_id, t.ticket_key
                FROM sprint_tracker_tickets t
                WHERE t.theme_id IN ({placeholders})""",
            tuple(theme_ids),
        )
        existing_rows = cursor.fetchall() or []
        existing_by_key = {str(r["ticket_key"]).strip().upper(): r for r in existing_rows if r.get("ticket_key")}

        all_issues = []
        seen_keys = set()
        start_at = 0
        next_page_token = None
        page_safety = 0
        while page_safety < 80:
            page_safety += 1
            params = {
                "jql": jql,
                "maxResults": 100,
                "fields": "summary,status,customfield_10077,labels"
            }
            if next_page_token:
                params["nextPageToken"] = next_page_token
            else:
                params["startAt"] = start_at
            r = requests.get(
                f"{jira_domain}/rest/api/3/search/jql",
                headers=jira_headers,
                params=params,
                timeout=45,
            )
            if not r.ok:
                return jsonify({"error": f"Jira API error: {r.text}"}), r.status_code
            data = r.json()
            issues = data.get("issues", [])
            if not issues:
                break
            new_count = 0
            for issue in issues:
                key = (issue.get("key") or "").strip().upper()
                if not key or key in seen_keys:
                    continue
                seen_keys.add(key)
                all_issues.append(issue)
                new_count += 1
            if new_count == 0:
                break
            next_page_token = data.get("nextPageToken")
            if next_page_token:
                continue
            if len(issues) < 100:
                break
            start_at += 100

        now_ts = datetime.utcnow()
        updated = 0
        inserted = 0
        errors = []

        cursor.execute("SELECT COALESCE(MAX(sort_order), -1) AS max_order FROM sprint_tracker_tickets WHERE theme_id = %s", (sync_theme_id,))
        max_order = (cursor.fetchone() or {}).get("max_order", -1)
        next_ticket_order = int(max_order) + 1

        active_keys = set()
        for issue in all_issues:
            try:
                key = (issue.get("key") or "").strip().upper()
                if not key:
                    continue
                active_keys.add(key)
                f = issue.get("fields") or {}
                status_name = ((f.get("status") or {}).get("name")) or "Open"
                existing = existing_by_key.get(key)
                if existing:
                    cursor.execute(
                        """UPDATE sprint_tracker_tickets
                           SET status = %s, last_synced_at = %s
                           WHERE id = %s""",
                        (status_name, now_ts, existing["id"]),
                    )
                    updated += 1
                else:
                    summary = f.get("summary") or ""
                    customers = _extract_jira_customers(f.get("customfield_10077"))
                    labels = f.get("labels") or []
                    lb = 1 if _is_launch_blocker(labels) else 0
                    cursor.execute(
                        """INSERT INTO sprint_tracker_tickets
                           (theme_id, ticket_key, summary, status, customers, lb, description_bullets, last_synced_at, sort_order)
                           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                        (sync_theme_id, key, summary, status_name, json.dumps(customers), lb, json.dumps([]), now_ts, next_ticket_order),
                    )
                    next_ticket_order += 1
                    inserted += 1
            except Exception as e:
                errors.append(str(e))

        # Remove only tickets from JQL sync bucket that are no longer in JQL scope.
        # Keep manually curated row tickets intact.
        removed = 0
        for r in existing_rows:
            key = str(r.get("ticket_key") or "").strip().upper()
            theme_id = r.get("theme_id")
            if theme_id == sync_theme_id and key and key not in active_keys:
                cursor.execute("DELETE FROM sprint_tracker_tickets WHERE id = %s", (r["id"],))
                removed += 1

        conn.commit()
        return jsonify({
            "success": True,
            "total_fetched": len(all_issues),
            "updated": updated,
            "inserted": inserted,
            "removed": removed,
            "errors": errors[:10],
            "sync_theme_id": sync_theme_id,
        })
    finally:
        conn.close()


@app.route("/api/sprint_tracker/tickets/<int:ticket_id>/rewrite", methods=["POST"])
@login_required
@page_permission_required("sprint_tracker")
@sprint_tracker_write_required("edit")
def sprint_tracker_ticket_rewrite(ticket_id):
    payload = request.json or {}
    anthropic_key = (payload.get("anthropic_api_key") or "").strip() or _get_app_config_value("anthropic_api_key")
    if not anthropic_key:
        return jsonify({"error": "Anthropic API key missing. Save it in Settings (key: anthropic_api_key)."}), 400

    conn, cursor = get_db_connection(dictionary=True)
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        cursor.execute("SELECT id, ticket_key FROM sprint_tracker_tickets WHERE id = %s", (ticket_id,))
        row = cursor.fetchone()
    finally:
        conn.close()
    if not row:
        return jsonify({"error": "Ticket not found"}), 404

    jira_headers = dict(HEADERS)
    if "Authorization" not in jira_headers:
        return jsonify({"error": "Jira credentials missing"}), 400
    jira_domain = str(JIRA_DOMAIN)

    try:
        r = requests.get(
            f"{jira_domain}/rest/api/3/issue/{row['ticket_key']}",
            headers=jira_headers,
            params={"fields": "description,summary,status,customfield_10077"},
            timeout=30,
        )
        if not r.ok:
            return jsonify({"error": f"Jira {r.status_code}"}), 400
        issue = r.json()
        fields = issue.get("fields") or {}
        desc_text = _adf_to_text(fields.get("description"))
        customers = _extract_jira_customers(fields.get("customfield_10077"))
        bullets = _claude_rewrite_description(desc_text, anthropic_key) if desc_text.strip() else []

        conn, cursor = get_db_connection()
        try:
            cursor.execute(
                "UPDATE sprint_tracker_tickets SET description_bullets = %s, customers = %s, last_synced_at = %s WHERE id = %s",
                (json.dumps(bullets), json.dumps(customers), datetime.utcnow(), ticket_id),
            )
            conn.commit()
        finally:
            conn.close()
        return jsonify({"success": True, "description_bullets": bullets})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/sprint_tracker/themes/generate_from_tickets", methods=["POST"])
@login_required
@page_permission_required("sprint_tracker")
@sprint_tracker_write_required("create")
def sprint_tracker_generate_theme_from_tickets():
    payload = request.json or {}
    ticket_keys = payload.get("ticket_keys") or []
    anthropic_key = (payload.get("anthropic_api_key") or "").strip() or _get_app_config_value("anthropic_api_key")
    model = payload.get("model") or "claude-sonnet-4-20250514"

    if not isinstance(ticket_keys, list) or not ticket_keys:
        return jsonify({"error": "ticket_keys is required"}), 400

    normalized_keys = []
    seen = set()
    for k in ticket_keys:
        kk = str(k).strip().upper()
        if kk and kk not in seen:
            seen.add(kk)
            normalized_keys.append(kk)
    if not normalized_keys:
        return jsonify({"error": "No valid ticket keys"}), 400

    if not anthropic_key:
        return jsonify({"error": "Anthropic API key missing. Save it in Settings (key: anthropic_api_key)."}), 400

    jira_headers = dict(HEADERS)
    if "Authorization" not in jira_headers:
        return jsonify({"error": "Jira credentials are not configured. Please set them in Settings."}), 400
    jira_domain = str(JIRA_DOMAIN)

    fetched_tickets = []
    for key in normalized_keys:
        try:
            r = requests.get(
                f"{jira_domain}/rest/api/3/issue/{key}",
                headers=jira_headers,
                params={"fields": "summary,status,description,customfield_10077"},
                timeout=30,
            )
            if not r.ok:
                fetched_tickets.append({
                    "ticket_key": key,
                    "summary": "",
                    "status": "Open",
                    "description_text": "",
                    "error": f"Jira {r.status_code}",
                })
                continue
            issue = r.json()
            fields = issue.get("fields") or {}
            fetched_tickets.append({
                "ticket_key": key,
                "summary": fields.get("summary") or "",
                "status": ((fields.get("status") or {}).get("name")) or "Open",
                "description_text": _adf_to_text(fields.get("description")).strip(),
                "customers": _extract_jira_customers(fields.get("customfield_10077")),
            })
        except Exception as e:
            fetched_tickets.append({
                "ticket_key": key,
                "summary": "",
                "status": "Open",
                "description_text": "",
                "customers": [],
                "error": str(e),
            })

    llm_input = []
    for t in fetched_tickets:
        llm_input.append({
            "key": t["ticket_key"],
            "summary": t.get("summary") or "",
            "status": t.get("status") or "Open",
            "description": (t.get("description_text") or "")[:2000],
        })

    prompt = (
        "You are generating a sprint theme from Jira tickets.\n"
        "Given ticket summaries and descriptions, infer one coherent theme.\n"
        "Return ONLY valid JSON with this exact shape:\n"
        "{\n"
        "  \"theme_key\": \"short-kebab-case-key\",\n"
        "  \"epic_name\": \"Readable theme title\",\n"
        "  \"sentence\": \"One line 'What we are building'\",\n"
        "  \"bullets\": [\"bullet 1\", \"bullet 2\", \"bullet 3\"]\n"
        "}\n"
        "Rules:\n"
        "- bullets max 3\n"
        "- concise, action-oriented\n"
        "- do not include markdown\n\n"
        f"Tickets:\n{json.dumps(llm_input, ensure_ascii=False)}"
    )

    suggestion = {
        "theme_key": "generated-theme",
        "epic_name": "Generated Theme",
        "sentence": "Theme generated from selected tickets",
        "bullets": [],
    }
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": anthropic_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": model,
                "max_tokens": 700,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=60,
        )
        if resp.ok:
            data = resp.json()
            text = "".join(b.get("text", "") for b in (data.get("content") or []) if b.get("type") == "text")
            m = re.search(r"\{[\s\S]*\}", text)
            if m:
                parsed = json.loads(m.group(0))
                suggestion = {
                    "theme_key": str(parsed.get("theme_key") or suggestion["theme_key"]).strip(),
                    "epic_name": str(parsed.get("epic_name") or suggestion["epic_name"]).strip(),
                    "sentence": str(parsed.get("sentence") or suggestion["sentence"]).strip(),
                    "bullets": [str(b).strip() for b in (parsed.get("bullets") or []) if str(b).strip()][:3],
                }
    except Exception as e:
        print(f"Theme generation error: {e}")

    tickets_payload = [{
        "ticket_key": t["ticket_key"],
        "summary": t.get("summary") or "",
        "status": t.get("status") or "Open",
        "customers": t.get("customers") or [],
        "lb": False,
    } for t in fetched_tickets]

    return jsonify({
        "success": True,
        "suggestion": suggestion,
        "tickets": tickets_payload,
        "fetched": len(tickets_payload),
    })


def _fetch_jira_issues_from_jql(jql_query):
    """Paginated Jira search for sprint-tracker previews (handles nextPageToken + startAt)."""
    jira_headers = dict(HEADERS)
    if "Authorization" not in jira_headers:
        raise ValueError("Jira credentials are not configured. Please set them in Settings.")
    jira_domain = str(JIRA_DOMAIN).rstrip("/")
    fields = "summary,status,description,labels,customfield_10077,customfield_10014,issuetype,parent"
    all_issues = []
    seen_keys = set()
    start_at = 0
    next_page_token = None
    search_url = f"{jira_domain}/rest/api/3/search/jql"

    for _page in range(120):
        params = {"jql": jql_query, "maxResults": 100, "fields": fields}
        if next_page_token:
            params["nextPageToken"] = next_page_token
        else:
            params["startAt"] = start_at

        resp = requests.get(search_url, headers=jira_headers, params=params, timeout=45)
        if resp.status_code in (404, 405):
            resp = requests.get(
                f"{jira_domain}/rest/api/3/search",
                headers=jira_headers,
                params=params,
                timeout=45,
            )
        if not resp.ok:
            raise ValueError(f"Jira search failed: {resp.status_code} {resp.text[:280]}")

        try:
            data = resp.json()
        except ValueError as e:
            raise ValueError(f"Invalid JSON from Jira search: {e}") from e

        issues = data.get("issues") or []
        if not issues:
            break

        new_count = 0
        for issue in issues:
            ik = issue.get("key")
            if ik and ik not in seen_keys:
                seen_keys.add(ik)
                all_issues.append(issue)
                new_count += 1

        if new_count == 0:
            break

        next_page_token = data.get("nextPageToken")
        if next_page_token:
            continue

        if len(issues) < 100:
            break

        total = data.get("total")
        start_at += len(issues)
        if isinstance(total, int) and start_at >= total:
            break
        if total is None and len(issues) < 100:
            break

    return all_issues


def _normalize_jira_epic_link(raw):
    if raw is None:
        return ""
    if isinstance(raw, dict):
        return str(raw.get("key") or raw.get("id") or "").strip()
    return str(raw).strip()


def _resolve_epic_summary(issue, jira_domain, jira_headers, epic_cache):
    fields = issue.get("fields") or {}
    issue_type = ((fields.get("issuetype") or {}).get("name") or "").strip().lower()
    if issue_type == "epic":
        return fields.get("summary") or issue.get("key") or "_no_epic"

    epic_key = _normalize_jira_epic_link(fields.get("customfield_10014"))
    if not epic_key and isinstance(fields.get("parent"), dict):
        parent = fields.get("parent") or {}
        p_fields = parent.get("fields") or {}
        p_type = ((p_fields.get("issuetype") or {}).get("name") or "").strip().lower()
        if p_type == "epic":
            return p_fields.get("summary") or parent.get("key") or "_no_epic"
        epic_key = _normalize_jira_epic_link(parent.get("key"))

    if not epic_key:
        return "_no_epic"
    epic_key = epic_key.strip()
    if not epic_key:
        return "_no_epic"
    if epic_key in epic_cache:
        return epic_cache[epic_key]
    try:
        r = requests.get(
            f"{jira_domain}/rest/api/3/issue/{epic_key}",
            headers=jira_headers,
            params={"fields": "summary"},
            timeout=30,
        )
        if r.ok:
            summary = (r.json().get("fields") or {}).get("summary") or epic_key
        else:
            summary = epic_key
    except Exception:
        summary = epic_key
    epic_cache[epic_key] = summary
    return summary


def _sprint_preview_anthropic_http_timeout():
    """(connect, read) timeouts for Claude during /api/sprint_tracker/jql/preview.

    Sync Gunicorn workers often use a 30s worker timeout. If this HTTP call blocks
    longer than that, the master aborts the worker (WORKER TIMEOUT). Default read
    timeout stays safely below a 30s worker unless you set env vars.

    - SPRINT_PREVIEW_ANTHROPIC_CONNECT_TIMEOUT (default 15)
    - SPRINT_PREVIEW_ANTHROPIC_READ_TIMEOUT: if unset, min(25, worker_limit - 4)
      where worker_limit comes from GUNICORN_WORKER_TIMEOUT or GUNICORN_TIMEOUT
      (default 30). If set explicitly, used as-is (must stay below real Gunicorn
      --timeout or workers will still be killed).
    """
    try:
        connect = float(os.environ.get("SPRINT_PREVIEW_ANTHROPIC_CONNECT_TIMEOUT", "15"))
    except ValueError:
        connect = 15.0
    explicit_read = os.environ.get("SPRINT_PREVIEW_ANTHROPIC_READ_TIMEOUT")
    if explicit_read is not None and str(explicit_read).strip() != "":
        try:
            read = float(explicit_read)
        except ValueError:
            read = 25.0
    else:
        try:
            worker_limit = float(
                os.environ.get("GUNICORN_WORKER_TIMEOUT")
                or os.environ.get("GUNICORN_TIMEOUT")
                or "30"
            )
        except ValueError:
            worker_limit = 30.0
        safety = 4.0
        read = min(25.0, max(8.0, worker_limit - safety))
    return (connect, read)


def _generate_themes_with_claude(grouped, anthropic_key):
    """Returns (theme_by_group_dict, optional_user_message).

    On Anthropic timeout or failure, dict may be {} and themes fall back to epic/summary text.
    """
    try:
        groups_blob = json.dumps(grouped, ensure_ascii=False)
    except (TypeError, ValueError) as ser_exc:
        print(f"Sprint tracker Claude prompt serialization error: {ser_exc}")
        return {}, None
    prompt = (
        "You are preparing sprint tracker themes from Jira ticket groups.\n"
        "For EACH group return:\n"
        "- theme_key: short kebab-case\n"
        "- epic_name: readable title\n"
        "- sentence: under 12 words, outcome-oriented\n"
        "- bullets: exactly 2 bullets, each under 12 words (if only 1 ticket, 1 bullet allowed)\n"
        "Return ONLY valid JSON:\n"
        "{ \"themes\": [ { \"group_key\": \"...\", \"theme_key\": \"...\", \"epic_name\": \"...\", \"sentence\": \"...\", \"bullets\": [\"...\"] } ] }\n"
        f"Groups:\n{groups_blob}"
    )
    timeout = _sprint_preview_anthropic_http_timeout()
    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": anthropic_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 1800,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=timeout,
        )
        if not resp.ok:
            print(f"Claude theme generation HTTP {resp.status_code}: {(resp.text or '')[:500]}")
            return {}, None
        data = resp.json()
        text = "".join(b.get("text", "") for b in (data.get("content") or []) if b.get("type") == "text")
        m = re.search(r"\{[\s\S]*\}", text)
        if not m:
            return {}, None
        parsed = json.loads(m.group(0))
        items = parsed.get("themes") or []
        by_group = {}
        for t in items:
            gk = str(t.get("group_key") or "").strip()
            if not gk:
                continue
            bullets = [str(b).strip() for b in (t.get("bullets") or []) if str(b).strip()]
            by_group[gk] = {
                "theme_key": str(t.get("theme_key") or "").strip(),
                "epic_name": str(t.get("epic_name") or "").strip(),
                "sentence": str(t.get("sentence") or "").strip(),
                "bullets": bullets[:3],
            }
        return by_group, None
    except requests.Timeout:
        print(
            "Claude theme generation timed out (Anthropic HTTP). "
            f"timeout={timeout!r} — raise Gunicorn --timeout or SPRINT_PREVIEW_ANTHROPIC_READ_TIMEOUT if needed."
        )
        return {}, (
            "Theme generation timed out talking to Anthropic; using automatic theme text. "
            "Increase Gunicorn worker timeout (e.g. --timeout 120) and optionally "
            "SPRINT_PREVIEW_ANTHROPIC_READ_TIMEOUT if Claude is often slow."
        )
    except requests.RequestException as e:
        print(f"Claude theme generation HTTP error: {e}")
        return {}, None
    except Exception as e:
        print(f"Claude theme generation error: {e}")
        return {}, None


@app.route("/api/sprint_tracker/jql/preview", methods=["POST"])
@login_required
@page_permission_required("sprint_tracker")
@sprint_tracker_write_required("create")
def sprint_tracker_jql_preview():
    payload = request.get_json(silent=True) or {}
    jql = (payload.get("jql") or "").strip()
    if not jql:
        return jsonify({"error": "JQL is required"}), 400
    anthropic_key = (payload.get("anthropic_api_key") or "").strip() or _get_app_config_value("anthropic_api_key")
    if not anthropic_key:
        return jsonify({"error": "Anthropic API key missing. Save it in settings first."}), 400

    jira_headers = dict(HEADERS)
    if "Authorization" not in jira_headers:
        return jsonify({"error": "Jira credentials are not configured. Please set them in Settings."}), 400
    jira_domain = str(JIRA_DOMAIN)

    try:
        try:
            issues = _fetch_jira_issues_from_jql(jql)
        except ValueError as ve:
            return jsonify({"error": str(ve)}), 400
        if not issues:
            return jsonify({"error": "No issues found for the provided JQL"}), 400

        epic_cache = {}
        grouped = {}
        for issue in issues:
            try:
                fields = issue.get("fields") or {}
                labels = [str(l).strip() for l in (fields.get("labels") or []) if str(l).strip()]
                lb = any(str(l).strip().lower() == "launch_blocker" for l in labels)
                field_customers = _extract_jira_customers(fields.get("customfield_10077"))
                label_customers = _customers_from_labels(labels)
                customers = _merge_customers(field_customers, label_customers)
                epic_summary = _resolve_epic_summary(issue, jira_domain, jira_headers, epic_cache)
                group_key = epic_summary if epic_summary else "_no_epic"
                grouped.setdefault(group_key, []).append({
                    "ticket_key": issue.get("key"),
                    "summary": fields.get("summary") or "",
                    "status": ((fields.get("status") or {}).get("name")) or "Open",
                    "labels": labels,
                    "description": _adf_to_text(fields.get("description")).strip()[:2500],
                    "customers": customers,
                    "lb": lb,
                })
            except Exception as row_exc:
                key = issue.get("key") or "?"
                return jsonify({"error": f"Failed processing issue {key}: {row_exc}"}), 400

        claude_groups = []
        for group_key, tickets in grouped.items():
            claude_groups.append({
                "group_key": group_key,
                "epic_summary": group_key,
                "ticket_count": len(tickets),
                "tickets": [{"key": t["ticket_key"], "summary": t["summary"], "description": t["description"]} for t in tickets],
            })

        generated, claude_note = _generate_themes_with_claude(claude_groups, anthropic_key)

        themes = []
        for idx, (group_key, tickets) in enumerate(grouped.items()):
            g = generated.get(group_key) or {}
            fallback_key = re.sub(r"[^a-z0-9]+", "-", group_key.lower()).strip("-") or f"theme-{idx+1}"
            fallback_sentence = (group_key if group_key != "_no_epic" else "General improvements").strip()
            bullet_seed = [t["summary"] for t in tickets if t.get("summary")]
            fallback_bullets = bullet_seed[:2] if len(tickets) > 1 else bullet_seed[:1]
            themes.append({
                "theme_key": g.get("theme_key") or fallback_key,
                "epic_name": g.get("epic_name") or (group_key if group_key != "_no_epic" else "No Epic"),
                "sentence": g.get("sentence") or fallback_sentence,
                "bullets": (g.get("bullets") or fallback_bullets)[:3],
                "tickets": [{
                    "ticket_key": t["ticket_key"],
                    "summary": t["summary"],
                    "status": t["status"],
                    "customers": t["customers"],
                    "lb": t["lb"],
                } for t in tickets],
            })

        top = sorted(themes, key=lambda th: len(th.get("tickets") or []), reverse=True)[:3]
        top_sentences = [str(t.get("sentence") or "").strip().lower().rstrip(".") for t in top if str(t.get("sentence") or "").strip()]
        if not top_sentences:
            inferred_goal = "Ship sprint commitments from selected Jira scope"
        elif len(top_sentences) == 1:
            inferred_goal = f"Ship {top_sentences[0]}"
        elif len(top_sentences) == 2:
            inferred_goal = f"Ship {top_sentences[0]} and {top_sentences[1]}"
        else:
            inferred_goal = f"Ship {top_sentences[0]}, {top_sentences[1]}, and {top_sentences[2]}"

        body = {
            "success": True,
            "jql": jql,
            "ticket_count": len(issues),
            "theme_count": len(themes),
            "sprint_goal": inferred_goal,
            "themes": themes,
        }
        if claude_note:
            body["claude_note"] = claude_note
        return jsonify(body)
    except Exception as exc:
        import traceback
        traceback.print_exc()
        return jsonify({"error": "Sprint preview failed unexpectedly.", "detail": str(exc)}), 500


@app.route("/api/sprint_tracker/sprints/from_generated", methods=["POST"])
@login_required
@page_permission_required("sprint_tracker")
@sprint_tracker_write_required("create")
def sprint_tracker_create_sprint_from_generated():
    payload = request.json or {}
    sprint_name = (payload.get("sprint_name") or "").strip()
    sprint_goal = (payload.get("sprint_goal") or "").strip()
    themes = payload.get("themes") or []
    if not sprint_name:
        return jsonify({"error": "sprint_name is required"}), 400
    if not isinstance(themes, list) or not themes:
        return jsonify({"error": "themes is required"}), 400

    conn, cursor = get_db_connection()
    if not conn:
        return jsonify({"error": "Database error"}), 500
    try:
        cursor.execute("SELECT COALESCE(MAX(sort_order), -1) + 1 FROM sprint_tracker_sprints")
        sprint_order = cursor.fetchone()[0] or 0
        cursor.execute(
            """INSERT INTO sprint_tracker_sprints (name, sprint_goal, goal_edited, divider_index, sort_order)
               VALUES (%s, %s, 1, %s, %s)""",
            (sprint_name, sprint_goal, len(themes), sprint_order),
        )
        sprint_id = cursor.lastrowid

        used_theme_keys = set()
        for t_idx, theme in enumerate(themes):
            raw_key = str(theme.get("theme_key") or "").strip() or f"theme-{t_idx+1}"
            key = raw_key
            suffix = 2
            while key in used_theme_keys:
                key = f"{raw_key}-{suffix}"
                suffix += 1
            used_theme_keys.add(key)

            cursor.execute(
                """INSERT INTO sprint_tracker_themes
                   (sprint_id, theme_key, epic_name, sentence, bullets, notes, sort_order)
                   VALUES (%s, %s, %s, %s, %s, '', %s)""",
                (
                    sprint_id,
                    key,
                    str(theme.get("epic_name") or "").strip(),
                    str(theme.get("sentence") or "").strip(),
                    json.dumps(theme.get("bullets") or []),
                    t_idx,
                ),
            )
            theme_id = cursor.lastrowid
            tickets = theme.get("tickets") or []
            for k_idx, tk in enumerate(tickets):
                cursor.execute(
                    """INSERT INTO sprint_tracker_tickets
                       (theme_id, ticket_key, summary, status, customers, lb, sort_order)
                       VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                    (
                        theme_id,
                        str(tk.get("ticket_key") or "").strip().upper(),
                        str(tk.get("summary") or "").strip(),
                        str(tk.get("status") or "Open").strip(),
                        json.dumps(tk.get("customers") or []),
                        1 if tk.get("lb") else 0,
                        k_idx,
                    ),
                )
        conn.commit()
        return jsonify({"success": True, "sprint_id": sprint_id})
    except Exception as e:
        conn.rollback()
        return jsonify({"error": str(e)}), 400
    finally:
        conn.close()



# ---------------------------------------------------------------------------
# Team Diagram Report
# ---------------------------------------------------------------------------

def _team_diagram_coerce_story_points(raw):
    """Numeric story points from Jira custom field (often customfield_10016); missing → 0."""
    if raw is None:
        return 0.0
    if isinstance(raw, bool):
        return 0.0
    if isinstance(raw, (int, float)):
        return float(raw)
    if isinstance(raw, str) and raw.strip():
        try:
            return float(raw.strip())
        except ValueError:
            return 0.0
    return 0.0


def _team_diagram_normalize_bullet_line(s):
    """Strip leading list markers and orphaned ':' left when UI strips PROJ-123 from 'PROJ-123: …'."""
    t = (s or "").strip()
    if not t:
        return ""
    for _ in range(6):
        prev = t
        t = re.sub(r"^[\s•·\u2022]+", "", t)
        t = re.sub(r"^[\-\u2013\u2014]\s*", "", t)
        t = re.sub(r"^\d+[.)]\s*", "", t)
        t = re.sub(r"^:\s*", "", t)
        t = t.strip()
        if t == prev:
            break
    return t


def _team_diagram_ai_cache_key(jql, sub_team_name):
    raw = f"{(jql or '').strip()}\n{(sub_team_name or '').strip()}\nai_pack_v4_prod_story_ticket_refs"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _team_diagram_ai_normalize_client_blob(d):
    """Ensure bullet arrays exist; map legacy prose fields if needed."""
    if not isinstance(d, dict):
        return {}

    def _clean_list(key_new, key_legacy):
        v = d.get(key_new)
        if isinstance(v, list):
            out = [_team_diagram_normalize_bullet_line(str(x).strip()) for x in v if x is not None and str(x).strip()]
            out = [x for x in out if x][:14]
            if out:
                return out
        s = (d.get(key_legacy) or "").strip()
        if not s:
            return []
        lines = [_team_diagram_normalize_bullet_line(ln.strip().lstrip("-•*\t ")) for ln in re.split(r"[\n\r]+", s) if ln.strip()]
        lines = [x for x in lines if x][:14]
        if len(lines) > 1:
            return lines
        one = _team_diagram_normalize_bullet_line(s[:1200])
        return [one] if one else []

    m = dict(d)
    m["stories_bullets"] = _clean_list("stories_bullets", "stories_summary")
    m["prod_bullets"] = _clean_list("prod_bullets", "prod_issues_summary")
    m["overall_bullets"] = _clean_list("overall_bullets", "overall_summary")
    return m


def _team_diagram_ai_cache_get(cache_key):
    conn, cursor = get_db_connection(dictionary=True)
    if not conn:
        return None
    try:
        cursor.execute(
            "SELECT response_json FROM team_diagram_ai_cache WHERE cache_key = %s",
            (cache_key,),
        )
        row = cursor.fetchone()
        if row and row.get("response_json"):
            return json.loads(row["response_json"])
    except Exception as e:
        print(f"team_diagram_ai_cache_get: {e}")
    finally:
        conn.close()
    return None


def _team_diagram_ai_pack_for_store(blob):
    """Persist only stable summary fields (no runtime flags)."""
    if not isinstance(blob, dict):
        return {}
    allow = (
        "stories_bullets",
        "prod_bullets",
        "overall_bullets",
        "story_keys_highlighted",
        "prod_keys_highlighted",
    )
    return {k: blob[k] for k in allow if k in blob}


def _team_diagram_ai_cache_put(cache_key, sub_team_name, data_obj):
    conn, cursor = get_db_connection()
    if not conn:
        return False
    try:
        blob = json.dumps(_team_diagram_ai_pack_for_store(data_obj), ensure_ascii=False)
        cursor.execute(
            """INSERT INTO team_diagram_ai_cache (cache_key, sub_team_name, response_json)
               VALUES (%s, %s, %s)
               ON DUPLICATE KEY UPDATE sub_team_name = VALUES(sub_team_name),
                 response_json = VALUES(response_json)""",
            (cache_key, (sub_team_name or "")[:512], blob),
        )
        conn.commit()
        return True
    except Exception as e:
        print(f"team_diagram_ai_cache_put: {e}")
        conn.rollback()
    finally:
        conn.close()
    return False


def _team_diagram_claude_summarize(sub_team_name, main_team, story_issues, prod_issues, bug_issues, anthropic_key, model):
    sj = json.dumps(story_issues[:55], ensure_ascii=False)
    pj = json.dumps(prod_issues[:55], ensure_ascii=False)
    bj = json.dumps(bug_issues[:65], ensure_ascii=False)
    prompt = (
        f'You summarize engineering work for sub-team "{sub_team_name}" under main team "{main_team}".\n'
        "Three JSON arrays follow:\n"
        "(1) Stories — Story / User Story issues.\n"
        "(2) Production-flagged — issues with production platform checked (may include bugs, tasks, etc.).\n"
        "(3) Bugs — all Bug-type issues from the same query (may overlap production list).\n\n"
        f"Stories:\n{sj}\n\nProduction-flagged:\n{pj}\n\nBugs:\n{bj}\n\n"
        "Return ONLY valid JSON (no markdown fences) with exactly these keys:\n"
        "{\n"
        '  "stories_bullets": ["Plain sentence about story work; cite ticket as (ABC-123) at end when referring to a specific story"],\n'
        '  "overall_bullets": ["2-5 bullets: executive rollup of STORY delivery only"],\n'
        '  "prod_bullets": ["Same style as stories_bullets: plain sentence, cite ticket as (ABC-123) at end when referring to a prod or bug issue"],\n'
        '  "story_keys_highlighted": ["KEY1"],\n'
        '  "prod_keys_highlighted": ["KEY2"]\n'
        "}\n"
        "Rules:\n"
        "- stories_bullets: 4-10 items when stories exist; else one bullet noting none. "
        "When a bullet ties to a specific story from the Stories JSON, end that bullet with (KEY) like prod_bullets.\n"
        "- overall_bullets: themes across stories only (not bugs).\n"
        "- prod_bullets: cover production-flagged work AND non-story bugs (dedupe overlaps). "
        "4-12 items when there is any prod or bug data; if both lists empty use one bullet stating none.\n"
        "- prod_bullets MUST cite Jira keys the same way as stories_bullets: put the issue key in parentheses at the end "
        'of the sentence or clause for each bullet that ties to a specific ticket (e.g. "... export edge case (TIM-27800)"). '
        "Include keys from the Production-flagged and Bugs JSON lists so readers can match summaries to tickets.\n"
        "- Each bullet one line, plain text, no markdown.\n"
        '- Do not start bullets with issue keys or \"KEY:\" — that produces stray colons when keys are linked at the end.\n'
        "- story_keys_highlighted: up to 10 story keys that matter.\n"
        "- prod_keys_highlighted: up to 12 keys mixing prod-flagged and bug keys that matter most."
    )
    try:
        res = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "x-api-key": anthropic_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model": model,
                "max_tokens": 2048,
                "messages": [{"role": "user", "content": prompt}],
            },
            timeout=120,
        )
        if not res.ok:
            return None, f"Claude HTTP {res.status_code}"
        body = res.json()
        raw = (body.get("content") or [{}])[0].get("text", "").strip()
        raw = raw.strip("`").strip()
        if raw.lower().startswith("json"):
            raw = raw[4:].lstrip()
        data = json.loads(raw)
        if not isinstance(data, dict):
            return None, "Invalid JSON shape"
        def _coerce_bullets(v):
            if isinstance(v, list):
                out = [_team_diagram_normalize_bullet_line(str(x).strip()) for x in v if x is not None and str(x).strip()]
                return [x for x in out if x][:14]
            if isinstance(v, str) and v.strip():
                one = _team_diagram_normalize_bullet_line(v.strip())
                return [one] if one else []
            return []

        out = {
            "stories_bullets": _coerce_bullets(data.get("stories_bullets")),
            "prod_bullets": _coerce_bullets(data.get("prod_bullets")),
            "overall_bullets": _coerce_bullets(data.get("overall_bullets")),
            "story_keys_highlighted": data.get("story_keys_highlighted") or [],
            "prod_keys_highlighted": data.get("prod_keys_highlighted") or [],
        }
        # Accept legacy prose-only responses
        if not out["stories_bullets"]:
            out["stories_bullets"] = _coerce_bullets(data.get("stories_summary"))
        if not out["prod_bullets"]:
            out["prod_bullets"] = _coerce_bullets(data.get("prod_issues_summary"))
        if not out["overall_bullets"]:
            out["overall_bullets"] = _coerce_bullets(data.get("overall_summary"))
        if not isinstance(out["story_keys_highlighted"], list):
            out["story_keys_highlighted"] = []
        if not isinstance(out["prod_keys_highlighted"], list):
            out["prod_keys_highlighted"] = []
        out["story_keys_highlighted"] = [str(x).strip() for x in out["story_keys_highlighted"][:10] if x]
        out["prod_keys_highlighted"] = [str(x).strip() for x in out["prod_keys_highlighted"][:10] if x]
        return out, None
    except json.JSONDecodeError as e:
        return None, f"JSON parse error: {e}"
    except Exception as e:
        return None, str(e)


def _team_diagram_resolve_ai(main_team, sub_team_name, jql, story_issues, prod_issues, bug_issues, ai_cache_only, anthropic_key, model):
    cache_key = _team_diagram_ai_cache_key(jql, sub_team_name)
    cached = _team_diagram_ai_cache_get(cache_key)
    if cached and isinstance(cached, dict):
        merged = _team_diagram_ai_normalize_client_blob(dict(cached))
        merged["from_cache"] = True
        merged["cache_key"] = cache_key
        return merged

    if ai_cache_only:
        return _team_diagram_ai_normalize_client_blob(
            {
                "stories_bullets": [],
                "prod_bullets": [],
                "overall_bullets": [],
                "story_keys_highlighted": [],
                "prod_keys_highlighted": [],
                "from_cache": False,
                "cache_miss": True,
                "cache_key": cache_key,
                "error": "No cached AI summary yet. Run Generate once without 'AI summaries: cache only'.",
            }
        )

    if not anthropic_key:
        return _team_diagram_ai_normalize_client_blob(
            {
                "stories_bullets": [],
                "prod_bullets": [],
                "overall_bullets": [],
                "story_keys_highlighted": [],
                "prod_keys_highlighted": [],
                "from_cache": False,
                "cache_key": cache_key,
                "error": "Anthropic API key not configured in Settings.",
            }
        )

    if not story_issues and not prod_issues and not bug_issues:
        empty = {
            "stories_bullets": ["No Story-type tickets returned by this JQL."],
            "prod_bullets": ["No production-flagged or Bug-type tickets in this result set."],
            "overall_bullets": ["No tickets in this result set to summarize."],
            "story_keys_highlighted": [],
            "prod_keys_highlighted": [],
            "from_cache": False,
            "cache_key": cache_key,
        }
        _team_diagram_ai_cache_put(cache_key, sub_team_name, _team_diagram_ai_pack_for_store(empty))
        return _team_diagram_ai_normalize_client_blob(empty)

    parsed, err = _team_diagram_claude_summarize(
        sub_team_name, main_team, story_issues, prod_issues, bug_issues, anthropic_key, model
    )
    if err or not parsed:
        return _team_diagram_ai_normalize_client_blob(
            {
                "stories_bullets": [],
                "prod_bullets": [],
                "overall_bullets": [],
                "story_keys_highlighted": [],
                "prod_keys_highlighted": [],
                "from_cache": False,
                "cache_key": cache_key,
                "error": err or "Claude failed",
            }
        )
    parsed["from_cache"] = False
    parsed["cache_key"] = cache_key
    _team_diagram_ai_cache_put(cache_key, sub_team_name, parsed)
    return _team_diagram_ai_normalize_client_blob(parsed)


@app.route("/team_diagram")
@admin_required
def team_diagram_page():
    return render_template("team_diagram.html")


@app.route("/api/team_diagram/fetch", methods=["POST"])
@admin_required
def team_diagram_fetch():
    """Fetch aggregated Jira metrics for each sub-team JQL."""
    payload = request.get_json(silent=True) or {}
    main_team = str(payload.get("main_team") or "").strip()
    sub_teams = payload.get("sub_teams") or []

    if not main_team:
        return jsonify({"error": "Main team name is required"}), 400
    if not sub_teams:
        return jsonify({"error": "At least one sub-team is required"}), 400

    headers_dict = dict(HEADERS)
    if "Authorization" not in headers_dict:
        return jsonify({"error": "Jira credentials not configured. Please set them in Settings."}), 401

    jira_domain = str(JIRA_DOMAIN)
    _plat_cf = _get_platform_checkboxes_field_id()
    fields = "summary,status,issuetype,labels,customfield_10016,customfield_10077"
    if _plat_cf:
        fields += f",{_plat_cf}"

    def _opt_nonneg_int(v):
        if v is None:
            return None
        try:
            return max(0, int(v))
        except (TypeError, ValueError):
            return None

    ai_model = str(payload.get("model") or "").strip() or "claude-sonnet-4-20250514"
    ai_cache_only = bool(payload.get("ai_cache_only"))
    anthropic_key_global = (_get_app_config_value("anthropic_api_key") or "").strip()

    results = []
    for st in sub_teams:
        name = str(st.get("name") or "").strip()
        jql = str(st.get("jql") or "").strip()
        if not name or not jql:
            continue

        all_issues = []
        seen_keys = set()
        start_at = 0
        next_page_token = None
        page_safety = 0
        error_msg = None

        while page_safety < 80:
            page_safety += 1
            params = {"jql": jql, "maxResults": 100, "fields": fields}
            if next_page_token:
                params["nextPageToken"] = next_page_token
            else:
                params["startAt"] = start_at
            try:
                jira_res = requests.get(
                    f"{jira_domain}/rest/api/3/search/jql",
                    headers=headers_dict,
                    params=params,
                    timeout=30,
                )
            except Exception as e:
                error_msg = str(e)
                break

            if jira_res.status_code != 200:
                error_msg = f"Jira API error ({jira_res.status_code}): {jira_res.text[:200]}"
                break

            data = jira_res.json()
            issues = data.get("issues", [])
            if not issues:
                break
            new_count = 0
            for issue in issues:
                key = issue.get("key")
                if not key or key in seen_keys:
                    continue
                seen_keys.add(key)
                all_issues.append(issue)
                new_count += 1
            if new_count == 0:
                break
            next_page_token = data.get("nextPageToken")
            if next_page_token:
                continue
            if len(issues) < 100:
                break
            start_at += 100

        total_stories = done_stories = total_bugs = done_bugs = prod_issues = 0
        done_total = 0
        story_tasks_total = story_tasks_done = 0
        prod_issues_done = 0
        total_test_cases = automated_test_cases = 0
        total_story_points = 0.0
        story_issues = []
        prod_issues_ai = []
        bug_issues_ai = []
        customer_names_local = set()
        for issue in all_issues:
            f = issue.get("fields") or {}
            itype_raw = (f.get("issuetype") or {}).get("name") or ""
            itype = itype_raw.strip().lower()
            status_obj = f.get("status") or {}
            status_disp = (status_obj.get("name") or "").strip()
            status_category = ((status_obj.get("statusCategory") or {}).get("key") or "").strip()
            is_done = _is_done_like_status(status_disp, status_category)
            _plat_raw = f.get(_plat_cf) if _plat_cf else None
            is_prod = _jira_platform_includes_production(_plat_raw)

            if is_done:
                done_total += 1

            story_like = itype in ("story", "task", "sub-task", "subtask")
            if story_like:
                story_tasks_total += 1
                if is_done:
                    story_tasks_done += 1

            if itype == "story":
                total_stories += 1
                if is_done:
                    done_stories += 1
            elif itype == "bug":
                total_bugs += 1
                if is_done:
                    done_bugs += 1
            if is_prod:
                prod_issues += 1
                if is_done:
                    prod_issues_done += 1
            if _is_test_case_issue_type(itype_raw):
                total_test_cases += 1
                if _jira_issue_is_automated_test(itype_raw, f):
                    automated_test_cases += 1

            total_story_points += _team_diagram_coerce_story_points(f.get("customfield_10016"))

            for cust in _extract_customer_values(f.get("customfield_10077")):
                customer_names_local.add(cust)

            ik = issue.get("key")
            sum_txt = (f.get("summary") or "").strip()
            if ik and _issue_type_bucket(itype_raw) == "Story":
                story_issues.append({"key": ik, "summary": sum_txt[:240], "status": status_disp})
            if ik and is_prod:
                prod_issues_ai.append(
                    {"key": ik, "summary": sum_txt[:240], "status": status_disp, "type": itype_raw}
                )
            if ik and _is_bug_type(itype_raw):
                bug_issues_ai.append(
                    {"key": ik, "summary": sum_txt[:240], "status": status_disp, "production": bool(is_prod)}
                )

        jira_ttc = total_test_cases

        prod_bug_rows = []
        _seen_pb = set()
        for row in prod_issues_ai:
            k = row.get("key")
            if k and k not in _seen_pb:
                _seen_pb.add(k)
                prod_bug_rows.append(
                    {
                        "key": k,
                        "summary": (row.get("summary") or "")[:280],
                        "status": (row.get("status") or ""),
                        "issue_type": (row.get("type") or ""),
                        "prod_platform": True,
                    }
                )
            if len(prod_bug_rows) >= 40:
                break
        if len(prod_bug_rows) < 40:
            for row in bug_issues_ai:
                k = row.get("key")
                if k and k not in _seen_pb:
                    _seen_pb.add(k)
                    prod_bug_rows.append(
                        {
                            "key": k,
                            "summary": (row.get("summary") or "")[:280],
                            "status": (row.get("status") or ""),
                            "issue_type": "Bug",
                            "prod_platform": bool(row.get("production")),
                        }
                    )
                if len(prod_bug_rows) >= 40:
                    break

        prod_bug_ticket_keys_ordered = [r["key"] for r in prod_bug_rows]

        team_member_count = _opt_nonneg_int(st.get("team_member_count"))
        total_story_points_r = round(total_story_points, 2)
        estimated_calendar_days = None
        if team_member_count is not None and team_member_count > 0:
            estimated_calendar_days = round(total_story_points / float(team_member_count), 2)

        manual_total_i = _opt_nonneg_int(st.get("test_cases_total"))
        manual_auto_i = _opt_nonneg_int(st.get("test_cases_automated"))
        use_manual_tests = manual_total_i is not None or manual_auto_i is not None
        if use_manual_tests:
            if manual_total_i is not None:
                total_test_cases = manual_total_i
            else:
                total_test_cases = max(jira_ttc, manual_auto_i or 0)
            if manual_auto_i is not None:
                automated_test_cases = manual_auto_i
            else:
                automated_test_cases = 0
            automated_test_cases = min(automated_test_cases, total_test_cases)

        test_auto_pct = (
            round(automated_test_cases / total_test_cases * 100) if total_test_cases else 0
        )

        ai_block = _team_diagram_resolve_ai(
            main_team,
            name,
            jql,
            story_issues,
            prod_issues_ai,
            bug_issues_ai,
            ai_cache_only,
            anthropic_key_global,
            ai_model,
        )

        results.append({
            "name": name,
            "jql": jql,
            "total": len(all_issues),
            "done_total": done_total,
            "story_tasks_total": story_tasks_total,
            "story_tasks_done": story_tasks_done,
            "team_member_count": team_member_count,
            "total_story_points": total_story_points_r,
            "estimated_calendar_days": estimated_calendar_days,
            "total_stories": total_stories,
            "done_stories": done_stories,
            "total_bugs": total_bugs,
            "done_bugs": done_bugs,
            "prod_issues": prod_issues,
            "prod_issues_done": prod_issues_done,
            "total_test_cases": total_test_cases,
            "automated_test_cases": automated_test_cases,
            "test_automation_pct": test_auto_pct,
            "test_metrics_manual": use_manual_tests,
            "story_ticket_keys": [x["key"] for x in story_issues[:24]],
            "prod_ticket_keys": [x["key"] for x in prod_issues_ai[:24]],
            "bug_ticket_keys": [x["key"] for x in bug_issues_ai[:28]],
            "prod_bug_ticket_keys": prod_bug_ticket_keys_ordered[:32],
            "prod_bug_rows": prod_bug_rows[:40],
            "customer_names": sorted(customer_names_local, key=lambda x: str(x).lower()),
            "ai": ai_block,
            "error": error_msg,
        })

    portfolio_customers = set()
    for row in results:
        for c in row.get("customer_names") or []:
            portfolio_customers.add(c)
    portfolio_customer_names = sorted(portfolio_customers, key=lambda x: str(x).lower())

    return jsonify(
        {
            "main_team": main_team,
            "sub_teams": results,
            "portfolio_customer_names": portfolio_customer_names,
        }
    )


if __name__ == "__main__":
    app.run(debug=True)
